#!/usr/bin/env python3
"""Build the BK24 synthetic corporate-law demonstration document pack."""

from __future__ import annotations

import json
import os
import re
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from demo_scenario_library import ADDITIONAL_SCENARIOS

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "demo" / "cybersafeadvisor-corporate-pack"
INK = RGBColor(31, 45, 61)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(96, 105, 115)
DISCLAIMER = "SYNTHETIC DEMO - NOT LEGAL ADVICE"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

SOURCES = {
    "sec_msa": (
        "SEC EDGAR material-contract exhibits (structure research only): "
        "https://www.sec.gov/Archives/edgar/data/2026767/000149315226029398/ex10-12.htm"
    ),
    "nist": (
        "NIST Cybersecurity Framework 2.0: "
        "https://www.nist.gov/publications/nist-cybersecurity-framework-csf-20"
    ),
    "delaware": (
        "Delaware General Corporation Law, 8 Del. C. Section 141: "
        "https://delcode.delaware.gov/title8/c001/sc04/index.html#141"
    ),
    "aba_fees": (
        "ABA Model Rule 1.5 (fees): "
        "https://www.americanbar.org/groups/professional_responsibility/"
        "publications/model_rules_of_professional_conduct/rule_1_5_fees/"
    ),
    "synthetic_design": (
        "Scenario design note: all people, organizations, facts, and amounts are fictional; "
        "the material is for product demonstration only and is not legal advice."
    ),
}


LEGACY_SCENARIOS = [
    {
        "external_key": "demo-northstar-saas-review",
        "name": "Northstar Analytics - SaaS Vendor Review",
        "client": "Northstar Analytics, Inc. (fictional)",
        "lead": "A customer-side SaaS renewal with AI-training, security, and liability issues.",
        "demo_prompt": "Review the MSA and data addendum for customer-side risks, cite the source clauses, and propose follow-up tasks for the renewal and data-use issues.",
        "suggested_tasks": [
            "Negotiate AI training restriction",
            "Calendar renewal notice",
            "Resolve security liability carve-out",
        ],
    },
    {
        "external_key": "demo-harborlight-seed-financing",
        "name": "HarborLight Robotics - Series Seed Financing",
        "client": "HarborLight Robotics, Inc. (fictional)",
        "lead": "A founder-friendly seed financing with approvals, cap-table, and closing work.",
        "demo_prompt": "Compare the term sheet to the board consent, identify missing approvals and closing conditions, and prepare a prioritized closing checklist.",
        "suggested_tasks": [
            "Verify approval thresholds",
            "Prepare stockholder consent",
            "Reconcile capitalization certificate",
        ],
    },
    {
        "external_key": "demo-redwood-ogc-retainer",
        "name": "Redwood Outdoor Supply - Outside General Counsel Retainer",
        "client": "Redwood Outdoor Supply, Inc. (fictional)",
        "lead": "A lively monthly outside-counsel portfolio with contract, product, and governance work.",
        "demo_prompt": "Triage the August legal requests, show the next notice deadlines, and draft the monthly retainer portfolio summary with decisions needed.",
        "suggested_tasks": [
            "Review PinePeak exclusivity",
            "Calendar Evergreen notice",
            "Prepare September board materials",
        ],
    },
]

ALL_SCENARIOS = [*LEGACY_SCENARIOS, *ADDITIONAL_SCENARIOS]


def _client_profile(index: int, scenario: dict) -> dict:
    """Return fictional contact details that make each demo matter feel lived in."""

    contacts = [
        ("Avery", "Nguyen", "Chief Operating Officer"),
        ("Morgan", "Ellis", "Founder and CEO"),
        ("Riley", "Santos", "General Manager"),
        ("Jamie", "Patel", "Director of Operations"),
        ("Casey", "Morgan", "Owner"),
        ("Taylor", "Brooks", "Chief of Staff"),
        ("Jordan", "Kim", "Vice President of Product"),
        ("Robin", "Flores", "Founder"),
        ("Cameron", "Wells", "AI Governance Program Manager"),
        ("Drew", "Hart", "Director of Brand Compliance"),
        ("Avery", "Riverstone", "Client"),
        ("Jordan", "Parker", "Client"),
        ("Emerson", "Lane", "Managing Member"),
        ("Skyler", "Bennett", "Program Director"),
        ("Alex", "Monroe", "Mediation Participant"),
        ("Maya", "Chen", "Client"),
        ("Nico", "Alvarez", "Managing Member"),
        ("Priya", "Raman", "Chief Executive Officer"),
    ]
    secondary_contacts = [
        ("Dana", "Price", "Finance and Billing"),
        ("Elliot", "Shaw", "Board Secretary"),
        ("Sam", "Ortiz", "Operations Coordinator"),
        ("Lee", "Wallace", "Risk Manager"),
        ("Ari", "Foster", "Office Manager"),
        ("Jules", "Grant", "HR Business Partner"),
        ("Mina", "Park", "Product Counsel Liaison"),
        ("Noah", "Bishop", "Marketing Director"),
        ("Remy", "Cole", "People Operations"),
        ("Asha", "Bell", "Quality Manager"),
        ("Micah", "Stone", "Family Coordinator"),
        ("Tess", "Jordan", "Emergency Contact"),
        ("Ira", "Sloan", "Bookkeeper"),
        ("Lena", "Marsh", "Estate Accountant"),
        ("Owen", "Dale", "Mediation Coordinator"),
        ("Sofia", "Lin", "Records Contact"),
        ("Marco", "Diaz", "Restaurant Manager"),
        ("Elena", "Voss", "Finance Director"),
    ]
    counterparties = [
        ("BlueHarbor Cloud Systems, Inc. (fictional)", "Blake Carter"),
        ("Horizon Seed Ventures, LP (fictional)", "Rowan Carter"),
        ("PinePeak Distribution, LLC (fictional)", "Hayden Carter"),
        ("Riverside Roof Gardens, LLC (fictional)", "Sage Carter"),
        ("FreshForm Nutrition Partners, Inc. (fictional)", "Reese Carter"),
        ("Morgan Hale (fictional)", "Morgan Hale"),
        ("CampTrail Retail Group, Inc. (fictional)", "Avery Dawson"),
        ("Sparkle Crumbs Confections, LLC (fictional)", "Devon Price"),
        ("RecruitSight AI, Inc. (fictional)", "Marlowe Grant"),
        ("MarketSquare Grocers, Inc. (fictional)", "Carmen Bell"),
        ("Morgan Riverstone Household (fictional)", "Morgan Riverstone"),
        ("Lakeview Municipal Prosecutor (fictional)", "Taylor Brooks"),
        ("Harbor District Properties, LLC (fictional)", "Rene Patel"),
        ("Marigold Family Beneficiaries (fictional)", "Jamie Marigold"),
        ("Bell Household (fictional)", "Casey Bell"),
        ("QuickCrate Delivery, Inc. (fictional)", "Dana Walsh"),
        ("Greenline Municipal Licensing Bureau (fictional)", "Elliot Reed"),
        ("Cameron Vale (fictional)", "Cameron Vale"),
    ]
    streets = [
        ("184 Lantern Way", "Chicago", "IL", "60607"),
        ("42 Harbor Loop", "Wilmington", "DE", "19801"),
        ("77 Trailhead Avenue", "Columbus", "OH", "43215"),
        ("310 Juniper Street", "Cleveland", "OH", "44113"),
        ("18 Orchard Lane", "Ann Arbor", "MI", "48104"),
    ]
    first, last, title = contacts[index % len(contacts)]
    secondary_first, secondary_last, secondary_title = secondary_contacts[
        index % len(secondary_contacts)
    ]
    street, city, state, postal_code = streets[index % len(streets)]
    local = f"{first}.{last}".lower()
    secondary_local = f"{secondary_first}.{secondary_last}".lower()
    counterparty_organization, counterparty_contact = counterparties[
        index % len(counterparties)
    ]
    return {
        "organization": scenario["client"],
        "address": {
            "street": street,
            "city": city,
            "state": state,
            "zip": postal_code,
            "country": "US",
        },
        "client_since": f"20{19 + (index % 7)}-{(index % 9) + 1:02d}-15",
        "preferred_contact_method": "email" if index % 3 else "phone",
        "preferred_contact_window": [
            "Weekdays, 9:00 a.m.-noon Central",
            "Weekdays, 1:00-4:00 p.m. Eastern",
            "Weekdays after 3:00 p.m. local time",
        ][index % 3],
        "billing_terms": "Monthly itemized invoice; electronic delivery; client approval required before third-party spend.",
        "primary_contact": {
            "first_name": first,
            "last_name": last,
            "title": title,
            "email": f"{local}@example.invalid",
            "phone": f"+1-202-555-01{index + 10:02d}",
        },
        "secondary_contact": {
            "first_name": secondary_first,
            "last_name": secondary_last,
            "title": secondary_title,
            "email": f"{secondary_local}@example.invalid",
            "phone": f"+1-202-555-02{index + 10:02d}",
        },
        "opposing_party": {
            "organization": counterparty_organization,
            "contact_name": counterparty_contact,
            "email": f"counterparty-{index + 1}@example.invalid",
        },
    }


DOCUMENTS = [
    {
        "filename": "01-northstar-cloud-master-services-agreement.docx",
        "title": "Master Services Agreement",
        "subtitle": "BlueHarbor Cloud Systems, Inc. and Northstar Analytics, Inc.",
        "matter": "Northstar Analytics - SaaS Vendor Review",
        "date": "August 1, 2026",
        "lead": (
            "This synthetic agreement intentionally contains negotiation issues for an AI-assisted "
            "customer-side review: a narrow liability cap, broad provider data-use rights, automatic "
            "renewal, an asymmetric indemnity, and a short suspension right."
        ),
        "sections": [
            (
                "Services and Order Forms",
                [
                    "BlueHarbor will provide the hosted analytics platform described in each signed Order Form. An Order Form may add services, usage limits, support levels, and fees, but it may not amend this Agreement unless it identifies the provision being amended and is signed by authorized representatives of both parties.",
                    "BlueHarbor may update non-material platform features during the term. A change that materially reduces core functionality requires at least 30 days' notice and gives Northstar a termination right for the affected service if the reduction is not cured within 15 days after written notice.",
                ],
            ),
            (
                "Fees, Invoicing, and Taxes",
                [
                    "Northstar will pay undisputed invoices within 30 days. Disputed amounts must be identified with reasonable detail within 15 days after receipt, and the parties will work in good faith to resolve them. Overdue undisputed amounts accrue interest at 1.5% per month or the maximum lawful rate, whichever is lower.",
                    "Fees increase by 8% at each renewal unless the applicable Order Form states otherwise. Northstar is responsible for sales and use taxes other than taxes measured by BlueHarbor's net income.",
                ],
            ),
            (
                "Customer Data and Product Improvement",
                [
                    "Northstar owns Customer Data. Northstar grants BlueHarbor a worldwide license during the term to host, copy, transmit, and otherwise process Customer Data to provide and secure the services.",
                    "BlueHarbor may also use Customer Data, including prompts, uploaded documents, and resulting usage signals, to train, tune, evaluate, and improve BlueHarbor and third-party machine-learning models, provided BlueHarbor removes direct identifiers before that use. This product-improvement license survives termination in perpetuity.",
                ],
            ),
            (
                "Confidentiality",
                [
                    "Each recipient will use the other party's Confidential Information only to perform or receive the services, protect it using at least reasonable care, and disclose it only to personnel and contractors bound by confidentiality duties. Standard exclusions apply for information independently developed, rightfully received, public without breach, or previously known without restriction.",
                    "A compelled disclosure is permitted after prompt notice when legally allowed and reasonable cooperation at the disclosing party's expense.",
                ],
            ),
            (
                "Security and Incident Response",
                [
                    "BlueHarbor will maintain a written security program aligned in material respects with the NIST Cybersecurity Framework and the Data and Security Addendum. BlueHarbor will notify Northstar without undue delay and no later than 72 hours after confirming unauthorized access to unencrypted Customer Data.",
                    "BlueHarbor may suspend access immediately when it reasonably believes continued use creates a security risk, violates law, or threatens the platform. BlueHarbor is not required to provide advance notice and has no express obligation to limit a suspension to the affected account or function.",
                ],
            ),
            (
                "Intellectual Property",
                [
                    "BlueHarbor owns the platform, documentation, usage analytics, generalized know-how, and all improvements. Northstar owns Customer Data and Northstar materials. Feedback is assigned to BlueHarbor without payment or restriction.",
                ],
            ),
            (
                "Warranties",
                [
                    "BlueHarbor warrants that the service will materially conform to the documentation and that professional services will be performed in a professional manner. Northstar's exclusive remedy is re-performance or, if BlueHarbor cannot cure, termination of the affected Order Form and refund of prepaid unused fees.",
                    "Except for the express warranties above, the service is provided as is and BlueHarbor disclaims implied warranties to the maximum extent permitted by law.",
                ],
            ),
            (
                "Indemnification",
                [
                    "Northstar will defend and indemnify BlueHarbor from third-party claims arising from Customer Data, Northstar's products, or use of the service in violation of this Agreement.",
                    "BlueHarbor will defend Northstar only against a third-party claim that the unmodified service directly infringes a United States patent, copyright, or trademark. BlueHarbor has no express indemnity obligation for a security incident, privacy claim, or BlueHarbor's violation of law.",
                ],
            ),
            (
                "Limitation of Liability",
                [
                    "Neither party is liable for indirect, incidental, special, consequential, exemplary, or punitive damages. BlueHarbor's aggregate liability arising from this Agreement will not exceed fees paid for the affected service during the three months preceding the event giving rise to liability.",
                    "Northstar's payment obligations, indemnity obligations, and breach of BlueHarbor's intellectual-property rights are uncapped. There is no separate cap for confidentiality, security, privacy, gross negligence, or willful misconduct.",
                ],
            ),
            (
                "Term, Renewal, and Termination",
                [
                    "The initial term ends July 31, 2027 and automatically renews for successive one-year terms unless either party gives at least 60 days' written notice before the current term ends. Either party may terminate for an uncured material breach after 30 days' notice, reduced to 10 days for nonpayment.",
                    "On termination, Northstar may export Customer Data for 15 days. BlueHarbor may then delete it under its standard retention practices. Prepaid fees are nonrefundable except for BlueHarbor's uncured breach.",
                ],
            ),
            (
                "General",
                [
                    "Delaware law governs without regard to conflicts principles, and state or federal courts in New Castle County, Delaware have exclusive jurisdiction. Neither party may assign without consent, except BlueHarbor may assign to an affiliate or in connection with a merger, financing, or sale of substantially all relevant assets. Notices must be sent by nationally recognized overnight courier and email to the contacts in the Order Form.",
                ],
            ),
        ],
        "sources": ["sec_msa", "nist"],
    },
    {
        "filename": "02-northstar-cloud-data-security-addendum.docx",
        "title": "Data and Security Addendum",
        "subtitle": "Attachment to the BlueHarbor / Northstar Master Services Agreement",
        "matter": "Northstar Analytics - SaaS Vendor Review",
        "date": "August 1, 2026",
        "lead": (
            "This synthetic addendum creates a realistic conflict with the MSA's broad product-"
            "improvement license and includes gaps around subprocessors, deletion evidence, and audit."
        ),
        "sections": [
            (
                "Scope and Roles",
                [
                    "This Addendum applies when BlueHarbor processes Personal Data in Customer Data for Northstar. Northstar acts as controller or business, and BlueHarbor acts as processor or service provider, except when law requires otherwise.",
                ],
            ),
            (
                "Processing Instructions",
                [
                    "BlueHarbor will process Personal Data only to provide, secure, support, and improve the services as documented by Northstar's use and this Addendum. BlueHarbor will notify Northstar if an instruction appears unlawful.",
                    "The parties acknowledge that service improvement may include automated evaluation of prompts and outputs. This Addendum does not expressly prohibit using de-identified Customer Data to train generalized models, and it does not define a de-identification standard.",
                ],
            ),
            (
                "Security Program",
                [
                    "BlueHarbor will maintain administrative, technical, and physical safeguards reasonably designed around the Govern, Identify, Protect, Detect, Respond, and Recover functions of the NIST Cybersecurity Framework 2.0.",
                    "Minimum measures include multifactor authentication for privileged access, encryption in transit and at rest, logging of administrative access, annual penetration testing, secure development practices, vulnerability management, employee confidentiality obligations, and tested incident-response and business-continuity plans.",
                ],
            ),
            (
                "Security Incidents",
                [
                    "BlueHarbor will notify Northstar without undue delay and no later than 72 hours after confirming a Security Incident affecting Personal Data. Notice will include known scope, data categories, likely consequences, mitigation, and a response contact, with supplemental information as investigation continues.",
                    "BlueHarbor is not required to notify Northstar of unsuccessful attempts that do not compromise Personal Data. BlueHarbor controls external communications unless applicable law requires Northstar to communicate.",
                ],
            ),
            (
                "Subprocessors",
                [
                    "Northstar gives general authorization for subprocessors. BlueHarbor will maintain an online subprocessor list and provide 10 days' notice of a new subprocessor. Northstar may object on reasonable data-protection grounds, but its sole remedy is termination of the affected service without refund.",
                ],
            ),
            (
                "Individual Requests and Government Demands",
                [
                    "Taking account of the nature of processing, BlueHarbor will provide reasonable assistance with verified individual-rights requests. BlueHarbor may charge its then-current professional-services rates when assistance requires more than four hours in a month.",
                    "Unless prohibited, BlueHarbor will notify Northstar of a government demand for Personal Data and reasonably redirect the authority to Northstar.",
                ],
            ),
            (
                "Return and Deletion",
                [
                    "At termination, BlueHarbor will make Customer Data available for export for 15 days and then delete active copies within 90 days. Encrypted backup copies may remain until overwritten in the ordinary course, not to exceed 18 months. No deletion certificate or backup isolation evidence is required.",
                ],
            ),
            (
                "Audit and Compliance",
                [
                    "Once per year, BlueHarbor will provide its then-current SOC 2 Type II report under confidentiality. Northstar has no onsite audit right unless a regulator requires one after a confirmed Security Incident, and any audit must avoid access to other customers' information.",
                ],
            ),
            (
                "Order of Precedence",
                [
                    "This Addendum controls over the MSA only for a direct conflict concerning Personal Data protection. The MSA controls all other issues, including liability caps, remedies, and product-improvement rights.",
                ],
            ),
        ],
        "sources": ["nist", "sec_msa"],
    },
    {
        "filename": "03-harborlight-series-seed-term-sheet.docx",
        "title": "Series Seed Preferred Stock Term Sheet",
        "subtitle": "HarborLight Robotics, Inc. - Proposed $4.5 Million Financing",
        "matter": "HarborLight Robotics - Series Seed Financing",
        "date": "August 8, 2026",
        "lead": (
            "This non-binding synthetic term sheet is designed for issue extraction, approval mapping, "
            "and closing-task generation. Dollar amounts and parties are entirely fictional."
        ),
        "sections": [
            (
                "Offering",
                [
                    "HarborLight Robotics, Inc., a Delaware corporation, will issue Series Seed Preferred Stock for aggregate gross proceeds of $4,500,000 at $2.25 per share. Crescent Ventures I, L.P. will invest $3,000,000 as lead investor; other accredited investors will purchase the balance.",
                    "The pre-money valuation is $13,500,000 on a fully diluted basis, including an unallocated employee option pool equal to 12% of post-closing capitalization. The company will bear up to $40,000 of the lead investor's reasonable legal fees at closing.",
                ],
            ),
            (
                "Liquidation Preference",
                [
                    "The Series Seed will carry a one-times non-participating liquidation preference senior to common stock, plus declared but unpaid dividends. Holders may instead convert to common stock if that produces a greater return.",
                ],
            ),
            (
                "Dividends and Conversion",
                [
                    "Non-cumulative dividends will be payable when and if declared. Series Seed will convert voluntarily at the holder's option and automatically upon a qualified public offering or approval of holders of at least 60% of the Series Seed.",
                ],
            ),
            (
                "Protective Provisions",
                [
                    "Approval of at least 60% of the Series Seed is required to amend the charter adversely, authorize senior or pari passu securities, redeem shares, declare dividends, incur debt above $1,000,000 outside the approved budget, sell substantially all assets, change the board size, or enter a related-party transaction above $100,000.",
                ],
            ),
            (
                "Board Composition",
                [
                    "At closing the board will have five seats: two designated by common holders, one designated by Crescent Ventures, the chief executive officer, and one independent director mutually approved by the company and Crescent Ventures within 90 days after closing.",
                ],
            ),
            (
                "Investor Rights",
                [
                    "Major Investors purchasing at least $500,000 will receive customary quarterly and annual information rights, inspection rights on reasonable notice, pro rata participation rights, and registration rights. Information rights are subject to privilege and competitively sensitive-information protections to be negotiated.",
                ],
            ),
            (
                "Founder Matters",
                [
                    "Each founder will sign a proprietary-information and inventions agreement. Fifty percent of each founder's currently vested common shares will become subject to monthly reverse vesting over 24 months, with double-trigger acceleration of 25% following a qualifying termination after a change of control.",
                ],
            ),
            (
                "Conditions to Closing",
                [
                    "Closing is targeted for September 15, 2026 and is conditioned on satisfactory diligence, definitive documents, board and stockholder approvals, charter filing, securities-law compliance, a capitalization certificate, invention assignments, key-person employment arrangements, and no material adverse change.",
                ],
            ),
            (
                "Exclusivity and Confidentiality",
                [
                    "For 45 days after signing, the company will not solicit or negotiate another equity financing or sale transaction. The term sheet and negotiations are confidential, except disclosures to representatives and as required by law.",
                ],
            ),
            (
                "Effect",
                [
                    "Only exclusivity, confidentiality, expenses, governing law, and this Effect section are binding. Delaware law governs. Definitive agreements and required corporate approvals control the financing.",
                ],
            ),
        ],
        "sources": ["delaware"],
    },
    {
        "filename": "04-harborlight-board-written-consent.docx",
        "title": "Unanimous Written Consent of the Board of Directors",
        "subtitle": "HarborLight Robotics, Inc. - Series Seed Financing Approvals",
        "matter": "HarborLight Robotics - Series Seed Financing",
        "date": "August 20, 2026",
        "lead": (
            "This synthetic consent assumes the corporation's charter and bylaws do not restrict "
            "unanimous board action by written consent. Counsel must verify governing documents and "
            "the final financing instruments before use."
        ),
        "sections": [
            (
                "Background",
                [
                    "The directors reviewed the August 8, 2026 Series Seed Preferred Stock Term Sheet, a draft amended and restated certificate of incorporation, a stock purchase agreement, investor-rights documents, capitalization materials, and management's financing presentation.",
                ],
            ),
            (
                "Financing Approval",
                [
                    "RESOLVED, that the proposed issuance and sale of up to 2,000,000 shares of Series Seed Preferred Stock for aggregate gross proceeds of up to $4,500,000 is approved, subject to final terms not materially less favorable to the corporation than those presented to the board.",
                ],
            ),
            (
                "Charter and Stockholder Approval",
                [
                    "RESOLVED, that the amended and restated certificate of incorporation is approved and recommended to the stockholders for adoption; and that the officers are authorized to solicit the stockholder consent required by the existing certificate, bylaws, and applicable law before filing.",
                ],
            ),
            (
                "Definitive Agreements",
                [
                    "RESOLVED, that the forms of stock purchase agreement, investors' rights agreement, voting agreement, right of first refusal and co-sale agreement, and related instruments are approved, with changes an authorized officer determines necessary or advisable, provided material economic changes return to the board.",
                ],
            ),
            (
                "Option Pool and Capitalization",
                [
                    "RESOLVED, that management is authorized to reserve the shares necessary for a 12% post-closing unallocated option pool, subject to verification of the capitalization schedule and the board's separate approval of individual equity grants.",
                ],
            ),
            (
                "Officer Authority",
                [
                    "RESOLVED, that the chief executive officer and chief financial officer, acting individually, may execute approved documents, make required filings, pay transaction expenses, and take actions reasonably necessary to close the financing; provided that investor counsel fees paid by the corporation may not exceed $40,000 without further approval.",
                ],
            ),
            (
                "Conflicts and Recordkeeping",
                [
                    "RESOLVED, that each director has disclosed any material interest known to that director in the financing; and that this consent and the final documents will be filed with the minutes maintained by the corporation.",
                ],
            ),
            (
                "Effectiveness and Counterparts",
                [
                    "This consent is effective when all directors have signed and may be executed in counterparts and by electronic transmission. If a director's signature is conditioned or delayed, counsel must confirm effectiveness before any filing or closing action.",
                ],
            ),
        ],
        "signatories": [
            "Avery Chen, Director",
            "Jordan Ellis, Director",
            "Morgan Reyes, Director",
        ],
        "sources": ["delaware"],
    },
    {
        "filename": "05-redwood-ogc-engagement-letter.docx",
        "title": "Outside General Counsel Engagement Letter",
        "subtitle": "Summit Corporate Counsel, PLLC and Redwood Outdoor Supply, Inc.",
        "matter": "Redwood Outdoor Supply - Outside General Counsel Retainer",
        "date": "August 1, 2026",
        "lead": (
            "This synthetic engagement demonstrates a recurring corporate retainer with scoped work, "
            "monthly reporting, excluded matters, renewal review, and approval rules."
        ),
        "sections": [
            (
                "Client and Scope",
                [
                    "Summit Corporate Counsel, PLLC will represent Redwood Outdoor Supply, Inc. as outside general corporate counsel. The engagement includes routine contract review and negotiation, corporate governance support, commercial-policy advice, monthly legal-portfolio reporting, and coordination with specialized counsel.",
                    "The firm does not represent Redwood's affiliates, stockholders, officers, directors, employees, or investors unless a separate writing expressly says so.",
                ],
            ),
            (
                "Monthly Retainer",
                [
                    "Redwood will pay a monthly retainer of $12,500, invoiced on the first business day of each month and due within 15 days. The retainer covers up to 45 attorney hours and 10 paralegal hours per month. Unused hours do not roll over.",
                    "Work beyond the included hours requires written approval from Redwood's chief financial officer and is billed at $425 per attorney hour and $185 per paralegal hour. The firm will provide notice when monthly usage reaches 80% of included hours.",
                ],
            ),
            (
                "Excluded and Separately Scoped Work",
                [
                    "Litigation, tax opinions, patent prosecution, public-company reporting, regulated-industry licensing, merger or financing transactions above $1,000,000, and investigations are excluded. The firm will not begin excluded work without a written scope and fee arrangement.",
                ],
            ),
            (
                "Working Cadence",
                [
                    "The parties will hold a 30-minute legal operations call each Tuesday. Redwood will send new requests to legal@redwood-outdoor.example with the business owner, desired date, counterparty, value, and risk context. The firm will acknowledge urgent requests within four business hours and routine requests within one business day.",
                    "By the fifth business day of each month, the firm will deliver a portfolio summary listing open contracts, decisions needed, approaching renewals, notice dates, governance items, retainer utilization, and recommended priorities.",
                ],
            ),
            (
                "Client Decisions and Approval",
                [
                    "Redwood retains authority over business and settlement decisions. The firm may not send a final agreement, consent, notice, filing, or legal conclusion to a third party without approval from an authorized Redwood contact, except for ministerial communications expressly delegated in writing.",
                ],
            ),
            (
                "Confidentiality and Technology",
                [
                    "The firm will protect confidential information under applicable professional duties and maintain reasonable administrative, technical, and physical safeguards. The firm may use vetted technology vendors to provide document management, research, and drafting assistance, but attorneys remain responsible for review and client confidentiality.",
                ],
            ),
            (
                "Conflicts",
                [
                    "The firm represents other businesses, including companies in outdoor products and retail. The firm will not represent another client directly adverse to Redwood in the same or a substantially related matter without informed written consent where permitted.",
                ],
            ),
            (
                "Term and Review",
                [
                    "The engagement begins August 1, 2026 and continues month to month. Either party may terminate on 30 days' written notice, subject to professional obligations. The parties will review scope, utilization, and pricing by July 1, 2027; any change must be in a signed writing.",
                ],
            ),
            (
                "File Retention and Governing Rules",
                [
                    "At conclusion, the firm will return client property on request and retain the file under its written retention policy and applicable professional obligations. This letter is governed by applicable law and professional-conduct rules; it does not waive any non-waivable client right.",
                ],
            ),
        ],
        "signatories": [
            "Taylor Morgan, Managing Attorney",
            "Riley Bennett, Chief Financial Officer",
        ],
        "sources": ["aba_fees"],
    },
    {
        "filename": "06-redwood-monthly-request-and-contract-calendar.docx",
        "title": "Monthly Legal Request and Contract Calendar",
        "subtitle": "Redwood Outdoor Supply, Inc. - August 2026",
        "matter": "Redwood Outdoor Supply - Outside General Counsel Retainer",
        "date": "August 12, 2026",
        "lead": (
            "This synthetic client update is designed for recurring triage, task creation, renewal "
            "tracking, and a monthly outside-counsel portfolio summary."
        ),
        "sections": [
            (
                "New Requests",
                [
                    "PinePeak Distribution Amendment - Business owner: Sales. Counterparty requests exclusive Midwest distribution, a three-year term, and a $2.2 million annual minimum. Desired signature date: August 28. Flag exclusivity, forecast remedies, termination rights, and channel conflict.",
                    "TrailForge Influencer Agreement - Business owner: Marketing. Review content approval, FTC disclosure responsibility, morality clause, name/image rights, and usage period. Campaign launches September 10. Budget: $85,000.",
                    "Warehouse Automation Pilot - Business owner: Operations. Vendor requests access to inventory and employee workflow data. Pilot starts October 1. Confirm data restrictions, security review, equipment risk, insurance, and pilot-to-production conversion terms.",
                ],
            ),
            (
                "Contract Calendar",
                [
                    "September 1, 2026 - RidgeLine Logistics annual rate notice due. Owner: Operations. Action: compare proposed 6.5% increase with contract index and service credits.",
                    "September 15, 2026 - Evergreen Packaging non-renewal notice deadline. Owner: Procurement. Contract renews December 14 for one year unless 90 days' notice is given.",
                    "October 5, 2026 - Summit Retail private-label quality certification expires. Owner: Quality. Action: obtain renewed certificate and confirm recall insurance.",
                    "November 2, 2026 - Canyon ERP renewal decision. Owner: IT. Auto-renews for three years unless notice is sent. Current annual fee: $420,000; vendor proposes an 11% increase.",
                    "July 1, 2027 - Outside general counsel annual scope and pricing review. Owners: CFO and General Counsel liaison.",
                ],
            ),
            (
                "Governance and Policy Items",
                [
                    "Prepare third-quarter board materials by September 22, including contract concentration, product recall readiness, and the warehouse automation pilot. Update delegation-of-authority thresholds before the holiday purchasing cycle.",
                    "Confirm whether the company should adopt an AI acceptable-use policy before marketing and procurement teams begin a generative-content pilot.",
                ],
            ),
            (
                "Retainer Status",
                [
                    "August usage through August 11: 21.4 attorney hours and 3.0 paralegal hours. Forecast: 43 attorney hours if current requests remain in scope. No overage is approved. The PinePeak amendment is the highest-priority matter for the next Tuesday call.",
                ],
            ),
            (
                "Requested Decisions",
                [
                    "CFO approval is needed before negotiating any PinePeak exclusivity longer than 12 months. Marketing must select a maximum content-usage period for TrailForge. Operations must confirm whether employee-level workflow data is necessary for the warehouse pilot.",
                ],
            ),
        ],
        "sources": ["aba_fees", "nist"],
    },
]


def _scenario_document(scenario: dict) -> dict:
    """Turn compact scenario data into a structured, demo-ready source document."""

    return {
        "filename": scenario["filename"],
        "title": scenario["title"],
        "subtitle": scenario["subtitle"],
        "matter": scenario["name"],
        "date": "August 18, 2026",
        "lead": scenario["lead"],
        "sections": [
            ("Background and Working Facts", scenario["facts"]),
            ("Risk and Issue Map", scenario["issues"]),
            ("Options and Decisions Needed", scenario["decisions"]),
            (
                "90-Day Action Plan",
                [
                    f"Priority {index + 1}: {task}."
                    for index, task in enumerate(scenario["suggested_tasks"])
                ],
            ),
            (
                "Guided Demo Question",
                [
                    scenario["demo_prompt"],
                    "Use this scenario to demonstrate document-grounded analysis, a matter timeline, "
                    "task creation, and an attorney review workflow. Confirm jurisdiction-specific "
                    "law and facts before acting outside this fictional demonstration.",
                ],
            ),
        ],
        "sources": ["synthetic_design"],
    }


DOCUMENTS.extend(_scenario_document(scenario) for scenario in ADDITIONAL_SCENARIOS)


def _supporting_document_filenames(external_key: str) -> list[str]:
    stem = external_key.removeprefix("demo-")
    return [
        f"support-{stem}-intake-and-contact-profile.docx",
        f"support-{stem}-communications-and-work-plan.docx",
        f"support-{stem}-chronology-and-document-index.docx",
    ]


def _supporting_documents(index: int, scenario: dict) -> list[dict]:
    """Create the client-facing and internal work papers for every demo file."""

    profile = _client_profile(index, scenario)
    contact = profile["primary_contact"]
    secondary = profile["secondary_contact"]
    intake_filename, workplan_filename, chronology_filename = (
        _supporting_document_filenames(scenario["external_key"])
    )
    return [
        {
            "filename": intake_filename,
            "title": f"{scenario['name']} | Client Intake and Contact Profile",
            "subtitle": "Synthetic client file · privileged demonstration work paper",
            "matter": scenario["name"],
            "date": "August 18, 2026",
            "lead": scenario["lead"],
            "sections": [
                (
                    "Client Profile",
                    [
                        f"Organization: {profile['organization']}",
                        f"Primary contact: {contact['first_name']} {contact['last_name']}, {contact['title']}",
                        f"Secondary contact: {secondary['first_name']} {secondary['last_name']}, {secondary['title']}",
                        f"Office: {profile['address']['street']}, {profile['address']['city']}, {profile['address']['state']} {profile['address']['zip']}",
                        f"Direct line: {contact['phone']} · Email: {contact['email']}",
                        f"Client since: {profile['client_since']} · Preferred contact: {profile['preferred_contact_method']} · {profile['preferred_contact_window']}",
                        f"Billing profile: {profile['billing_terms']}",
                    ],
                ),
                (
                    "Initial Intake",
                    [
                        scenario["lead"],
                        "Conflict check recorded as cleared for this fictional demonstration file.",
                        "Client asked for a practical work plan, source-grounded risk summary, and attorney-reviewed communications.",
                    ],
                ),
                (
                    "Other Known Parties",
                    [
                        f"Counterparty or stakeholder: {profile['opposing_party']['organization']}.",
                        f"Primary outside contact: {profile['opposing_party']['contact_name']} ({profile['opposing_party']['email']}).",
                        "Party information is fictional and is included solely to demonstrate a complete matter file.",
                    ],
                ),
                (
                    "Open Questions for Counsel",
                    [
                        "Confirm governing documents, venue, and jurisdiction-specific requirements before advice or external action.",
                        "Confirm the client representative's authority and any insurance, notice, or consent obligations.",
                        "Identify records that require preservation or privilege treatment.",
                    ],
                ),
            ],
            "sources": ["synthetic_design"],
        },
        {
            "filename": workplan_filename,
            "title": f"{scenario['name']} | Communications and 30-Day Work Plan",
            "subtitle": "Synthetic correspondence log and attorney work plan",
            "matter": scenario["name"],
            "date": "August 18, 2026",
            "lead": "A simulated client update, internal analysis, and review-ready next steps for this fictional matter.",
            "sections": [
                (
                    "Recent Communications",
                    [
                        f"Client email: {contact['first_name']} {contact['last_name']} asked counsel to prioritize the immediate decision and summarize remaining information gaps.",
                        f"Follow-up call: {secondary['first_name']} {secondary['last_name']} confirmed the records owner and the preferred review cadence.",
                        "Attorney reply: acknowledged receipt, confirmed that no external communication will be sent without review, and requested the listed records.",
                        "Internal call note: team agreed to preserve relevant records and return a concise risk-and-options memo.",
                    ],
                ),
                (
                    "30-Day Work Plan",
                    [
                        f"Week 1: {scenario['suggested_tasks'][0]}.",
                        f"Week 2: {scenario['suggested_tasks'][1]}.",
                        f"Weeks 3–4: {scenario['suggested_tasks'][2]}.",
                    ],
                ),
                (
                    "Decision Log",
                    [
                        "Decision pending: confirm client objective, authority, and acceptable risk tolerance.",
                        "Decision pending: select an attorney-reviewed communication posture.",
                        "Decision pending: approve the next milestone and accountable owner.",
                    ],
                ),
                (
                    "Guided Demo Question",
                    [
                        scenario["demo_prompt"],
                        "Demonstrate how the assistant connects source documents, communications, tasks, and the matter timeline while leaving legal judgment with counsel.",
                    ],
                ),
            ],
            "sources": ["synthetic_design"],
        },
        {
            "filename": chronology_filename,
            "title": f"{scenario['name']} | Chronology and Document Index",
            "subtitle": "Synthetic matter chronology · review and production tracker",
            "matter": scenario["name"],
            "date": "August 20, 2026",
            "lead": "A review-ready index that connects the fictional matter's key events, custodians, and source records.",
            "sections": [
                (
                    "Working Chronology",
                    [
                        f"Day 0 - Intake opened after {contact['first_name']} {contact['last_name']} reported the central issue: {scenario['lead']}",
                        f"Day 1 - Counsel requested the first preservation set and assigned {scenario['suggested_tasks'][0].lower()}.",
                        f"Day 3 - {secondary['first_name']} {secondary['last_name']} confirmed the internal records owner and outstanding information gaps.",
                        f"Next milestone - Complete {scenario['suggested_tasks'][1].lower()} before the scheduled attorney review.",
                    ],
                ),
                (
                    "Document and Evidence Index",
                    [
                        "INT-001 - Client intake and contact profile; source: primary client contact; status: reviewed.",
                        "COM-001 - Communications and 30-day work plan; source: matter team; status: working draft.",
                        "SRC-001 - Primary scenario document; source: synthetic client file; status: indexed for demo search.",
                        "REQ-001 - Outstanding records request; owner: secondary client contact; status: pending.",
                    ],
                ),
                (
                    "Custodians and Review Owners",
                    [
                        f"Client decision-maker: {contact['first_name']} {contact['last_name']} ({contact['title']}).",
                        f"Records and billing contact: {secondary['first_name']} {secondary['last_name']} ({secondary['title']}).",
                        f"External stakeholder: {profile['opposing_party']['contact_name']} for {profile['opposing_party']['organization']}.",
                        "Matter lead: Jordan Lee (synthetic); final legal conclusions and external communications require attorney approval.",
                    ],
                ),
                (
                    "Next Review",
                    [
                        f"Use the guided question: {scenario['demo_prompt']}",
                        f"Immediate action: {scenario['suggested_tasks'][0]}.",
                        "Confirm document completeness, privilege, and governing-law assumptions before reliance or production.",
                    ],
                ),
            ],
            "sources": ["synthetic_design"],
        },
    ]


DOCUMENTS.extend(
    document
    for index, scenario in enumerate(ALL_SCENARIOS)
    for document in _supporting_documents(index, scenario)
)


def _font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style(style, *, size, color, before, after, line=1.1, bold=False):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def _numbering(
    doc: Document, num_format: str, text: str, indent: int, hanging: int
) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_format)
    level.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    level.append(lvl_text)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(indent))
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), str(indent))
    indentation.set(qn("w:hanging"), str(hanging))
    paragraph_properties.append(indentation)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "120")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.append(spacing)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_number(paragraph, num_id: int):
    properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_properties.append(level)
    num_properties.append(number)
    properties.append(num_properties)


def _page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)
    _font(run, size=9, color=MUTED)


def _scrub_revision_ids(path: Path) -> None:
    """Remove Word editing-session IDs from the generated package."""

    with tempfile.NamedTemporaryFile(
        delete=False, suffix=".docx", dir=path.parent
    ) as handle:
        temporary = Path(handle.name)
    try:
        with (
            ZipFile(path, "r") as source,
            ZipFile(temporary, "w", compression=ZIP_DEFLATED) as target,
        ):
            for item in source.infolist():
                payload = source.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    # Do not parse/re-serialize with stdlib ElementTree here. It
                    # rewrites lxml's namespace map (for example to ns0/ns1) but
                    # leaves mc:Ignorable values such as ``w14 wp14`` untouched;
                    # Word then rejects the package as corrupt. A byte-level
                    # removal of the generated rsid attributes preserves the
                    # original OOXML namespace declarations and formatting.
                    payload = re.sub(
                        rb'\s+w:rsid[A-Za-z0-9]+=(?:"[^"]*"|\'[^\']*\')',
                        b"",
                        payload,
                    )
                    # python-docx can also emit standalone revision-id elements
                    # in styles/settings (for example w:rsids and w:rsidRoot).
                    # Remove those with byte-level patterns so namespace maps
                    # remain untouched while preserving the privacy scrubber's
                    # prior contract that no w:rsid material ships.
                    revision_elements = rb"rsids?|rsidRoot|rsidDel|rsidP|rsidR|rsidSect"
                    payload = re.sub(
                        rb"<w:(" + revision_elements + rb")(?:\s[^>]*)?>.*?</w:\1>",
                        b"",
                        payload,
                        flags=re.DOTALL,
                    )
                    payload = re.sub(
                        rb"<w:(" + revision_elements + rb")(?:\s[^>]*)?/>",
                        b"",
                        payload,
                    )
                target.writestr(item, payload)
        os.replace(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)


def build_document(spec: dict) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    _set_style(styles["Normal"], size=11, color=INK, before=0, after=6, line=1.1)
    _set_style(styles["Heading 1"], size=16, color=BLUE, before=16, after=8, bold=True)
    _set_style(styles["Heading 2"], size=13, color=BLUE, before=12, after=6, bold=True)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    _font(header.add_run(DISCLAIMER), size=8.5, color=RGBColor(139, 37, 0), bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    _font(footer.add_run(f"{spec['title']}  |  Page "), size=9, color=MUTED)
    _page_field(footer)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(4)
    _font(title.add_run(spec["title"]), size=24, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _font(subtitle.add_run(spec["subtitle"]), size=13, color=MUTED)
    for label, value in (
        ("Matter", spec["matter"]),
        ("Document date", spec["date"]),
        ("Status", DISCLAIMER),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        _font(paragraph.add_run(f"{label}: "), size=10.5, color=INK, bold=True)
        _font(paragraph.add_run(value), size=10.5, color=INK)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(10)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.left_indent = Inches(0.2)
    lead.paragraph_format.right_indent = Inches(0.2)
    _font(
        lead.add_run(spec["lead"]), size=10.5, color=RGBColor(80, 65, 25), italic=True
    )

    heading_num = _numbering(doc, "decimal", "%1.", 540, 270)
    bullet_num = _numbering(doc, "bullet", "-", 540, 270)
    for heading, paragraphs in spec["sections"]:
        h = doc.add_paragraph(style="Heading 1")
        _apply_number(h, heading_num)
        h.add_run(heading)
        for text in paragraphs:
            doc.add_paragraph(text)

    if spec.get("signatories"):
        heading = doc.add_paragraph(
            "Acknowledged by the synthetic signatories", style="Heading 1"
        )
        for signatory in spec["signatories"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            _font(p.add_run(f"/s/ {signatory}"), size=11, color=INK)

    source_heading = doc.add_paragraph("Research basis", style="Heading 1")
    source_heading.paragraph_format.keep_with_next = True
    source_note = doc.add_paragraph(
        "These sources informed the fictional issue pattern and document structure. They are not "
        "incorporated into the agreement and are not a substitute for jurisdiction-specific review."
    )
    source_note.paragraph_format.keep_with_next = True
    for source_key in spec["sources"]:
        p = doc.add_paragraph()
        _apply_number(p, bullet_num)
        p.add_run(SOURCES[source_key])

    doc.core_properties.title = spec["title"]
    doc.core_properties.subject = DISCLAIMER
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = (
        "Fictional BK24 corporate-law demonstration document."
    )
    output = OUT / spec["filename"]
    doc.save(output)
    _scrub_revision_ids(output)
    return output


def build_manifest():
    matters = [
        {
            "external_key": "demo-northstar-saas-review",
            "primary_plugin": "commercial-legal",
            "matter_type": "commercial",
            "jurisdiction": "Delaware",
            "name": "Northstar Analytics - SaaS Vendor Review",
            "practice_area": "Corporate / Commercial Contracts",
            "status": "Open - Negotiation",
            "client": "Northstar Analytics, Inc. (fictional)",
            "description": "A customer-side SaaS renewal with AI-training, security, and liability issues.",
            "documents": [DOCUMENTS[0]["filename"], DOCUMENTS[1]["filename"]],
            "demo_prompt": "Review the MSA and data addendum for customer-side risks, cite the source clauses, and propose follow-up tasks for the renewal and data-use issues.",
            "suggested_tasks": [
                "Negotiate AI training and de-identification restriction",
                "Propose a 12-month fee cap and calendar the renewal notice",
                "Add security/privacy liability carve-out and balanced indemnity",
            ],
            "renewal": {"date": "2027-07-31", "notice_days": 60},
        },
        {
            "external_key": "demo-harborlight-seed-financing",
            "primary_plugin": "corporate-legal",
            "matter_type": "corporate",
            "jurisdiction": "Delaware",
            "name": "HarborLight Robotics - Series Seed Financing",
            "practice_area": "Corporate / Financing",
            "status": "Open - Pre-Closing",
            "client": "HarborLight Robotics, Inc. (fictional)",
            "description": "A founder-friendly seed financing with approvals, cap-table, and closing work.",
            "documents": [DOCUMENTS[2]["filename"], DOCUMENTS[3]["filename"]],
            "demo_prompt": "Compare the term sheet to the board consent, identify missing approvals and closing conditions, and prepare a prioritized closing checklist.",
            "suggested_tasks": [
                "Verify charter/bylaw approval thresholds and director conflicts",
                "Prepare stockholder consent and charter filing package",
                "Reconcile option pool and capitalization certificate",
                "Collect invention assignments and key-person agreements",
            ],
            "target_close": "2026-09-15",
        },
        {
            "external_key": "demo-redwood-ogc-retainer",
            "primary_plugin": "commercial-legal",
            "matter_type": "commercial",
            "jurisdiction": "Ohio",
            "name": "Redwood Outdoor Supply - Outside General Counsel Retainer",
            "practice_area": "Outside General Counsel",
            "status": "Open - Recurring",
            "client": "Redwood Outdoor Supply, Inc. (fictional)",
            "description": "A lively monthly outside-counsel portfolio with contract, product, and governance work.",
            "documents": [DOCUMENTS[4]["filename"], DOCUMENTS[5]["filename"]],
            "demo_prompt": "Triage the August legal requests, show the next notice deadlines, and draft the monthly retainer portfolio summary with decisions needed.",
            "suggested_tasks": [
                "Review PinePeak exclusivity and forecast remedies",
                "Calendar Evergreen Packaging non-renewal notice",
                "Review warehouse pilot data and security terms",
                "Prepare September board-materials legal section",
            ],
            "retainer": {"monthly_usd": 12500, "included_attorney_hours": 45},
        },
    ]
    matters.extend(
        {
            "external_key": scenario["external_key"],
            "primary_plugin": scenario["primary_plugin"],
            "matter_type": scenario["matter_type"],
            "jurisdiction": scenario["jurisdiction"],
            "name": scenario["name"],
            "practice_area": scenario["practice_area"],
            "status": scenario["status"],
            "client": scenario["client"],
            "description": scenario["lead"],
            "documents": [scenario["filename"]],
            "demo_prompt": scenario["demo_prompt"],
            "suggested_tasks": scenario["suggested_tasks"],
        }
        for scenario in ADDITIONAL_SCENARIOS
    )
    for index, matter in enumerate(matters):
        matter["documents"].extend(
            _supporting_document_filenames(matter["external_key"])
        )
        matter["client_profile"] = _client_profile(index, matter)
    manifest = {
        "schema_version": 3,
        "pack_version": "demo-scenario-library-v2",
        "tenant_domain": "cybersafeadvisor.com",
        "synthetic": True,
        "warning": DISCLAIMER,
        "matters": matters,
    }
    (OUT / "manifest.json").write_text(
        json.dumps(manifest, indent=2) + "\n", encoding="utf-8"
    )


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    outputs = [build_document(spec) for spec in DOCUMENTS]
    build_manifest()
    for output in outputs:
        print(output.relative_to(ROOT))


if __name__ == "__main__":
    main()
