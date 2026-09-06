"""Durable, agent-local lexical index for private file-share content.

The derived text index stays beside the agent ledger on the
customer-controlled host. This module does not wire query results or excerpts
into the outbound task relay.
"""

from __future__ import annotations

import asyncio
import inspect
import io
import json
import logging
import re
import time
from collections.abc import Awaitable, Callable
from pathlib import Path, PureWindowsPath

import aiosqlite
from docx import Document
from pypdf import PdfReader

from clarity_agent.config import _restrict
from clarity_agent.native_acl import authorize_acl, capture_windows_acl


logger = logging.getLogger("clarity_agent.local_index")

SCHEMA = """
CREATE TABLE IF NOT EXISTS index_files (
    path TEXT PRIMARY KEY,
    share_id TEXT NOT NULL,
    ext TEXT NOT NULL DEFAULT '',
    content_hash TEXT,
    size_bytes INTEGER,
    modified_time TEXT,
    status TEXT NOT NULL,
    attempts INTEGER NOT NULL DEFAULT 0,
    lease_until REAL,
    next_attempt_at REAL,
    page_count INTEGER,
    extraction_error TEXT,
    indexed_at TEXT,
    acl_json TEXT,
    acl_state TEXT NOT NULL DEFAULT 'unknown',
    acl_version TEXT,
    acl_captured_at INTEGER
);
CREATE VIRTUAL TABLE IF NOT EXISTS index_fts USING fts5(
    text,
    path UNINDEXED,
    share_id UNINDEXED,
    page_no UNINDEXED,
    ordinal UNINDEXED,
    ext UNINDEXED,
    tokenize='porter unicode61 remove_diacritics 2'
);
CREATE INDEX IF NOT EXISTS idx_index_files_share ON index_files(share_id);
CREATE INDEX IF NOT EXISTS idx_index_files_queue
    ON index_files(status, next_attempt_at, lease_until);
"""

SUPPORTED = {".pdf", ".docx", ".docm", ".txt", ".rtf"}
MAX_QUERY_CHARS = 1000
MAX_QUERY_TERMS = 32
MAX_RESULTS = 50
MAX_SNIPPET_CHARS = 1000
MAX_CHUNK_CHARS = 20_000
CHUNK_OVERLAP_CHARS = 256
DEFAULT_MAX_TEXT_CHARS = 5_000_000
MAX_ATTEMPTS = 3
MAX_WORKERS = 4
SCHEMA_VERSION = 2

REQUIRED_FILE_COLUMNS = {
    "path",
    "share_id",
    "ext",
    "content_hash",
    "size_bytes",
    "modified_time",
    "status",
    "attempts",
    "lease_until",
    "next_attempt_at",
    "page_count",
    "extraction_error",
    "indexed_at",
    "acl_json",
    "acl_state",
    "acl_version",
    "acl_captured_at",
}


def _validated_local_db_path(raw_path: str) -> Path:
    """Return a resolved local database path or reject unsafe locations."""
    windows_path = PureWindowsPath(raw_path)
    if str(windows_path.anchor).startswith("\\\\"):
        raise ValueError("local index path must be on a local disk, not a UNC path")
    db_file = Path(raw_path)
    if not db_file.is_absolute():
        raise ValueError("local index path must be absolute")
    if db_file.parent == Path(db_file.anchor):
        raise ValueError("local index must use a dedicated subdirectory")
    return db_file.resolve()


REQUIRED_FTS_COLUMNS = {"text", "path", "share_id", "page_no", "ordinal", "ext"}

CLAIMED_JOB_PREDICATE = """
path=? AND share_id=? AND ext=? AND content_hash IS ? AND size_bytes IS ?
AND modified_time IS ? AND status='running' AND attempts=? AND lease_until IS ?
"""


def _claimed_job_params(job: dict) -> tuple:
    """Identity of one exact queue claim, including its lease generation."""
    return (
        str(job["path"]),
        str(job["share_id"]),
        str(job.get("ext") or ""),
        job.get("content_hash"),
        int(job.get("size_bytes") or 0),
        job.get("modified_time"),
        int(job["_claim_attempt"]),
        float(job["_claim_lease_until"]),
    )


# These words are search-request scaffolding, not legal concepts. Dropping
# them before constructing the FTS query keeps a natural question such as
# "find matters discussing negligent spoliation" from requiring every UI word
# to appear verbatim in the source document.
QUERY_STOP_WORDS = {
    "a",
    "about",
    "all",
    "an",
    "and",
    "document",
    "documents",
    "every",
    "file",
    "files",
    "find",
    "for",
    "from",
    "in",
    "into",
    "matter",
    "matters",
    "of",
    "on",
    "or",
    "please",
    "show",
    "that",
    "the",
    "these",
    "this",
    "those",
    "through",
    "to",
    "where",
    "which",
    "with",
}


class PermanentIndexError(ValueError):
    """An indexing failure that retrying cannot repair."""


def _bounded_chunks(
    text: str,
    *,
    page_no: int | None,
    start_ordinal: int,
    remaining_chars: int,
) -> tuple[list[tuple[int | None, int, str]], int, int]:
    """Split text into bounded FTS rows while retaining page provenance."""
    text = text[:remaining_chars]
    if not text.strip():
        return [], start_ordinal, remaining_chars
    rows: list[tuple[int | None, int, str]] = []
    offset = 0
    ordinal = start_ordinal
    while offset < len(text):
        end = min(len(text), offset + MAX_CHUNK_CHARS)
        chunk = text[offset:end]
        if chunk.strip():
            rows.append((page_no, ordinal, chunk))
            ordinal += 1
        if end >= len(text):
            break
        offset = max(offset + 1, end - CHUNK_OVERLAP_CHARS)
    return rows, ordinal, remaining_chars - len(text)


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _rtf_text(content: bytes) -> str:
    """Extract conservative plain text without executing embedded content."""
    source = _decode_text(content)

    def hex_escape(match: re.Match[str]) -> str:
        try:
            return bytes.fromhex(match.group(1)).decode("cp1252")
        except (UnicodeDecodeError, ValueError):
            return " "

    source = re.sub(r"\\'([0-9a-fA-F]{2})", hex_escape, source)
    source = source.replace(r"\{", "{").replace(r"\}", "}")
    source = source.replace(r"\\", "\\")
    source = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", source)
    source = re.sub(r"[{}]", " ", source)
    return re.sub(r"\s+", " ", source).strip()


def _extract_chunks(
    path: str,
    content: bytes,
    max_text_chars: int,
) -> tuple[list[tuple[int | None, int, str]], int | None]:
    ext = PureWindowsPath(path).suffix.lower()
    rows: list[tuple[int | None, int, str]] = []
    ordinal = 0
    remaining = max_text_chars

    if ext == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        for page_number, page in enumerate(reader.pages, 1):
            if remaining <= 0:
                break
            page_rows, ordinal, remaining = _bounded_chunks(
                page.extract_text() or "",
                page_no=page_number,
                start_ordinal=ordinal,
                remaining_chars=remaining,
            )
            rows.extend(page_rows)
        return rows, len(reader.pages)

    if ext in {".docx", ".docm"}:
        document = Document(io.BytesIO(content))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for table_row in table.rows:
                blocks.extend(cell.text for cell in table_row.cells)
        for block in blocks:
            if remaining <= 0:
                break
            block_rows, ordinal, remaining = _bounded_chunks(
                block,
                page_no=None,
                start_ordinal=ordinal,
                remaining_chars=remaining,
            )
            rows.extend(block_rows)
        return rows, None

    text = _rtf_text(content) if ext == ".rtf" else _decode_text(content)
    rows, _, _ = _bounded_chunks(
        text,
        page_no=None,
        start_ordinal=0,
        remaining_chars=remaining,
    )
    return rows, None


def _plain_fts_queries(query: str) -> tuple[str, ...]:
    """Return strict then recall-oriented, non-programmable FTS queries."""
    terms = re.findall(r"[^\W_]+(?:['’][^\W_]+)?", query, flags=re.UNICODE)
    terms = terms[:MAX_QUERY_TERMS]
    if not terms:
        raise ValueError("query must contain at least one searchable term")
    meaningful = [term for term in terms if term.casefold() not in QUERY_STOP_WORDS]
    terms = meaningful or terms
    quoted = [f'"{term}"' for term in terms]
    strict = " AND ".join(quoted)
    relaxed = " OR ".join(quoted)
    return tuple(dict.fromkeys((strict, relaxed)))


def _relative_under_root(root: str, path: str) -> str | None:
    normalized_root = (root or "").replace("/", "\\").rstrip("\\")
    normalized_path = (path or "").replace("/", "\\").strip()
    prefix = normalized_root + "\\"
    if not normalized_root or not normalized_path.casefold().startswith(
        prefix.casefold()
    ):
        return None
    relative = normalized_path[len(prefix) :].strip("\\")
    parts = [part for part in relative.split("\\") if part]
    if not parts or any(part in {".", ".."} for part in parts):
        return None
    return "\\".join(parts)


def _escape_like(value: str) -> str:
    return value.replace("!", "!!").replace("%", "!%").replace("_", "!_")


def _database_bytes(db_path: str) -> int:
    return sum(
        candidate.stat().st_size
        for candidate in (
            Path(db_path),
            Path(db_path + "-wal"),
            Path(db_path + "-shm"),
        )
        if candidate.exists()
    )


async def _table_columns(db: aiosqlite.Connection, table: str) -> set[str]:
    cursor = await db.execute(f"PRAGMA table_info({table})")
    return {str(row["name"]) for row in await cursor.fetchall()}


async def _validate_schema(db: aiosqlite.Connection) -> None:
    version_cursor = await db.execute("PRAGMA user_version")
    version_row = await version_cursor.fetchone()
    version = int((version_row[0] if version_row else 0) or 0)
    if version != SCHEMA_VERSION:
        raise RuntimeError(f"unsupported local index schema version {version}")
    file_columns = await _table_columns(db, "index_files")
    fts_columns = await _table_columns(db, "index_fts")
    if not REQUIRED_FILE_COLUMNS.issubset(file_columns):
        raise RuntimeError("local index file manifest schema is incomplete")
    if not REQUIRED_FTS_COLUMNS.issubset(fts_columns):
        raise RuntimeError("local index search schema is incomplete")


async def _initialize_schema(db: aiosqlite.Connection) -> None:
    version_cursor = await db.execute("PRAGMA user_version")
    version_row = await version_cursor.fetchone()
    version = int((version_row[0] if version_row else 0) or 0)
    if version > SCHEMA_VERSION:
        raise RuntimeError(f"unsupported local index schema version {version}")
    await db.executescript(SCHEMA)
    columns = await _table_columns(db, "index_files")
    if "ext" not in columns:
        # Version 0 was used only during development. Preserve its rows while
        # allowing subsequent scans to populate the extension accurately.
        await db.execute(
            "ALTER TABLE index_files ADD COLUMN ext TEXT NOT NULL DEFAULT ''"
        )
    if "acl_json" not in columns:
        await db.execute("ALTER TABLE index_files ADD COLUMN acl_json TEXT")
        await db.execute(
            "ALTER TABLE index_files ADD COLUMN acl_state TEXT NOT NULL DEFAULT 'unknown'"
        )
        await db.execute("ALTER TABLE index_files ADD COLUMN acl_version TEXT")
        await db.execute("ALTER TABLE index_files ADD COLUMN acl_captured_at INTEGER")
    cursor = await db.execute("SELECT path FROM index_files WHERE ext='' ORDER BY path")
    while rows := await cursor.fetchmany(500):
        await db.executemany(
            "UPDATE index_files SET ext=? WHERE path=?",
            [
                (PureWindowsPath(str(row["path"])).suffix.lower(), str(row["path"]))
                for row in rows
            ],
        )
    await db.execute(
        "CREATE INDEX IF NOT EXISTS idx_index_files_ext_status "
        "ON index_files(ext, status)"
    )
    await db.execute(f"PRAGMA user_version = {SCHEMA_VERSION}")
    await _validate_schema(db)


async def _aggregate_stats(db: aiosqlite.Connection) -> tuple[dict, dict, int, dict]:
    cursor = await db.execute(
        """SELECT status, count(*) AS files,
                  coalesce(sum(size_bytes), 0) AS source_bytes
           FROM index_files GROUP BY status ORDER BY status"""
    )
    rows = await cursor.fetchall()
    extension_cursor = await db.execute(
        """SELECT CASE WHEN ext='' THEN '(unknown)' ELSE ext END AS extension,
                  status, count(*) AS files,
                  coalesce(sum(size_bytes), 0) AS source_bytes
           FROM index_files
           GROUP BY extension, status
           ORDER BY extension, status"""
    )
    extension_rows = await extension_cursor.fetchall()
    fts_cursor = await db.execute("SELECT count(*) FROM index_fts")
    fts_row = await fts_cursor.fetchone()
    acl_cursor = await db.execute(
        "SELECT acl_state, count(*) AS files FROM index_files GROUP BY acl_state"
    )
    acl_rows = await acl_cursor.fetchall()
    statuses = {
        str(row["status"]): {
            "files": int(row["files"] or 0),
            "source_bytes": int(row["source_bytes"] or 0),
        }
        for row in rows
    }
    by_extension: dict[str, dict[str, dict[str, int]]] = {}
    for row in extension_rows:
        by_extension.setdefault(str(row["extension"]), {})[str(row["status"])] = {
            "files": int(row["files"] or 0),
            "source_bytes": int(row["source_bytes"] or 0),
        }
    acl_states = {
        str(row["acl_state"] or "unknown"): int(row["files"] or 0) for row in acl_rows
    }
    return statuses, by_extension, int((fts_row[0] if fts_row else 0) or 0), acl_states


async def read_index_stats(db_path: str) -> dict:
    """Read aggregate index coverage without mutating or rebuilding the index."""
    db_file = Path(db_path)
    if not db_file.is_file():
        return {
            "available": False,
            "database_bytes": 0,
            "fts_rows": 0,
            "statuses": {},
            "by_extension": {},
            "acl_states": {},
        }
    try:
        uri = db_file.resolve().as_uri() + "?mode=ro"
        async with aiosqlite.connect(uri, uri=True) as db:
            db.row_factory = aiosqlite.Row
            await db.execute("PRAGMA query_only=ON")
            await _validate_schema(db)
            statuses, by_extension, fts_rows, acl_states = await _aggregate_stats(db)
    except Exception:
        return {
            "available": False,
            "database_bytes": _database_bytes(db_path),
            "fts_rows": 0,
            "statuses": {},
            "by_extension": {},
            "acl_states": {},
        }
    return {
        "available": True,
        "database_bytes": _database_bytes(db_path),
        "fts_rows": fts_rows,
        "statuses": statuses,
        "by_extension": by_extension,
        "acl_states": acl_states,
    }


class LocalSearchIndex:
    """Private, optional lexical index. Originals never leave the agent host."""

    def __init__(
        self,
        db_path: str,
        max_file_bytes: int = 25 * 1024 * 1024,
        max_text_chars: int = DEFAULT_MAX_TEXT_CHARS,
        acl_refresh_seconds: int = 3600,
    ):
        self.db_path = db_path
        self.max_file_bytes = max_file_bytes
        self.max_text_chars = max(1, min(int(max_text_chars), DEFAULT_MAX_TEXT_CHARS))
        self.acl_refresh_seconds = max(60, int(acl_refresh_seconds))
        self._db: aiosqlite.Connection | None = None
        self._workers: list[asyncio.Task] = []
        self._fetcher: Callable[[dict], Awaitable[bytes]] | None = None
        self._path_validator: Callable[[dict], Awaitable[bool]] | None = None
        self._acl_loader: Callable = lambda job: capture_windows_acl(job["path"])
        self._db_lock = asyncio.Lock()
        self.available = False
        self._wake = asyncio.Event()
        self._readonly = False

    async def init(self) -> None:
        if self._db is not None or self._workers:
            raise RuntimeError("local index is already initialized")
        self._readonly = False
        db_file = _validated_local_db_path(self.db_path)
        self.db_path = str(db_file)
        db_file.parent.mkdir(parents=True, exist_ok=True)
        _restrict(db_file.parent, required=True)
        self._db = await aiosqlite.connect(self.db_path)
        self._db.row_factory = aiosqlite.Row
        try:
            # WAL lets searches proceed while a document is committed. A
            # bounded page cache avoids turning an old host into a swap storm.
            await self._db.execute("PRAGMA journal_mode=WAL")
            await self._db.execute("PRAGMA synchronous=NORMAL")
            await self._db.execute("PRAGMA busy_timeout=5000")
            await self._db.execute("PRAGMA cache_size=-32768")
            await self._db.execute("PRAGMA temp_store=MEMORY")
            await _initialize_schema(self._db)
            await self._db.commit()
            _restrict(db_file, required=True)
            self.available = True
            await self._db.execute(
                "UPDATE index_files SET status='pending', lease_until=NULL "
                "WHERE status='running'"
            )
            await self._db.commit()
        except Exception as exc:
            await self._db.rollback()
            await self._db.close()
            self._db = None
            self.available = False
            logger.warning(
                "Local full-text index is unavailable: %s", type(exc).__name__
            )

    async def init_readonly(self) -> None:
        """Open an existing index without creating, migrating, or writing it.

        Benchmark and diagnostic callers must not use :meth:`init`, since that
        method initializes the schema and recovers abandoned leases.
        """
        if self._db is not None or self._workers:
            raise RuntimeError("local index is already initialized")
        db_file = _validated_local_db_path(self.db_path)
        if not db_file.is_file():
            raise FileNotFoundError(str(db_file))
        wal_file = Path(str(db_file) + "-wal")
        if wal_file.exists() and wal_file.stat().st_size:
            raise RuntimeError(
                "read-only evaluation requires a stopped, checkpointed index"
            )
        # ``immutable=1`` avoids creating SQLite WAL/SHM sidecars. It is safe
        # only after the writer is stopped and its WAL is absent, checked above.
        uri = db_file.as_uri() + "?mode=ro&immutable=1"
        self._db = await aiosqlite.connect(uri, uri=True)
        self._db.row_factory = aiosqlite.Row
        try:
            await self._db.execute("PRAGMA query_only=ON")
            await self._db.execute("PRAGMA busy_timeout=5000")
            await _validate_schema(self._db)
        except Exception:
            await self._db.close()
            self._db = None
            raise
        self.db_path = str(db_file)
        self._readonly = True
        self.available = True

    async def close(self) -> None:
        for worker in self._workers:
            worker.cancel()
        if self._workers:
            await asyncio.gather(*self._workers, return_exceptions=True)
            self._workers.clear()
        if self._db:
            await self._db.close()
            self._db = None
        self.available = False

    def _require_writable(self) -> None:
        if self._readonly:
            raise RuntimeError("local index was opened read-only")

    def start(
        self,
        fetcher: Callable[[dict], Awaitable[bytes]],
        *,
        path_validator: Callable[[dict], Awaitable[bool]],
        acl_loader: Callable | None = None,
        worker_count: int = 1,
    ) -> None:
        self._require_writable()
        self._fetcher = fetcher
        self._path_validator = path_validator
        if acl_loader is not None:
            self._acl_loader = acl_loader
        if self.available and not self._workers:
            count = max(1, min(int(worker_count or 1), MAX_WORKERS))
            self._workers = [
                asyncio.create_task(
                    self._run(), name=f"lawhand-local-index-{number + 1}"
                )
                for number in range(count)
            ]

    async def enqueue(self, file_info: dict) -> None:
        await self.enqueue_many([file_info])

    async def enqueue_many(
        self, files: list[dict], *, only_if_missing: bool = False
    ) -> None:
        """Queue changed files, or seed missing rows from an older scan ledger."""
        if not self.available or not self._db or not files:
            return
        self._require_writable()
        async with self._db_lock:
            await self._db.execute("BEGIN IMMEDIATE")
            try:
                for file_info in files:
                    path = str(file_info["path"])
                    ext = str(
                        file_info.get("ext") or PureWindowsPath(path).suffix
                    ).lower()
                    size = int(file_info.get("size_bytes") or 0)
                    status = "pending"
                    error = None
                    if ext not in SUPPORTED:
                        status, error = "unsupported", "unsupported_format"
                    elif size > self.max_file_bytes:
                        status, error = "error", "file_too_large"

                    values = (
                        path,
                        str(file_info["share_id"]),
                        ext,
                        file_info.get("content_hash"),
                        size,
                        file_info.get("modified_time"),
                        status,
                        error,
                    )
                    if only_if_missing:
                        await self._db.execute(
                            """INSERT OR IGNORE INTO index_files(
                                   path,share_id,ext,content_hash,size_bytes,modified_time,
                                   status,extraction_error
                               ) VALUES(?,?,?,?,?,?,?,?)""",
                            values,
                        )
                    else:
                        await self._db.execute(
                            "DELETE FROM index_fts WHERE path=?", (path,)
                        )
                        await self._db.execute(
                            """INSERT INTO index_files(
                                   path,share_id,ext,content_hash,size_bytes,modified_time,
                                   status,extraction_error
                               ) VALUES(?,?,?,?,?,?,?,?)
                               ON CONFLICT(path) DO UPDATE SET
                                   share_id=excluded.share_id,
                                   ext=excluded.ext,
                                   content_hash=excluded.content_hash,
                                   size_bytes=excluded.size_bytes,
                                   modified_time=excluded.modified_time,
                                   status=excluded.status,
                                   attempts=0,
                                   lease_until=NULL,
                                   next_attempt_at=NULL,
                                   page_count=NULL,
                                   extraction_error=excluded.extraction_error,
                                   indexed_at=NULL,
                                   acl_json=NULL,
                                   acl_state='pending',
                                   acl_version=NULL,
                                   acl_captured_at=NULL""",
                            values,
                        )
                # An unchanged file can still have a changed DACL.  Queue a
                # bounded refresh without deleting its old text; searches deny
                # the stale ACL until the refreshed record commits atomically.
                cutoff = int(time.time()) - self.acl_refresh_seconds
                await self._db.execute(
                    """UPDATE index_files
                       SET status='pending', acl_state='pending', attempts=0,
                           lease_until=NULL, next_attempt_at=NULL
                       WHERE status='ready' AND
                             (acl_captured_at IS NULL OR acl_captured_at<?)""",
                    (cutoff,),
                )
                await self._db.commit()
            except Exception:
                await self._db.rollback()
                raise
        self._wake.set()

    async def delete(self, paths: list[str]) -> None:
        if not self.available or not self._db:
            return
        self._require_writable()
        if not paths:
            return
        async with self._db_lock:
            await self._db.execute("BEGIN IMMEDIATE")
            try:
                for path in paths:
                    await self._db.execute(
                        "DELETE FROM index_fts WHERE path=?", (path,)
                    )
                    await self._db.execute(
                        "DELETE FROM index_files WHERE path=?", (path,)
                    )
                await self._db.commit()
            except Exception:
                await self._db.rollback()
                raise

    async def _claim(self) -> dict | None:
        assert self._db
        now = time.time()
        async with self._db_lock:
            await self._db.execute("BEGIN IMMEDIATE")
            try:
                cursor = await self._db.execute(
                    """SELECT * FROM index_files
                       WHERE (status='pending' AND
                              (next_attempt_at IS NULL OR next_attempt_at<=?))
                          OR (status='running' AND lease_until<?)
                       ORDER BY rowid LIMIT 1""",
                    (now, now),
                )
                row = await cursor.fetchone()
                if not row:
                    await self._db.commit()
                    return None
                lease_until = now + 300
                claimed_attempt = int(row["attempts"] or 0) + 1
                claimed = await self._db.execute(
                    """UPDATE index_files
                       SET status='running', attempts=?, lease_until=?,
                           next_attempt_at=NULL
                       WHERE path=? AND share_id=? AND ext=?
                         AND content_hash IS ? AND size_bytes IS ?
                         AND modified_time IS ? AND attempts=?
                         AND ((status='pending' AND
                               (next_attempt_at IS NULL OR next_attempt_at<=?))
                              OR (status='running' AND lease_until<?))""",
                    (
                        claimed_attempt,
                        lease_until,
                        row["path"],
                        row["share_id"],
                        row["ext"],
                        row["content_hash"],
                        row["size_bytes"],
                        row["modified_time"],
                        row["attempts"],
                        now,
                        now,
                    ),
                )
                await self._db.commit()
                if claimed.rowcount != 1:
                    return None
                job = dict(row)
                job["_claim_attempt"] = claimed_attempt
                job["_claim_lease_until"] = lease_until
                return job
            except Exception:
                await self._db.rollback()
                raise

    async def _next_wait_seconds(self) -> float:
        assert self._db
        async with self._db_lock:
            cursor = await self._db.execute(
                "SELECT min(next_attempt_at) FROM index_files WHERE status='pending'"
            )
            row = await cursor.fetchone()
        if not row or row[0] is None:
            return 60.0
        return max(0.05, min(60.0, float(row[0]) - time.time()))

    async def _run(self) -> None:
        while True:
            # Clear before claiming so an enqueue racing with the empty check
            # leaves the event set instead of losing the wake-up.
            self._wake.clear()
            job = await self._claim()
            if not job:
                try:
                    await asyncio.wait_for(
                        self._wake.wait(), timeout=await self._next_wait_seconds()
                    )
                except asyncio.TimeoutError:
                    pass
                continue
            try:
                if self._fetcher is None:
                    raise RuntimeError("local index fetcher is unavailable")
                if self._path_validator is None or not await self._path_validator(job):
                    raise PermanentIndexError("path_outside_assigned_share")
                acl_record = self._acl_loader(job)
                if inspect.isawaitable(acl_record):
                    acl_record = await acl_record
                content = await self._fetcher(job)
                if len(content) > self.max_file_bytes:
                    raise PermanentIndexError("file_too_large")
                rows, page_count = await self._extract_text(job, content)
                if not rows:
                    raise PermanentIndexError("no_extractable_text")
                assert self._db
                async with self._db_lock:
                    await self._db.execute("BEGIN IMMEDIATE")
                    try:
                        current = await self._db.execute(
                            "SELECT 1 FROM index_files WHERE " + CLAIMED_JOB_PREDICATE,
                            _claimed_job_params(job),
                        )
                        if not await current.fetchone():
                            # A rescan, delete, or newer lease superseded this
                            # extraction. Never publish its stale text.
                            await self._db.commit()
                            continue
                        await self._db.execute(
                            "DELETE FROM index_fts WHERE path=?", (job["path"],)
                        )
                        acl_record = await self._publish_text(job, rows, acl_record)
                        completed = await self._db.execute(
                            """UPDATE index_files
                               SET status='ready', page_count=?, extraction_error=NULL,
                                   lease_until=NULL, next_attempt_at=NULL,
                                   indexed_at=datetime('now'), acl_json=?, acl_state=?,
                                   acl_version=?, acl_captured_at=?
                               WHERE """
                            + CLAIMED_JOB_PREDICATE,
                            (
                                page_count,
                                json.dumps(
                                    acl_record, sort_keys=True, separators=(",", ":")
                                ),
                                str(acl_record.get("state") or "unknown")[:40],
                                str(acl_record.get("version") or "")[:128] or None,
                                int(acl_record.get("captured_at") or 0) or None,
                                *_claimed_job_params(job),
                            ),
                        )
                        if completed.rowcount != 1:
                            raise RuntimeError(
                                "local index claim changed during commit"
                            )
                        await self._db.commit()
                    except Exception:
                        await self._db.rollback()
                        raise
            except asyncio.CancelledError:
                raise
            except PermanentIndexError as exc:
                await self._record_failure(job, str(exc), retry=False)
            except Exception as exc:
                attempts = int(job.get("attempts") or 0) + 1
                # Parser/library messages frequently contain source paths or
                # fragments. Persist the class only; operators need a stable
                # failure category, not document data in the queue ledger.
                reason = f"transient:{type(exc).__name__}"
                await self._record_failure(job, reason, retry=attempts < MAX_ATTEMPTS)

    async def _extract_text(self, job: dict, content: bytes):
        return await asyncio.to_thread(
            _extract_chunks, job["path"], content, self.max_text_chars
        )

    async def _publish_text(self, job: dict, rows: list, acl_record: dict) -> dict:
        """Commit derived text while the manifest claim is fenced by its lock."""
        assert self._db
        ext = PureWindowsPath(job["path"]).suffix.lower()
        await self._db.executemany(
            """INSERT INTO index_fts(
                   text,path,share_id,page_no,ordinal,ext
               ) VALUES(?,?,?,?,?,?)""",
            [
                (
                    text,
                    job["path"],
                    job["share_id"],
                    page,
                    ordinal,
                    ext,
                )
                for page, ordinal, text in rows
            ],
        )
        return acl_record

    async def _record_failure(self, job: dict, reason: str, *, retry: bool) -> None:
        if not self._db:
            return
        status = "pending" if retry else "error"
        attempts = int(job.get("attempts") or 0) + 1
        next_attempt = time.time() + min(60, 2**attempts) if retry else None
        async with self._db_lock:
            updated = await self._db.execute(
                """UPDATE index_files
                   SET status=?, extraction_error=?, lease_until=NULL,
                       next_attempt_at=? WHERE """
                + CLAIMED_JOB_PREDICATE,
                (
                    status,
                    reason[:500],
                    next_attempt,
                    *_claimed_job_params(job),
                ),
            )
            await self._db.commit()
        if updated.rowcount == 1:
            self._wake.set()

    async def wait_until_idle(self) -> None:
        """Wait until the durable queue has no runnable or leased work."""
        if not self.available or not self._db:
            return
        while True:
            async with self._db_lock:
                cursor = await self._db.execute(
                    """SELECT count(*) FROM index_files
                       WHERE status IN ('pending','running')"""
                )
                row = await cursor.fetchone()
            if not row or int(row[0]) == 0:
                return
            await asyncio.sleep(0.05)

    async def stats(self) -> dict:
        """Return aggregate coverage telemetry without paths or document text."""
        if not self.available or not self._db:
            return {
                "available": False,
                "database_bytes": 0,
                "fts_rows": 0,
                "statuses": {},
                "by_extension": {},
            }
        async with self._db_lock:
            statuses, by_extension, fts_rows, acl_states = await _aggregate_stats(
                self._db
            )
        return {
            "available": True,
            "database_bytes": _database_bytes(self.db_path),
            "fts_rows": fts_rows,
            "statuses": statuses,
            "by_extension": by_extension,
            "acl_states": acl_states,
        }

    async def search(
        self,
        query: str,
        scopes: list[dict],
        assigned_shares: list[dict],
        extensions: list[str] | None,
        limit: int,
        authorization=None,
        acl_max_age_seconds: int = 3600,
    ) -> dict:
        if not self.available or not self._db:
            return {
                "hits": [],
                "index_state": "unavailable",
                "indexed_files": 0,
                "pending_files": 0,
            }
        query = (query or "").strip()
        if not query or len(query) > MAX_QUERY_CHARS:
            raise ValueError("query is required and must be at most 1000 characters")
        if not assigned_shares or not scopes:
            raise ValueError("at least one assigned search scope is required")
        limit = max(1, min(int(limit or 20), MAX_RESULTS))
        allowed = {str(item.get("share_id")): item for item in assigned_shares}

        scope_clauses: list[str] = []
        scope_params: list[str] = []
        for scope in scopes:
            share_id = str(scope.get("share_id") or "")
            share = allowed.get(share_id)
            if share is None:
                raise ValueError("search scope is not assigned to this agent")
            root = str(share.get("share_path") or "").replace("/", "\\").rstrip("\\")
            if not root:
                raise ValueError("assigned share is missing its canonical path")
            folder = str(scope.get("folder_path") or "").replace("\\", "/").strip("/")
            parts = [part for part in folder.split("/") if part and part != "."]
            if any(part == ".." for part in parts):
                raise ValueError("search folder is outside its assigned share")
            if parts:
                absolute = root + "\\" + "\\".join(parts)
                scope_clauses.append(
                    "(index_fts.share_id=? AND index_fts.path LIKE ? ESCAPE '!')"
                )
                scope_params.extend(
                    [share_id, _escape_like(absolute.rstrip("\\") + "\\") + "%"]
                )
            else:
                scope_clauses.append("(index_fts.share_id=?)")
                scope_params.append(share_id)

        extension_clause = ""
        extension_params: list[str] = []
        if extensions:
            normalized = []
            for extension in extensions:
                value = str(extension).strip().lower()
                if not value:
                    continue
                normalized.append(value if value.startswith(".") else "." + value)
            normalized = list(dict.fromkeys(normalized))[:50]
            if normalized:
                placeholders = ",".join("?" for _ in normalized)
                extension_clause = f" AND index_fts.ext IN ({placeholders})"
                extension_params.extend(normalized)

        sql = (
            "SELECT index_fts.*, index_files.acl_json AS acl_json, bm25(index_fts) AS rank, "
            "snippet(index_fts,0,'','', '…',48) AS match_snippet "
            "FROM index_fts JOIN index_files ON index_files.path=index_fts.path "
            "WHERE index_fts MATCH ? AND ("
            + " OR ".join(scope_clauses)
            + ")"
            + extension_clause
            + " ORDER BY rank LIMIT ?"
        )
        # The FTS table stores chunks, but callers request distinct documents.
        # Overfetch enough rows for every allowed chunk in ``limit`` documents
        # so one long pleading cannot crowd all other files out of the window.
        stride = MAX_CHUNK_CHARS - CHUNK_OVERLAP_CHARS
        max_chunks_per_file = 1 + max(
            0,
            (self.max_text_chars - MAX_CHUNK_CHARS + stride - 1) // stride,
        )
        candidate_limit = limit * max_chunks_per_file
        ranked_rows: list[tuple[aiosqlite.Row, int]] = []
        async with self._db_lock:
            for tier, fts_query in enumerate(_plain_fts_queries(query)):
                params = [
                    fts_query,
                    *scope_params,
                    *extension_params,
                    candidate_limit,
                ]
                cursor = await self._db.execute(sql, params)
                ranked_rows.extend((row, tier) for row in await cursor.fetchall())
            stats_cursor = await self._db.execute(
                """SELECT
                       sum(status='ready') AS ready,
                       sum(status IN ('pending','running')) AS pending,
                       sum(status IN ('error','unsupported')) AS failed,
                       count(*) AS total
                   FROM index_files"""
            )
            stats = await stats_cursor.fetchone()

        hits = []
        seen_paths: set[str] = set()
        for row, tier in ranked_rows:
            share = allowed.get(str(row["share_id"]))
            if share is None:
                continue
            full_path = str(row["path"])
            relative = _relative_under_root(
                str(share.get("share_path") or ""), full_path
            )
            if relative is None or full_path.casefold() in seen_paths:
                continue
            if authorization is not None:
                try:
                    acl_record = json.loads(row["acl_json"] or "null")
                except (TypeError, json.JSONDecodeError):
                    acl_record = None
                decision = authorize_acl(
                    acl_record,
                    authorization.principal_sids,
                    max_age_seconds=max(60, int(acl_max_age_seconds)),
                )
                if not decision.allowed:
                    continue
            seen_paths.add(full_path.casefold())
            snippet = str(row["match_snippet"] or "").strip()[:MAX_SNIPPET_CHARS]
            hits.append(
                {
                    "share_id": str(row["share_id"]),
                    "relative_path": relative,
                    "filename": PureWindowsPath(full_path).name,
                    "ext": PureWindowsPath(full_path).suffix.lower(),
                    "snippet": snippet,
                    "page_number": row["page_no"],
                    # SQLite BM25 is lower-is-better. Negate it so the relay's
                    # public score remains conventional higher-is-better.
                    # Strict all-term matches sort ahead of recall-oriented OR
                    # fallback matches. SQLite BM25 is lower-is-better, so the
                    # public score uses a positive higher-is-better convention.
                    "score": float(2 - tier) + max(0.0, -float(row["rank"])),
                }
            )
            if len(hits) >= limit:
                break

        ready = int((stats["ready"] if stats else 0) or 0)
        pending = int((stats["pending"] if stats else 0) or 0)
        failed = int((stats["failed"] if stats else 0) or 0)
        if pending:
            index_state = "partial" if ready else "building"
        elif failed:
            index_state = "partial" if ready else "degraded"
        elif ready:
            index_state = "ready"
        else:
            index_state = "empty"
        return {
            "hits": hits,
            "index_state": index_state,
            "indexed_files": ready,
            "pending_files": pending,
        }

    async def authorize_path(
        self, path: str, authorization, *, acl_max_age_seconds: int = 3600
    ):
        """Revalidate the current indexed ACL before preview/open content release."""
        if not self.available or not self._db:
            return authorize_acl(
                None, authorization.principal_sids, max_age_seconds=acl_max_age_seconds
            )
        async with self._db_lock:
            cursor = await self._db.execute(
                "SELECT share_id, acl_json FROM index_files WHERE path=? AND status='ready'",
                (path,),
            )
            row = await cursor.fetchone()
        if row is None or str(row["share_id"]) not in authorization.source_ids:
            return authorize_acl(
                None, authorization.principal_sids, max_age_seconds=acl_max_age_seconds
            )
        try:
            record = json.loads(row["acl_json"] or "null")
        except (TypeError, json.JSONDecodeError):
            record = None
        return authorize_acl(
            record,
            authorization.principal_sids,
            max_age_seconds=max(60, int(acl_max_age_seconds)),
        )
