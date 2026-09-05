"""The Word authoring outline: paragraphs addressed as fields address them."""

import io

import pytest
from docx import Document

from app.services.docx_outline import (
    MAX_OUTLINE_PARAGRAPHS,
    MAX_PARAGRAPH_CHARACTERS,
    docx_outline,
)
from app.services.docx_templates import (
    TemplateDocxError,
    fill_docx_template,
    iter_docx_paragraphs,
    _open_docx,
)


def _docx(build):
    document = Document()
    build(document)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestOrdinalContract:
    def test_ordinals_match_the_iterator_that_fills_the_template(self):
        # This is the whole contract: a span selected against the outline must
        # address the same paragraph at generation time. Reconstructing
        # ordinals from a rendered page is exactly what this avoids.
        def build(document):
            document.add_heading("ENGAGEMENT LETTER", level=1)
            document.add_paragraph("Client: Ada Lovelace")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Party"
            table.cell(0, 1).text = "Role"
            document.sections[0].footer.paragraphs[0].text = "Firm footer"

        source = _docx(build)
        outline = docx_outline(source)
        from_iterator = [
            paragraph.text for paragraph in iter_docx_paragraphs(_open_docx(source))
        ]
        assert [p["text"] for p in outline["paragraphs"]] == from_iterator
        assert [p["ordinal"] for p in outline["paragraphs"]] == list(
            range(len(from_iterator))
        )

    def test_an_anchor_taken_from_the_outline_fills_correctly(self):
        source = _docx(lambda d: [d.add_paragraph("x"), d.add_paragraph("Client: Ada Lovelace")])
        outline = docx_outline(source)
        target = next(p for p in outline["paragraphs"] if "Ada" in p["text"])
        start = target["text"].index("Ada Lovelace")

        filled = fill_docx_template(
            source,
            variable_schema={
                "fields": [
                    {
                        "name": "client_name",
                        "source_text": "Ada Lovelace",
                        "docx_anchor": {
                            "paragraph_ordinal": target["ordinal"],
                            "start": start,
                            "end": start + len("Ada Lovelace"),
                        },
                    }
                ]
            },
            variables={"client_name": "Grace Hopper"},
        )
        rendered = [p.text for p in iter_docx_paragraphs(_open_docx(filled))]
        assert "Client: Grace Hopper" in rendered


class TestParagraphDetail:
    def test_containers_are_reported_so_a_footer_reads_as_a_footer(self):
        def build(document):
            document.add_paragraph("Body line")
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Cell line"
            document.sections[0].footer.paragraphs[0].text = "Footer line"

        by_text = {p["text"]: p for p in docx_outline(_docx(build))["paragraphs"]}
        assert by_text["Body line"]["container"] == "body"
        assert by_text["Cell line"]["container"] == "table"
        assert by_text["Footer line"]["container"] == "footer"

    def test_styles_are_reported_for_headings(self):
        source = _docx(lambda d: d.add_heading("Title here", level=1))
        assert docx_outline(source)["paragraphs"][0]["style"] == "Heading 1"

    def test_runs_carry_formatting_and_their_own_offsets(self):
        def build(document):
            paragraph = document.add_paragraph()
            paragraph.add_run("Fee: ").bold = True
            paragraph.add_run("$250")

        runs = docx_outline(_docx(build))["paragraphs"][0]["runs"]
        assert [(r["text"], r["bold"], r["start"], r["end"]) for r in runs] == [
            ("Fee: ", True, 0, 5),
            ("$250", False, 5, 9),
        ]

    def test_logic_markers_are_identified_not_left_as_prose(self):
        def build(document):
            document.add_paragraph("{{#if is_entity}}")
            document.add_paragraph("Authority clause.")
            document.add_paragraph("{{/if}}")
            document.add_paragraph("{{#each parties}}")
            document.add_paragraph("{{/each}}")

        paragraphs = docx_outline(_docx(build))["paragraphs"]
        assert paragraphs[0]["marker"] == {
            "kind": "open",
            "keyword": "if",
            "name": "is_entity",
        }
        assert paragraphs[2]["marker"] == {"kind": "close", "keyword": "if", "name": ""}
        assert paragraphs[3]["marker"]["keyword"] == "each"
        assert "marker" not in paragraphs[1]

    def test_a_marker_sharing_a_paragraph_with_prose_is_prose(self):
        # It is prose to the renderer too, so calling it a marker here would
        # draw a region the generated document does not have.
        source = _docx(lambda d: d.add_paragraph("Clause {{#if a}} inline"))
        assert "marker" not in docx_outline(source)["paragraphs"][0]


class TestBounds:
    def test_paragraph_count_is_capped_and_reported(self):
        def build(document):
            for index in range(MAX_OUTLINE_PARAGRAPHS + 5):
                document.add_paragraph(f"line {index}")

        outline = docx_outline(_docx(build))
        assert outline["paragraph_count"] == MAX_OUTLINE_PARAGRAPHS
        assert outline["truncated"] is True

    def test_a_short_document_is_not_marked_truncated(self):
        outline = docx_outline(_docx(lambda d: d.add_paragraph("one")))
        assert outline["truncated"] is False

    def test_an_enormous_paragraph_is_clipped(self):
        source = _docx(lambda d: d.add_paragraph("x" * (MAX_PARAGRAPH_CHARACTERS + 100)))
        assert len(docx_outline(source)["paragraphs"][0]["text"]) == MAX_PARAGRAPH_CHARACTERS


class TestSafety:
    def test_a_damaged_file_is_a_customer_error(self):
        with pytest.raises(TemplateDocxError):
            docx_outline(b"not a docx at all")
