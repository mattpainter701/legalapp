import hashlib
import json
import uuid
from datetime import datetime, timedelta, timezone
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi import HTTPException
from sqlalchemy import select

import app.routers.esignature as esignature_router
import app.routers.matter_documents as matter_documents_router
import app.services.matter_document_revisions as revision_module
from app.models.conversation import UsageRecord
from app.models.matter_document import MatterDocument
from app.models.matter_document_revision import MatterDocumentRevision
from app.models.plugin import Matter
from app.models.signature import SignatureRequest, SignatureSigner
from app.schemas.matter_document import MatterDocumentUpdate
from app.schemas.matter_document_revision import (
    MatterDocumentRevisionApprove,
    MatterDocumentRevisionCreate,
    SignatureReplacementPrepare,
)
from app.schemas.signature import SignatureRequestCreate, SignerCreate
from app.services.llm_routing import LLMRoute
from app.services.matter_document_revisions import (
    DocumentRevisionServiceError,
    MatterDocumentRevisionService,
    assert_no_legacy_assistant_derivative_release,
)
from app.services.matter_file_store import MatterFileIntegrityError, StorageResult


def _docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class _MemoryFileStore:
    def __init__(self, source_path: str, source: bytes):
        self.files = {source_path: source}
        self.uploads: list[dict] = []
        self.deletions: list[str | None] = []

    async def read_matter_file_bytes(
        self,
        *,
        document,
        expected_sha256=None,
        expected_size=None,
        **_kwargs,
    ):
        content = self.files[document.storage_path]
        if expected_size is not None:
            if len(content) != expected_size:
                raise MatterFileIntegrityError("size mismatch")
        if expected_sha256 is not None:
            if hashlib.sha256(content).hexdigest() != expected_sha256:
                raise MatterFileIntegrityError("hash mismatch")
        return content

    async def store_matter_file_result(self, **kwargs):
        path = f"memory://{len(self.uploads) + 1}/{kwargs['filename']}"
        self.files[path] = kwargs["content"]
        self.uploads.append(kwargs)
        return StorageResult(
            provider="local",
            backend="local",
            storage_path=path,
        )

    async def delete_stored_result(self, *, result, **_kwargs):
        self.deletions.append(result.storage_path)
        self.files.pop(result.storage_path, None)


class _QueuedLLM:
    def __init__(self, *payloads: dict):
        self.payloads = list(payloads)
        self.calls: list[dict] = []

    async def complete(self, **kwargs):
        self.calls.append(kwargs)
        return json.dumps(self.payloads.pop(0)), 120, 40


def test_stale_processing_detection_is_pure_until_atomic_database_update():
    now = datetime.now(timezone.utc)
    row = MatterDocumentRevision(
        id=uuid.uuid4(),
        tenant_id=uuid.uuid4(),
        matter_id=uuid.uuid4(),
        root_document_id=uuid.uuid4(),
        source_document_id=uuid.uuid4(),
        client_request_id=uuid.uuid4(),
        version_no=1,
        instruction="Revise the client name",
        status="processing",
        source_sha256="a" * 64,
        requested_model_tier="auto",
        created_at=now - timedelta(minutes=11),
        updated_at=now - timedelta(minutes=11),
    )

    assert revision_module._is_stale_processing(row, now=now) is True
    assert row.status == "processing"
    assert row.error_code is None

    fresh = MatterDocumentRevision(
        id=uuid.uuid4(),
        tenant_id=row.tenant_id,
        matter_id=row.matter_id,
        root_document_id=row.root_document_id,
        source_document_id=row.source_document_id,
        client_request_id=uuid.uuid4(),
        version_no=2,
        instruction="Revise the date",
        status="processing",
        source_sha256="b" * 64,
        requested_model_tier="auto",
        created_at=now,
        updated_at=now,
    )
    assert revision_module._is_stale_processing(fresh, now=now) is False
    assert fresh.status == "processing"


@pytest.mark.asyncio
async def test_replacement_ancestor_chain_excludes_sibling_outputs():
    root_id = uuid.uuid4()
    ancestor_output_id = uuid.uuid4()
    sibling_output_id = uuid.uuid4()
    ancestor_id = uuid.uuid4()
    ancestor = SimpleNamespace(
        id=ancestor_id,
        root_document_id=root_id,
        source_document_id=root_id,
        output_document_id=ancestor_output_id,
        source_revision_id=None,
        version_no=1,
    )
    selected = SimpleNamespace(
        root_document_id=root_id,
        source_document_id=ancestor_output_id,
        source_revision_id=ancestor_id,
        version_no=3,
    )

    class _AncestorDB:
        async def scalar(self, _statement):
            return ancestor

    service = MatterDocumentRevisionService(
        llm=SimpleNamespace(), file_store=SimpleNamespace()
    )
    document_ids = await service._ancestor_document_ids(
        _AncestorDB(),
        tenant_id=uuid.uuid4(),
        matter_id=uuid.uuid4(),
        revision=selected,
    )

    assert root_id in document_ids
    assert ancestor_output_id in document_ids
    assert sibling_output_id not in document_ids


@pytest.mark.asyncio
async def test_ambiguous_commit_preserves_superseded_output_artifact():
    tenant_id = uuid.uuid4()
    matter_id = uuid.uuid4()
    revision_id = uuid.uuid4()
    persisted = SimpleNamespace(
        id=revision_id,
        tenant_id=tenant_id,
        matter_id=matter_id,
        status="superseded",
        output_document_id=uuid.uuid4(),
        output_sha256="a" * 64,
    )

    class _ReconcileDB:
        async def rollback(self):
            return None

        async def execute(self, *_args, **_kwargs):
            return SimpleNamespace()

        async def scalar(self, _statement):
            return persisted

    class _NoDeleteStore:
        def __init__(self):
            self.deleted = False

        async def delete_stored_result(self, **_kwargs):
            self.deleted = True

    class _ReconcileService(MatterDocumentRevisionService):
        async def artifact(self, *_args, **_kwargs):
            return b"verified", SimpleNamespace()

        async def _to_response(self, *_args, **_kwargs):
            return "preserved"

    file_store = _NoDeleteStore()
    service = _ReconcileService(llm=SimpleNamespace(), file_store=file_store)
    result = await service._reconcile_persistence_failure(
        db=_ReconcileDB(),
        user=SimpleNamespace(tenant_id=tenant_id),
        matter_id=matter_id,
        revision_id=revision_id,
        stored=StorageResult(
            provider="local",
            backend="local",
            storage_path="memory://superseded.docx",
        ),
        cause=RuntimeError("commit acknowledgement lost"),
    )

    assert result == "preserved"
    assert file_store.deleted is False


@pytest.mark.asyncio
async def test_timed_out_worker_cannot_publish_after_state_changed():
    failed_row = SimpleNamespace(status="failed")

    class _StateDB:
        async def scalar(self, _statement):
            return failed_row

    service = MatterDocumentRevisionService(
        llm=SimpleNamespace(), file_store=SimpleNamespace()
    )
    with pytest.raises(DocumentRevisionServiceError) as error:
        await service._locked_processing_revision(
            _StateDB(), uuid.uuid4(), uuid.uuid4(), uuid.uuid4()
        )

    assert error.value.code == "revision_no_longer_processing"


@pytest.mark.asyncio
async def test_revision_lifecycle_followup_and_non_executable_esign_preview(
    db_session,
    test_tenant,
    test_user,
    monkeypatch,
):
    source_bytes = _docx_bytes("This agreement is between Jane Doe and Acme LLC.")
    source_path = "memory://source/client-agreement.docx"
    file_store = _MemoryFileStore(source_path, source_bytes)
    llm = _QueuedLLM(
        {
            "outcome": "change_plan",
            "summary": "Corrected the client's surname.",
            "warnings": [],
            "operations": [
                {
                    "type": "replace_text",
                    "block_id": "body/p/0",
                    "target_text": "Jane Doe",
                    "replacement_text": "Jane Roe",
                    "rationale": "Attorney-requested client name correction.",
                }
            ],
        },
        {
            "outcome": "change_plan",
            "summary": "Updated the company name.",
            "warnings": [],
            "operations": [
                {
                    "type": "replace_text",
                    "block_id": "body/p/0",
                    "target_text": "Acme LLC",
                    "replacement_text": "Acme Corporation",
                    "rationale": None,
                }
            ],
        },
        {
            "outcome": "needs_input",
            "question": "What effective date should replace the current date?",
        },
    )
    service = MatterDocumentRevisionService(llm=llm, file_store=file_store)

    async def no_budget_check(_db, _user):
        return None

    async def standard_route(_db, _tenant_id, **_kwargs):
        return LLMRoute(
            requested_route="standard",
            resolved_route="standard",
            gateway_alias="test-standard",
        )

    monkeypatch.setattr(revision_module, "check_token_budget", no_budget_check)
    monkeypatch.setattr(revision_module, "resolve_llm_route", standard_route)
    test_user.tenant = test_tenant

    matter = Matter(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        user_id=test_user.id,
        slug="jane-doe-contract",
        matter_name="Jane Doe Contract",
        matter_type="contract",
    )
    source = MatterDocument(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        matter_id=matter.id,
        uploaded_by_user_id=test_user.id,
        filename="Client Agreement.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument." "wordprocessingml.document"
        ),
        file_size=len(source_bytes),
        storage_path=source_path,
        storage_provider="local",
        storage_backend="local",
        document_category="contract",
        portal_visible=False,
    )
    db_session.add(matter)
    await db_session.flush()
    db_session.add(source)
    await db_session.commit()

    first_request = MatterDocumentRevisionCreate(
        instruction="Change Jane Doe to Jane Roe.",
        client_request_id=uuid.uuid4(),
        model_tier="auto",
    )
    first = await service.create_revision(
        db_session, test_user, matter.id, source.id, first_request
    )

    assert first.status == "ready_for_review"
    assert first.source_document_id == source.id
    assert first.root_document_id == source.id
    assert first.output_document_id is not None
    assert (
        first.output_sha256
        == hashlib.sha256(file_store.uploads[0]["content"]).hexdigest()
    )
    assert first.operations[0].rationale
    assert first.artifact_url.endswith(f"/{first.id}/artifact")
    output = await db_session.scalar(
        select(MatterDocument).where(MatterDocument.id == first.output_document_id)
    )
    assert output is not None
    assert output.portal_visible is False
    assert output.document_category == "assistant_revision"

    original_output_bytes = file_store.files[output.storage_path]
    file_store.files[output.storage_path] = original_output_bytes + b"tampered"
    with pytest.raises(DocumentRevisionServiceError) as tampered_artifact:
        await service.artifact(db_session, test_tenant.id, matter.id, first.id)
    assert tampered_artifact.value.code == "artifact_integrity_failed"
    file_store.files[output.storage_path] = original_output_bytes

    # Same idempotency key returns the persisted result without a second model call.
    retry = await service.create_revision(
        db_session, test_user, matter.id, source.id, first_request
    )
    assert retry.id == first.id
    assert len(llm.calls) == 1
    with pytest.raises(DocumentRevisionServiceError) as reused:
        await service.create_revision(
            db_session,
            test_user,
            matter.id,
            source.id,
            first_request.model_copy(update={"instruction": "Different request"}),
        )
    assert reused.value.code == "idempotency_key_reused"

    # Legacy portal-sharing and e-sign creation paths both fail closed.
    async def current_user(_request, _db):
        return test_user

    monkeypatch.setattr(matter_documents_router, "get_current_user", current_user)
    monkeypatch.setattr(esignature_router, "get_current_user", current_user)
    with pytest.raises(HTTPException) as portal_error:
        await matter_documents_router.update_matter_document(
            str(matter.id),
            str(output.id),
            MatterDocumentUpdate(portal_visible=True),
            SimpleNamespace(),
            db_session,
        )
    assert portal_error.value.status_code == 409
    assert (
        portal_error.value.detail["code"] == "assistant_revision_legacy_release_blocked"
    )

    with pytest.raises(HTTPException) as category_error:
        await matter_documents_router.update_matter_document(
            str(matter.id),
            str(output.id),
            MatterDocumentUpdate(document_category="contract"),
            SimpleNamespace(),
            db_session,
        )
    assert category_error.value.status_code == 409
    assert (
        category_error.value.detail["code"] == "assistant_revision_category_immutable"
    )

    with pytest.raises(HTTPException) as esign_error:
        await esignature_router.create_signature_request(
            str(matter.id),
            SignatureRequestCreate(
                document_id=str(output.id),
                signers=[SignerCreate(name="Jane Roe", email="jane@example.com")],
            ),
            SimpleNamespace(),
            db_session,
        )
    assert esign_error.value.status_code == 409
    assert (
        esign_error.value.detail["code"] == "assistant_revision_legacy_release_blocked"
    )

    for retained_document in (source, output):
        with pytest.raises(HTTPException) as delete_error:
            await matter_documents_router.delete_matter_document(
                str(matter.id),
                str(retained_document.id),
                SimpleNamespace(),
                db_session,
            )
        assert delete_error.value.status_code == 409
        assert delete_error.value.detail["code"] == "document_has_revision_lineage"

    wrong_hash = "0" * 64
    with pytest.raises(DocumentRevisionServiceError) as wrong_approval:
        await service.approve(
            db_session,
            test_user,
            matter.id,
            first.id,
            MatterDocumentRevisionApprove(reviewed_output_sha256=wrong_hash),
        )
    assert wrong_approval.value.code == "reviewed_hash_mismatch"

    followup = await service.create_revision(
        db_session,
        test_user,
        matter.id,
        output.id,
        MatterDocumentRevisionCreate(
            instruction="Change Acme LLC to Acme Corporation.",
            client_request_id=uuid.uuid4(),
        ),
    )
    assert followup.version_no == 2
    assert followup.root_document_id == source.id
    assert followup.source_document_id == output.id
    assert followup.source_revision_id == first.id
    history_page = await service.list_revisions(
        db_session,
        test_tenant.id,
        matter.id,
        source.id,
        limit=1,
        offset=0,
    )
    assert history_page.total == 2
    assert history_page.limit == 1
    assert len(history_page.items) == 1
    assert history_page.items[0].output_text_preview == []
    superseded = await service.get_revision(
        db_session, test_tenant.id, matter.id, first.id
    )
    assert superseded.status == "superseded"
    with pytest.raises(DocumentRevisionServiceError) as stale_approval:
        await service.approve(
            db_session,
            test_user,
            matter.id,
            first.id,
            MatterDocumentRevisionApprove(reviewed_output_sha256=first.output_sha256),
        )
    assert stale_approval.value.code == "invalid_revision_status"

    approved = await service.approve(
        db_session,
        test_user,
        matter.id,
        followup.id,
        MatterDocumentRevisionApprove(reviewed_output_sha256=followup.output_sha256),
    )
    assert approved.status == "approved"
    approved_retry = await service.approve(
        db_session,
        test_user,
        matter.id,
        followup.id,
        MatterDocumentRevisionApprove(reviewed_output_sha256=followup.output_sha256),
    )
    assert approved_retry.approved_at == approved.approved_at
    followup_output = await db_session.scalar(
        select(MatterDocument).where(MatterDocument.id == followup.output_document_id)
    )
    assert followup_output is not None
    with pytest.raises(DocumentRevisionServiceError) as approved_legacy_release:
        await assert_no_legacy_assistant_derivative_release(
            db_session,
            tenant_id=test_tenant.id,
            matter_id=matter.id,
            document_id=followup_output.id,
        )
    assert approved_legacy_release.value.code == (
        "assistant_revision_legacy_release_blocked"
    )
    with pytest.raises(HTTPException) as approved_portal_error:
        await matter_documents_router.update_matter_document(
            str(matter.id),
            str(followup_output.id),
            MatterDocumentUpdate(portal_visible=True),
            SimpleNamespace(),
            db_session,
        )
    assert approved_portal_error.value.detail["code"] == (
        "assistant_revision_legacy_release_blocked"
    )
    with pytest.raises(HTTPException) as approved_esign_error:
        await esignature_router.create_signature_request(
            str(matter.id),
            SignatureRequestCreate(
                document_id=str(followup_output.id),
                signers=[SignerCreate(name="Jane Roe", email="jane@example.com")],
            ),
            SimpleNamespace(),
            db_session,
        )
    assert approved_esign_error.value.detail["code"] == (
        "assistant_revision_legacy_release_blocked"
    )

    signature_request = SignatureRequest(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        matter_id=matter.id,
        document_id=source.id,
        status="sent",
        provider="internal",
        source_document_sha256=hashlib.sha256(source_bytes).hexdigest(),
        source_document_size=len(source_bytes),
        source_document_filename=source.filename,
        created_by_user_id=test_user.id,
    )
    signer = SignatureSigner(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        request_id=signature_request.id,
        name="Jane Roe",
        email="jane@example.com",
        role="client",
        sign_order=0,
        status="pending",
    )
    db_session.add(signature_request)
    await db_session.flush()
    db_session.add(signer)
    await db_session.commit()

    prepared = await service.prepare_esign_replacement(
        db_session,
        test_user,
        matter.id,
        followup.id,
        SignatureReplacementPrepare(signature_request_id=signature_request.id),
    )
    preview = prepared.prepared_esign_preview
    assert preview is not None
    assert preview.executable is False
    assert preview.notification_will_be_sent is False
    assert preview.replacement_document_id == followup_output.id
    await db_session.refresh(signature_request)
    await db_session.refresh(signer)
    assert signature_request.document_id == source.id
    assert signature_request.status == "sent"
    assert signer.status == "pending"

    expired_request = SignatureRequest(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        matter_id=matter.id,
        document_id=source.id,
        status="sent",
        provider="internal",
        source_document_sha256=hashlib.sha256(source_bytes).hexdigest(),
        source_document_size=len(source_bytes),
        source_document_filename=source.filename,
        created_by_user_id=test_user.id,
        expires_at=datetime.now(timezone.utc) - timedelta(minutes=1),
    )
    expired_signer = SignatureSigner(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        request_id=expired_request.id,
        name="Jane Roe",
        email="jane@example.com",
        role="client",
        sign_order=0,
        status="pending",
    )
    db_session.add_all([expired_request, expired_signer])
    await db_session.commit()
    with pytest.raises(DocumentRevisionServiceError) as expired_preview:
        await service.prepare_esign_replacement(
            db_session,
            test_user,
            matter.id,
            followup.id,
            SignatureReplacementPrepare(signature_request_id=expired_request.id),
        )
    assert expired_preview.value.code == "signature_request_expired"

    signer.status = "signed"
    await db_session.commit()
    stale_prepared = await service.get_revision(
        db_session, test_tenant.id, matter.id, followup.id
    )
    assert stale_prepared.prepared_esign_preview is None

    second_source_bytes = _docx_bytes("Effective date: [DATE]")
    second_path = "memory://source/effective-date.docx"
    file_store.files[second_path] = second_source_bytes
    second_source = MatterDocument(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        matter_id=matter.id,
        uploaded_by_user_id=test_user.id,
        filename="Effective Date.docx",
        content_type=source.content_type,
        file_size=len(second_source_bytes),
        storage_path=second_path,
        storage_provider="local",
        storage_backend="local",
        document_category="contract",
        portal_visible=False,
    )
    db_session.add(second_source)
    await db_session.commit()
    needs_input = await service.create_revision(
        db_session,
        test_user,
        matter.id,
        second_source.id,
        MatterDocumentRevisionCreate(
            instruction="Update the effective date.",
            client_request_id=uuid.uuid4(),
        ),
    )
    assert needs_input.status == "needs_input"
    assert needs_input.output_document_id is None
    assert needs_input.clarification_question
    assert len(file_store.uploads) == 2

    async def unavailable_route(_db, _tenant_id, **_kwargs):
        raise RuntimeError("routing unavailable")

    monkeypatch.setattr(revision_module, "resolve_llm_route", unavailable_route)
    routing_request_id = uuid.uuid4()
    with pytest.raises(DocumentRevisionServiceError) as routing_error:
        await service.create_revision(
            db_session,
            test_user,
            matter.id,
            second_source.id,
            MatterDocumentRevisionCreate(
                instruction="Try the effective-date update again.",
                client_request_id=routing_request_id,
            ),
        )
    assert routing_error.value.code == "model_routing_failed"
    failed_routing_row = await db_session.scalar(
        select(MatterDocumentRevision).where(
            MatterDocumentRevision.tenant_id == test_tenant.id,
            MatterDocumentRevision.client_request_id == routing_request_id,
        )
    )
    assert failed_routing_row is not None
    assert failed_routing_row.status == "failed"
    assert failed_routing_row.error_code == "model_routing_failed"

    stale_row = MatterDocumentRevision(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        matter_id=matter.id,
        root_document_id=second_source.id,
        source_document_id=second_source.id,
        requested_by_user_id=test_user.id,
        client_request_id=uuid.uuid4(),
        version_no=3,
        instruction="Stale processing request",
        status="processing",
        source_sha256=hashlib.sha256(second_source_bytes).hexdigest(),
        requested_model_tier="auto",
        created_at=datetime.now(timezone.utc) - timedelta(minutes=11),
        updated_at=datetime.now(timezone.utc) - timedelta(minutes=11),
    )
    db_session.add(stale_row)
    await db_session.commit()
    stale_response = await service.get_revision(
        db_session, test_tenant.id, matter.id, stale_row.id
    )
    assert stale_response.status == "failed"
    assert stale_response.error_code == "processing_timeout"

    usage_rows = list(
        (
            await db_session.execute(
                select(UsageRecord).where(
                    UsageRecord.tenant_id == test_tenant.id,
                    UsageRecord.operation_type == "document_revision",
                )
            )
        )
        .scalars()
        .all()
    )
    assert len(usage_rows) == 3
    assert all(row.query_text is None for row in usage_rows)


@pytest.mark.asyncio
async def test_premium_revision_requires_an_enabled_licensed_user(
    db_session,
    test_tenant,
    test_user,
    monkeypatch,
):
    source_bytes = _docx_bytes("Original text")
    file_store = _MemoryFileStore("memory://source/premium.docx", source_bytes)
    service = MatterDocumentRevisionService(llm=_QueuedLLM(), file_store=file_store)

    async def no_budget_check(_db, _user):
        return None

    monkeypatch.setattr(revision_module, "check_token_budget", no_budget_check)
    matter = Matter(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        user_id=test_user.id,
        slug="premium-test",
        matter_name="Premium Test",
        matter_type="general",
    )
    source = MatterDocument(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        matter_id=matter.id,
        uploaded_by_user_id=test_user.id,
        filename="Premium.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument." "wordprocessingml.document"
        ),
        file_size=len(source_bytes),
        storage_path="memory://source/premium.docx",
        storage_provider="local",
        storage_backend="local",
        portal_visible=False,
    )
    db_session.add(matter)
    await db_session.flush()
    db_session.add(source)
    await db_session.commit()

    with pytest.raises(DocumentRevisionServiceError) as error:
        await service.create_revision(
            db_session,
            test_user,
            matter.id,
            source.id,
            MatterDocumentRevisionCreate(
                instruction="Make a change",
                client_request_id=uuid.uuid4(),
                model_tier="premium",
            ),
        )
    assert error.value.status_code == 403
    assert error.value.code == "premium_model_not_enabled"
