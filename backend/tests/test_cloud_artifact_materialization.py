from __future__ import annotations

import io

from docx import Document

from app.services.cloud_artifact_materialization import (
    canonical_docx_filename,
    render_revision_docx,
)


def test_canonical_docx_filename_is_safe_and_deterministic() -> None:
    value = canonical_docx_filename(
        "  Client / privileged: memo?.docx  ", revision_no=3
    )
    assert value == "Client - privileged- memo-r3.docx"
    assert "/" not in value and "?" not in value
    assert value.endswith(".docx")


def test_renderer_emits_valid_docx_and_preserves_paragraphs() -> None:
    raw = render_revision_docx(
        title="Review memo", content="First paragraph\nSecond paragraph"
    )
    assert raw.startswith(b"PK")
    document = Document(io.BytesIO(raw))
    assert document.core_properties.title == "Review memo"
    assert [p.text for p in document.paragraphs][-2:] == [
        "First paragraph",
        "Second paragraph",
    ]


def test_renderer_strips_xml_illegal_control_characters() -> None:
    raw = render_revision_docx(title="Draft", content="safe\x00 text")
    document = Document(io.BytesIO(raw))
    assert "safe text" in [p.text for p in document.paragraphs]
