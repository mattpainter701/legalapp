from __future__ import annotations

import hashlib
from io import BytesIO
import zipfile

import pytest
from docx import Document

from app.services.document_revision_engine import (
    CAPABILITY_BOUNDED_TEXT_REVISION,
    MAX_OPERATIONS,
    MAX_REPLACEMENT_TEXT_CHARS,
    DocumentCapabilityError,
    DocumentOperationError,
    apply_docx_revision,
    inspect_docx,
)


def _save(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _simple_docx(text: str = "Alpha beta gamma delta") -> bytes:
    document = Document()
    document.add_paragraph(text)
    return _save(document)


def _rewrite_package(
    source: bytes,
    *,
    replacements: dict[str, bytes] | None = None,
    additions: dict[str, bytes] | None = None,
) -> bytes:
    replacements = replacements or {}
    additions = additions or {}
    output = BytesIO()
    with (
        zipfile.ZipFile(BytesIO(source)) as incoming,
        zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as outgoing,
    ):
        for info in incoming.infolist():
            outgoing.writestr(
                info, replacements.get(info.filename, incoming.read(info))
            )
        for name, content in additions.items():
            outgoing.writestr(name, content)
    return output.getvalue()


def _replace_part(source: bytes, name: str, transform) -> bytes:
    with zipfile.ZipFile(BytesIO(source)) as archive:
        original = archive.read(name)
    return _rewrite_package(source, replacements={name: transform(original)})


def _append_before(content: bytes, closing: bytes, addition: bytes) -> bytes:
    assert closing in content
    return content.replace(closing, addition + closing, 1)


def test_inspection_extracts_stable_body_table_header_and_footer_blocks() -> None:
    document = Document()
    document.add_paragraph("Body paragraph")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Outer cell"
    nested = table.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested cell"
    document.sections[0].header.paragraphs[0].text = "Firm header"
    document.sections[0].footer.paragraphs[0].text = "Confidential footer"
    source = _save(document)
    snapshot = bytes(source)

    first = inspect_docx(source, filename="agreement.docx")
    second = inspect_docx(source, filename="agreement.docx")

    assert source == snapshot
    assert first.capability == CAPABILITY_BOUNDED_TEXT_REVISION
    assert first.source_sha256 == hashlib.sha256(source).hexdigest()
    assert first.source_size == len(source)
    assert [block.block_id for block in first.blocks] == [
        block.block_id for block in second.blocks
    ]
    by_text = {block.text: block for block in first.blocks}
    assert by_text["Body paragraph"].block_id == "body/p/0"
    assert by_text["Outer cell"].block_id == "body/tbl/0/row/0/cell/0/p/0"
    assert (
        by_text["Nested cell"].block_id
        == "body/tbl/0/row/0/cell/0/tbl/0/row/0/cell/0/p/0"
    )
    assert by_text["Firm header"].kind == "header_paragraph"
    assert by_text["Firm header"].block_id.startswith("header/word/header")
    assert by_text["Confidential footer"].kind == "footer_paragraph"
    assert by_text["Confidential footer"].block_id.startswith("footer/word/footer")
    assert all(block.editable for block in by_text.values())


def test_revision_preserves_source_and_unaffected_run_formatting() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    prefix = paragraph.add_run("Retainer: ")
    prefix.italic = True
    amount_first = paragraph.add_run("$2,")
    amount_first.bold = True
    amount_second = paragraph.add_run("500")
    amount_second.bold = True
    suffix = paragraph.add_run(" due on signing.")
    suffix.underline = True
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Unchanged table text"
    document.sections[0].header.paragraphs[0].text = "Unchanged header"
    document.sections[0].footer.paragraphs[0].text = "Unchanged footer"
    source = _save(document)
    snapshot = bytes(source)

    result = apply_docx_revision(
        source,
        [
            {
                "type": "replace_text",
                "block_id": "body/p/0",
                "target_text": "$2,500",
                "replacement_text": "$3,000",
            }
        ],
        filename="engagement.docx",
    )

    assert source == snapshot
    assert result.source_sha256 == hashlib.sha256(source).hexdigest()
    assert result.output_sha256 == hashlib.sha256(result.output_bytes).hexdigest()
    assert result.output_bytes != source
    assert result.operation_count == 1
    assert result.changes[0].before_text == "Retainer: $2,500 due on signing."
    assert result.changes[0].after_text == "Retainer: $3,000 due on signing."

    reopened = Document(BytesIO(result.output_bytes))
    revised = reopened.paragraphs[0]
    assert revised.text == "Retainer: $3,000 due on signing."
    assert revised.runs[0].text == "Retainer: "
    assert revised.runs[0].italic is True
    assert revised.runs[1].text == "$3,000"
    assert revised.runs[1].bold is True
    assert revised.runs[-1].text == " due on signing."
    assert revised.runs[-1].underline is True
    assert reopened.tables[0].cell(0, 0).text == "Unchanged table text"
    assert reopened.sections[0].header.paragraphs[0].text == "Unchanged header"
    assert reopened.sections[0].footer.paragraphs[0].text == "Unchanged footer"

    before_ids = [block.block_id for block in inspect_docx(source).blocks]
    after_ids = [block.block_id for block in result.blocks]
    assert after_ids == before_ids


def test_non_overlapping_operations_share_one_source_snapshot() -> None:
    source = _simple_docx()
    result = apply_docx_revision(
        source,
        [
            {
                "type": "replace_text",
                "block_id": "body/p/0",
                "target_text": "Alpha",
                "replacement_text": "First",
            },
            {
                "type": "replace_text",
                "block_id": "body/p/0",
                "target_text": "delta",
                "replacement_text": "last",
            },
        ],
    )

    assert Document(BytesIO(result.output_bytes)).paragraphs[0].text == (
        "First beta gamma last"
    )
    assert len(result.changes) == 1
    assert result.changes[0].before_text == "Alpha beta gamma delta"
    assert result.changes[0].after_text == "First beta gamma last"


def test_revision_can_target_table_header_and_footer_blocks() -> None:
    document = Document()
    document.add_paragraph("Body")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Client: Jane Doe"
    document.sections[0].header.paragraphs[0].text = "Draft agreement"
    document.sections[0].footer.paragraphs[0].text = "Page confidentiality notice"
    source = _save(document)
    inspection = inspect_docx(source)
    block_id_by_text = {block.text: block.block_id for block in inspection.blocks}

    result = apply_docx_revision(
        source,
        [
            {
                "type": "replace_text",
                "block_id": block_id_by_text["Client: Jane Doe"],
                "target_text": "Jane Doe",
                "replacement_text": "Jane Q. Doe",
            },
            {
                "type": "replace_text",
                "block_id": block_id_by_text["Draft agreement"],
                "target_text": "Draft",
                "replacement_text": "Approved",
            },
            {
                "type": "replace_text",
                "block_id": block_id_by_text["Page confidentiality notice"],
                "target_text": "confidentiality",
                "replacement_text": "privileged",
            },
        ],
    )

    reopened = Document(BytesIO(result.output_bytes))
    assert reopened.tables[0].cell(0, 0).text == "Client: Jane Q. Doe"
    assert reopened.sections[0].header.paragraphs[0].text == "Approved agreement"
    assert reopened.sections[0].footer.paragraphs[0].text == ("Page privileged notice")


@pytest.mark.parametrize(
    ("operation_factory", "expected_code"),
    [
        (lambda _block: [], "empty_operations"),
        (lambda _block: {"type": "replace_text"}, "invalid_operations"),
        (lambda _block: ["replace"], "invalid_operation"),
        (
            lambda block: [
                {
                    "type": "replace_text",
                    "block_id": block,
                    "target_text": "Alpha",
                    "replacement_text": "First",
                    "surprise": True,
                }
            ],
            "invalid_operation_keys",
        ),
        (
            lambda block: [
                {
                    "type": 7,
                    "block_id": block,
                    "target_text": "Alpha",
                    "replacement_text": "First",
                }
            ],
            "invalid_operation_value",
        ),
        (
            lambda block: [
                {
                    "type": "insert_after",
                    "block_id": block,
                    "target_text": "Alpha",
                    "replacement_text": "First",
                }
            ],
            "unsupported_operation_type",
        ),
        (
            lambda block: [
                {
                    "type": "replace_text",
                    "block_id": block,
                    "target_text": 7,
                    "replacement_text": "First",
                }
            ],
            "invalid_operation_value",
        ),
        (
            lambda block: [
                {
                    "type": "replace_text",
                    "block_id": block,
                    "target_text": "Alpha",
                    "replacement_text": "Alpha",
                }
            ],
            "no_effect_operation",
        ),
        (
            lambda _block: [
                {
                    "type": "replace_text",
                    "block_id": "body/p/999",
                    "target_text": "Alpha",
                    "replacement_text": "First",
                }
            ],
            "unknown_block",
        ),
        (
            lambda block: [
                {
                    "type": "replace_text",
                    "block_id": block,
                    "target_text": "missing",
                    "replacement_text": "First",
                }
            ],
            "target_not_found",
        ),
        (
            lambda block: [
                {
                    "type": "replace_text",
                    "block_id": block,
                    "target_text": "Alpha",
                    "replacement_text": "x" * (MAX_REPLACEMENT_TEXT_CHARS + 1),
                }
            ],
            "replacement_text_too_large",
        ),
        (
            lambda block: [
                {
                    "type": "replace_text",
                    "block_id": block,
                    "target_text": "Alpha",
                    "replacement_text": "invalid\x00text",
                }
            ],
            "invalid_replacement_text",
        ),
        (
            lambda block: [
                {
                    "type": "replace_text",
                    "block_id": block,
                    "target_text": "Alpha",
                    "replacement_text": f"First {index}",
                }
                for index in range(MAX_OPERATIONS + 1)
            ],
            "too_many_operations",
        ),
    ],
)
def test_strict_operation_validation(operation_factory, expected_code: str) -> None:
    source = _simple_docx()
    with pytest.raises(DocumentOperationError) as caught:
        apply_docx_revision(source, operation_factory("body/p/0"))
    assert caught.value.code == expected_code
    assert caught.value.capability == CAPABILITY_BOUNDED_TEXT_REVISION
    assert caught.value.as_dict()["message"]


def test_ambiguous_target_is_rejected() -> None:
    source = _simple_docx("same value and same value")
    with pytest.raises(DocumentOperationError) as caught:
        apply_docx_revision(
            source,
            [
                {
                    "type": "replace_text",
                    "block_id": "body/p/0",
                    "target_text": "same value",
                    "replacement_text": "different value",
                }
            ],
        )
    assert caught.value.code == "ambiguous_target"


def test_overlapping_operations_are_rejected_before_mutation() -> None:
    source = _simple_docx("abcdef")
    snapshot = bytes(source)
    with pytest.raises(DocumentOperationError) as caught:
        apply_docx_revision(
            source,
            [
                {
                    "type": "replace_text",
                    "block_id": "body/p/0",
                    "target_text": "abc",
                    "replacement_text": "ABC",
                },
                {
                    "type": "replace_text",
                    "block_id": "body/p/0",
                    "target_text": "bcd",
                    "replacement_text": "BCD",
                },
            ],
        )
    assert caught.value.code == "overlapping_operations"
    assert source == snapshot


def _with_macro_part(source: bytes) -> bytes:
    return _rewrite_package(source, additions={"word/vbaProject.bin": b"macro"})


def _with_tracked_revision(source: bytes) -> bytes:
    return _replace_part(
        source,
        "word/document.xml",
        lambda content: _append_before(
            content,
            b"</w:body>",
            b'<w:ins w:id="1" w:author="tester"><w:r><w:t>tracked</w:t></w:r></w:ins>',
        ),
    )


def _with_protection(source: bytes) -> bytes:
    return _replace_part(
        source,
        "word/settings.xml",
        lambda content: _append_before(
            content,
            b"</w:settings>",
            b'<w:documentProtection w:edit="readOnly" w:enforcement="1"/>',
        ),
    )


def _with_signature_part(source: bytes) -> bytes:
    return _rewrite_package(
        source,
        additions={"_xmlsignatures/sig1.xml": b"<Signature />"},
    )


def _with_ole_part(source: bytes) -> bytes:
    return _rewrite_package(
        source, additions={"word/embeddings/oleObject1.bin": b"ole"}
    )


def _with_activex_part(source: bytes) -> bytes:
    return _rewrite_package(source, additions={"word/activeX/activeX1.bin": b"control"})


def _with_altchunk(source: bytes) -> bytes:
    return _replace_part(
        source,
        "word/document.xml",
        lambda content: _append_before(
            content,
            b"</w:body>",
            b'<w:altChunk r:id="rIdImportedContent"/>',
        ),
    )


def _with_external_relationship(source: bytes) -> bytes:
    return _replace_part(
        source,
        "word/_rels/document.xml.rels",
        lambda content: _append_before(
            content,
            b"</Relationships>",
            (
                b'<Relationship Id="rIdExternal" '
                b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" '
                b'Target="https://example.test/" TargetMode="External"/>'
            ),
        ),
    )


@pytest.mark.parametrize(
    ("mutator", "expected_code"),
    [
        (_with_macro_part, "macro_enabled_document"),
        (_with_tracked_revision, "tracked_revisions"),
        (_with_protection, "protected_document"),
        (_with_signature_part, "digitally_signed_document"),
        (_with_ole_part, "embedded_ole_content"),
        (_with_activex_part, "activex_content"),
        (_with_altchunk, "altchunk_content"),
        (_with_external_relationship, "external_relationship"),
    ],
)
def test_unsafe_docx_capabilities_are_rejected(mutator, expected_code: str) -> None:
    source = mutator(_simple_docx())
    with pytest.raises(DocumentCapabilityError) as caught:
        inspect_docx(source, filename="source.docx")
    assert caught.value.code == expected_code
    assert caught.value.as_dict()["capability"] != CAPABILITY_BOUNDED_TEXT_REVISION


def test_macro_filename_and_encrypted_or_legacy_container_are_rejected() -> None:
    source = _simple_docx()
    with pytest.raises(DocumentCapabilityError) as macro_error:
        inspect_docx(source, filename="agreement.docm")
    assert macro_error.value.code == "macro_enabled_document"

    with pytest.raises(DocumentCapabilityError) as encrypted_error:
        inspect_docx(bytes.fromhex("D0CF11E0A1B11AE1") + b"encrypted package")
    assert encrypted_error.value.code == "encrypted_or_legacy_word_document"


def test_source_must_be_immutable_bytes() -> None:
    mutable_source = bytearray(_simple_docx())
    with pytest.raises(DocumentCapabilityError) as caught:
        inspect_docx(mutable_source)  # type: ignore[arg-type]
    assert caught.value.code == "invalid_source_type"
