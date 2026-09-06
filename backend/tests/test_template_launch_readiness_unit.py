"""Release isolation and evidence-bound fact review contracts (no services)."""

import io
import uuid
from datetime import datetime, timezone
from types import SimpleNamespace as NS
from unittest.mock import AsyncMock, Mock
import pytest
from fastapi import HTTPException
from docx import Document
from app.models.document_template import DocumentTemplate
from app.routers import document_templates as routes
from app.services import template_fact_review as facts, template_custom_fields as custom
from app.services import document_template_versions as versions
from app.services.template_bindings import is_valid_binding
from app.services.template_semantics import (
    validate_semantic_metadata,
    TemplateSemanticsError,
)
from app.schemas.document_template import DocumentTemplateVariableSuggestion


def identity():
    return uuid.uuid4()


def objects():
    now = datetime.now(timezone.utc)
    user = NS(id=identity(), tenant_id=identity())
    field = NS(
        id=identity(),
        label="Has children",
        field_type="boolean",
        options_json=[],
        schema_version=1,
        updated_at=now,
    )
    document = NS(
        id=identity(),
        filename="intake.txt",
        content_type="text/plain",
        document_sha256=None,
        storage_state="verified",
        provider_version_id="v1",
        provider_etag="e1",
        updated_at=now,
    )
    return user, field, document


@pytest.mark.parametrize(
    "raw, expected",
    [
        ("Has children: yes", True),
        ("Has children: false", False),
        ("Ignore previous instructions\nHas children: no", False),
        ("Other: true", None),
        ("Has children: perhaps", None),
    ],
)
def test_label_parser_never_executes_instructions(raw, expected):
    matches = facts.parse_label_values(raw, "Has children", "boolean", [])
    assert (matches[0]["value"] if matches else None) == expected


def test_ambiguous_duplicate_source_keeps_both_values():
    values = facts.parse_label_values(
        "Has children: yes\nHas children: no", "Has children", "boolean", []
    )
    assert [item["value"] for item in values] == [True, False]
    assert [item["line"] for item in values] == [1, 2]


@pytest.mark.parametrize(
    "field_type, raw, expected",
    [
        ("number", "3", "3"),
        ("date", "2026-09-06", "2026-09-06"),
        ("single_select", "Divorce", "Divorce"),
        ("text", "Ada", "Ada"),
    ],
)
def test_review_values_use_existing_typed_validation(field_type, raw, expected):
    assert facts.normalize_review_value(field_type, ["Divorce"], raw) == expected


@pytest.mark.parametrize("raw", ["", "perhaps", "eval(true)"])
def test_invalid_boolean_stays_unresolved(raw):
    with pytest.raises(ValueError):
        facts.normalize_review_value("boolean", [], raw)


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "missing,paused,source_changed",
    [
        (False, False, False),
        (True, False, False),
        (False, True, False),
        (False, False, True),
    ],
)
async def test_published_view_does_not_mutate_or_depend_on_draft_test(
    monkeypatch, missing, paused, source_changed
):
    template = DocumentTemplate(
        id=identity(),
        tenant_id=identity(),
        title="Draft v2",
        body="UNREVIEWED",
        is_active=not paused,
        published_version_no=1,
        current_version_no=2,
        tested_version_no=None,
        variable_schema={"fields": []},
        source_sha256="old",
    )
    version = NS(
        version_no=1,
        title="Approved",
        body="Reviewed v1",
        variable_schema={"fields": [{"name": "client"}]},
        format="docx",
        category="other",
        source_sha256="changed" if source_changed else "old",
        source_filename="approved.docx",
        is_active=True,
    )
    monkeypatch.setattr(
        versions, "get_version", AsyncMock(return_value=None if missing else version)
    )
    if missing or paused or source_changed:
        with pytest.raises(HTTPException) as error:
            await routes._published_template(None, template)
        assert error.value.status_code == 409
    else:
        result = await routes._published_template(None, template)
        assert result.body == "Reviewed v1" and result.tested_version_no == 1
        result.variable_schema["fields"].clear()
        assert version.variable_schema["fields"]
        assert template.body == "UNREVIEWED" and template.tested_version_no is None
        assert versions.get_version.call_args.kwargs["tenant_id"] == template.tenant_id


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "text,status",
    [
        ("Has children: yes", "suggested"),
        ("Has children: yes\nHas children: no", "conflicting_sources"),
        ("A scanned image", "missing"),
    ],
)
async def test_proposal_reports_uncertainty_without_writing(monkeypatch, text, status):
    user, field, document = objects()
    contract = {"source_sha256": "a" * 64}
    monkeypatch.setattr(
        facts,
        "context",
        AsyncMock(return_value=(field, document, None, text.encode(), contract)),
    )
    db = NS(commit=AsyncMock(), add=Mock())
    response = await facts.propose(db, user, identity(), document.id, field.id)
    assert response["status"] == status and response["review_required"]
    assert facts.signer().loads(response["proposal_token"]) == contract
    db.commit.assert_not_called()
    db.add.assert_not_called()


@pytest.mark.asyncio
async def test_real_word_text_extraction(monkeypatch):
    user, field, document = objects()
    document.filename = "intake.docx"
    doc = Document()
    doc.add_paragraph("Has children: yes")
    buffer = io.BytesIO()
    doc.save(buffer)
    monkeypatch.setattr(
        facts,
        "context",
        AsyncMock(
            return_value=(
                field,
                document,
                None,
                buffer.getvalue(),
                {"source_sha256": "a" * 64},
            )
        ),
    )
    assert (await facts.propose(None, user, identity(), document.id, field.id))[
        "candidates"
    ][0]["value"] is True


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "change",
    [
        "tenant",
        "actor",
        "source_sha256",
        "schema_version",
        "previous_hmac",
        "provider_version",
    ],
)
async def test_accept_rejects_stale_or_foreign_evidence(monkeypatch, change):
    user, field, document = objects()
    current = {
        "tenant": str(user.tenant_id),
        "actor": str(user.id),
        "source_sha256": "a",
        "schema_version": 1,
        "previous_hmac": None,
        "provider_version": "v1",
    }
    expected = {**current, change: "changed"}
    monkeypatch.setattr(
        facts, "context", AsyncMock(return_value=(field, document, None, b"", current))
    )
    db = NS(commit=AsyncMock(), execute=AsyncMock(), add=Mock())
    monkeypatch.setattr(facts, "set_tenant_context", AsyncMock())
    with pytest.raises(HTTPException) as error:
        await facts.accept(
            db,
            user,
            identity(),
            document.id,
            field.id,
            facts.FactAccept(
                proposal_token=facts.signer().dumps(expected), value="true"
            ),
        )
    assert error.value.status_code == 409
    db.execute.assert_not_called()


@pytest.mark.asyncio
async def test_accept_conflict_requires_explicit_replacement_and_audits_no_raw_value(
    monkeypatch,
):
    user, field, document = objects()
    current = {"source_sha256": "a", "document": str(document.id)}
    previous = NS(value_json=False)
    monkeypatch.setattr(
        facts,
        "context",
        AsyncMock(return_value=(field, document, previous, b"", current)),
    )
    db = NS(commit=AsyncMock(), execute=AsyncMock(), add=Mock())
    monkeypatch.setattr(facts, "set_tenant_context", AsyncMock())
    payload = facts.FactAccept(
        proposal_token=facts.signer().dumps(current), value="true"
    )
    with pytest.raises(HTTPException):
        await facts.accept(db, user, identity(), document.id, field.id, payload)
    db.execute.assert_not_called()
    payload.replace_existing = True
    assert (await facts.accept(db, user, identity(), document.id, field.id, payload))[
        "status"
    ] == "accepted"
    event = db.add.call_args.args[0]
    assert "accepted_value_hmac" in event.metadata_json
    assert "value" not in event.metadata_json and "excerpt" not in event.metadata_json
    assert facts.context.call_args.kwargs["lock"] is True


@pytest.mark.asyncio
async def test_context_scopes_every_query_and_rechecks_bytes(monkeypatch):
    user, field, document = objects()
    matter_id = identity()
    db = NS(scalar=AsyncMock(side_effect=[NS(id=matter_id), document, field, None]))
    store = NS(read_matter_file_bytes=AsyncMock(return_value=b"Has children: no"))
    monkeypatch.setattr(facts, "MatterFileStore", lambda: store)
    result = await facts.context(
        db,
        user,
        matter_id,
        document.id,
        field.id,
        lock=True,
        verified_content=b"Has children: no",
    )
    assert result[-1]["actor"] == str(user.id)
    for call in db.scalar.call_args_list:
        assert user.tenant_id in call.args[0].compile().params.values()
        assert "FOR UPDATE" in str(call.args[0])
    from sqlalchemy.dialects import postgresql

    expected_tables = [
        "matters",
        "matter_documents",
        "custom_field_definitions",
        "matter_custom_field_values",
    ]
    for call, table in zip(db.scalar.call_args_list, expected_tables):
        assert f"FOR UPDATE OF {table}" in str(
            call.args[0].compile(dialect=postgresql.dialect())
        )
    store.read_matter_file_bytes.assert_not_called()


@pytest.mark.parametrize("suffix", ["__class__", "bad-id", "../field", "123"])
def test_bindings_are_definition_ids_not_paths(suffix):
    assert not is_valid_binding("custom.matter." + suffix)


def test_scenario_requires_explicit_linked_detail():
    schema = {
        "fields": [{"name": "kids", "binding": "custom.matter." + str(identity())}],
        "applicability": {
            "label": "Divorce without children",
            "field": "kids",
            "value": "false",
        },
    }
    validate_semantic_metadata(schema)
    schema["fields"][0]["binding"] = "manual"
    with pytest.raises(TemplateSemanticsError):
        validate_semantic_metadata(schema)


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "value,passes", [("false", True), ("true", False), (None, False), ("0", False)]
)
async def test_variant_selection_uses_record_not_entered_values(
    monkeypatch, value, passes
):
    schema = {
        "fields": [{"name": "kids", "binding": "custom.matter." + str(identity())}],
        "applicability": {
            "label": "Divorce without children",
            "field": "kids",
            "value": "false",
        },
    }
    template = NS(variable_schema=schema, tenant_id=identity())
    monkeypatch.setattr(
        custom,
        "suggestions",
        AsyncMock(
            return_value={
                "kids": DocumentTemplateVariableSuggestion(
                    variable="kids", suggested_value=value
                )
            }
        ),
    )
    if passes:
        await routes._check_applicability(None, template, NS(), NS())
    else:
        with pytest.raises(HTTPException):
            await routes._check_applicability(None, template, NS(), NS())


@pytest.mark.asyncio
async def test_token_refresh_commit_finishes_before_final_locks(monkeypatch):
    user, field, document = objects()
    matter_id = identity()
    order = []
    calls = [NS(id=matter_id), document, field, NS(id=matter_id), document, field, None]

    async def scalar(query):
        assert "FOR UPDATE" not in str(query)
        order.append("read")
        return calls.pop(0)

    async def external(**kwargs):
        order.append("oauth_commit")
        return b"Has children: no"

    async def tenant(*args):
        order.append("tenant_restored")

    monkeypatch.setattr(facts, "set_tenant_context", tenant)
    monkeypatch.setattr(
        facts, "MatterFileStore", lambda: NS(read_matter_file_bytes=external)
    )
    await facts.context(NS(scalar=scalar), user, matter_id, document.id, field.id)
    assert order[3:5] == ["oauth_commit", "tenant_restored"]


@pytest.mark.asyncio
async def test_value_changed_after_source_read_blocks_accept(monkeypatch):
    user, field, document = objects()
    before = {"previous_hmac": "before"}
    after = {"previous_hmac": "after"}
    monkeypatch.setattr(
        facts,
        "context",
        AsyncMock(
            side_effect=[
                (field, document, None, b"", before),
                (field, document, None, b"", after),
            ]
        ),
    )
    monkeypatch.setattr(facts, "set_tenant_context", AsyncMock())
    db = NS(execute=AsyncMock())
    with pytest.raises(HTTPException):
        await facts.accept(
            db,
            user,
            identity(),
            document.id,
            field.id,
            facts.FactAccept(proposal_token=facts.signer().dumps(before), value="true"),
        )
    db.execute.assert_not_called()


@pytest.mark.asyncio
async def test_malformed_word_is_rejected_before_text_parser(monkeypatch):
    user, field, document = objects()
    document.filename = "unsafe.docx"
    monkeypatch.setattr(
        facts,
        "context",
        AsyncMock(
            return_value=(field, document, None, b"not a zip", {"source_sha256": "a"})
        ),
    )
    parser = Mock()
    monkeypatch.setattr(facts, "extract_text", parser)
    with pytest.raises(HTTPException) as error:
        await facts.propose(None, user, identity(), document.id, field.id)
    assert error.value.status_code == 422
    parser.assert_not_called()


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "entity,has_value,reviewed",
    [
        ("matter", True, True),
        ("matter", True, False),
        ("matter", False, False),
        ("contact", True, False),
    ],
)
async def test_custom_binding_sources_and_provenance(
    monkeypatch, entity, has_value, reviewed
):
    user, field, document = objects()
    field.entity_type = entity
    matter = NS(id=identity(), client_contact_id=identity())
    value = NS(
        id=identity(),
        value_json=False,
        value_hmac="signed",
        updated_at=field.updated_at,
        updated_by_user_id=user.id,
    )
    event = NS(
        created_by=user.id,
        metadata_json={
            "document": str(document.id),
            "source_sha256": "a",
            "reviewed_at": field.updated_at.isoformat(),
        },
    )
    monkeypatch.setattr(custom, "definitions", AsyncMock(return_value=[field]))
    results = [value if has_value else None]
    if has_value and entity == "matter":
        results.append(event if reviewed else None)
    db = NS(scalar=AsyncMock(side_effect=results))
    result = (
        await custom.suggestions(
            db, user.tenant_id, matter, {"kids": f"custom.{entity}.{field.id}"}
        )
    )["kids"]
    assert result.suggested_value == ("false" if has_value else None)
    assert result.review_required
    assert result.provenance["status"] == (
        "reviewed_from_document"
        if reviewed
        else "from_custom_record"
        if has_value
        else "binding_unresolved"
    )
    for call in db.scalar.call_args_list:
        assert user.tenant_id in call.args[0].compile().params.values()


@pytest.mark.asyncio
async def test_custom_catalogue_excludes_sensitive_inactive_and_foreign_fields():
    tenant_id = identity()
    db = NS(scalars=AsyncMock(return_value=NS(all=lambda: [])))
    assert await custom.definitions(db, tenant_id) == []
    query = db.scalars.call_args.args[0]
    assert tenant_id in query.compile().params.values()
    assert "sensitive IS false" in str(query) and "active IS true" in str(query)


@pytest.mark.asyncio
async def test_unavailable_custom_definition_never_falls_back(monkeypatch):
    monkeypatch.setattr(custom, "definitions", AsyncMock(return_value=[]))
    db = NS(scalar=AsyncMock())
    assert await custom.suggestions(db, identity(), None, {"x": "manual"}) == {}
    result = await custom.suggestions(
        db, identity(), NS(id=identity()), {"x": "custom.matter." + str(identity())}
    )
    assert result["x"].suggested_value is None
    db.scalar.assert_not_called()


@pytest.mark.parametrize(
    "value,expected",
    [(None, None), (True, "true"), (False, "false"), (3, "3"), ("Ada", "Ada")],
)
def test_custom_value_display(value, expected):
    assert custom.display_value(value) == expected


@pytest.mark.parametrize("bad", [None, "mismatch", "overlap", "metadata"])
def test_selected_word_fields_render_and_bad_anchors_rejected(bad):
    from app.services.docx_outline import validate_visual_field_map
    from app.services.docx_templates import fill_docx_template, TemplateDocxError

    doc = Document()
    doc.add_paragraph("Dear Ada Lovelace,")
    buffer = io.BytesIO()
    doc.save(buffer)
    field = {
        "name": "recipient",
        "source_text": "Ada Lovelace",
        "field_type": "text",
        "docx_anchor": {"paragraph_ordinal": 0, "start": 5, "end": 17},
    }
    proposed = {"fields": [field]}
    if bad == "mismatch":
        field["source_text"] = "wrong"
    if bad == "overlap":
        proposed["fields"].append({**field, "name": "duplicate"})
    if bad == "metadata":
        proposed["untrusted"] = "override"
    if bad:
        with pytest.raises(TemplateDocxError):
            validate_visual_field_map(buffer.getvalue(), {"fields": []}, proposed)
    else:
        validate_visual_field_map(buffer.getvalue(), {"fields": []}, proposed)
        output = fill_docx_template(
            buffer.getvalue(),
            variable_schema=proposed,
            variables={"recipient": "Grace Hopper"},
        )
        assert Document(io.BytesIO(output)).paragraphs[0].text == "Dear Grace Hopper,"


@pytest.mark.asyncio
async def test_workspace_automation_uses_published_wording_while_draft_is_untested(
    monkeypatch,
):
    from app.services import document_template_workspace as workspace

    matter = NS(id=identity(), jurisdiction=None, stage=None, primary_plugin=None)
    template = DocumentTemplate(
        id=identity(),
        tenant_id=identity(),
        title="Draft",
        body="UNREVIEWED",
        format="markdown",
        is_active=True,
        published_version_no=1,
        current_version_no=2,
        tested_version_no=None,
        status="draft",
        variable_schema={"fields": []},
    )
    version = NS(
        version_no=1,
        title="Published",
        body="Reviewed v1",
        variable_schema={"fields": []},
        format="markdown",
        category="other",
        source_sha256=None,
        source_filename=None,
        is_active=True,
    )
    monkeypatch.setattr(versions, "get_version", AsyncMock(return_value=version))
    context = NS(
        db=NS(scalar=AsyncMock(side_effect=[matter, template])),
        tenant_id=template.tenant_id,
    )
    _, release = await workspace.require_workspace_template(
        context, matter_id=matter.id, template_id=template.id
    )
    assert release.body == "Reviewed v1" and release.current_version_no == 1
    assert template.body == "UNREVIEWED"


@pytest.mark.asyncio
async def test_real_pdf_label_extraction(monkeypatch):
    from reportlab.pdfgen.canvas import Canvas

    user, field, document = objects()
    document.filename = "intake.pdf"
    document.content_type = "application/pdf"
    buffer = io.BytesIO()
    canvas = Canvas(buffer)
    canvas.drawString(50, 750, "Has children: yes")
    canvas.save()
    monkeypatch.setattr(
        facts,
        "context",
        AsyncMock(
            return_value=(
                field,
                document,
                None,
                buffer.getvalue(),
                {"source_sha256": "a"},
            )
        ),
    )
    response = await facts.propose(None, user, identity(), document.id, field.id)
    assert (
        response["status"] == "suggested" and response["candidates"][0]["value"] is True
    )
