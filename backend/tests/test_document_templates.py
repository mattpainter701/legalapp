import json
import uuid
from datetime import datetime, timedelta, timezone
from io import BytesIO
from decimal import Decimal
from types import SimpleNamespace

import pytest
from fastapi import HTTPException

from app.routers import document_templates
from app.services.template_intake import analyze_template_upload
from app.services.pdf_templates import (
    TemplatePdfError,
    discover_pdf_fields,
    fill_pdf_template,
    pdf_page_metadata,
    render_pdf_page_preview,
    validate_representative_pdf_variables,
)
from app.services.docx_templates import (
    TemplateDocxError,
    fill_docx_template,
    validate_docx_package,
)
from app.services.template_ai_assist import (
    AiFieldProposal,
    reconcile_ai_template_fields,
)
from app.services.template_ai_service import (
    TemplateAiAssistError,
    _redact_evidence,
    assist_template_mapping as run_template_ai_mapping,
)
from app.services.template_ocr import (
    OcrLine,
    PdfOcrResult,
    TemplateOcrError,
    image_to_pdf,
)


def _fillable_pdf() -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 720, "Client:")
    pdf.acroForm.textfield(name="Client Name", x=120, y=705, width=250, height=24)
    pdf.drawString(72, 680, "Approved:")
    pdf.acroForm.checkbox(name="Approved", x=130, y=670, size=15, fieldFlags="")
    pdf.drawString(72, 645, "Notes:")
    pdf.acroForm.textfield(
        name="Notes",
        x=120,
        y=585,
        width=250,
        height=50,
        fieldFlags="multiline",
    )
    pdf.save()
    return output.getvalue()


def _special_fields_pdf() -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 720, "Forum:")
    pdf.acroForm.choice(
        name="forum",
        value="State",
        options=[("State Court", "State"), ("Federal Court", "Federal")],
        x=130,
        y=705,
        width=180,
        height=24,
    )
    pdf.drawString(72, 670, "Filing type:")
    pdf.acroForm.radio(
        name="filing_type", value="Complaint", selected=True, x=150, y=660
    )
    pdf.acroForm.radio(name="filing_type", value="Motion", selected=False, x=270, y=660)
    pdf.save()
    return output.getvalue()


def _required_pdf() -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 740, "Required client review")
    pdf.acroForm.textfield(
        name="client_name",
        x=120,
        y=705,
        width=250,
        height=24,
        fieldFlags="required",
    )
    pdf.acroForm.checkbox(name="accepted", x=120, y=670, size=15, fieldFlags="required")
    pdf.save()
    return output.getvalue()


async def _grant_manage_documents(db_session, test_tenant, test_user) -> None:
    from app.models.rbac import Role, UserRole

    role = Role(
        tenant_id=test_tenant.id,
        name="Document managers",
        capabilities=["manage_documents"],
    )
    db_session.add(role)
    await db_session.flush()
    db_session.add(
        UserRole(
            user_id=test_user.id,
            role_id=role.id,
            tenant_id=test_tenant.id,
            source="manual",
        )
    )
    await db_session.commit()


@pytest.mark.asyncio
async def test_template_upload_read_is_bounded_before_size_rejection(monkeypatch):
    class OversizedUpload:
        filename = "oversized.docx"
        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        def __init__(self):
            self.requested_size = None

        async def read(self, size=-1):
            self.requested_size = size
            return b"x" * size

    upload = OversizedUpload()
    monkeypatch.setattr(document_templates.settings, "MAX_FILE_SIZE_MB", 1)

    with pytest.raises(HTTPException) as exc_info:
        await document_templates._read_template_sample(upload)

    assert getattr(exc_info.value, "status_code", None) == 413
    assert upload.requested_size == (1024 * 1024) + 1


async def _prepare_active_pdf_generation(
    *,
    client,
    db_session,
    test_tenant,
    test_user,
    tmp_path,
    monkeypatch,
    slug: str,
):
    from app.models.plugin import Matter
    from app.services import matter_file_store as matter_store_module

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    monkeypatch.setattr(matter_store_module.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)
    created = await client.post(
        "/api/templates/intake/create",
        files={"file": (f"{slug}.pdf", _fillable_pdf(), "application/pdf")},
        data={"title": f"{slug} form", "category": "other"},
    )
    assert created.status_code == 201, created.text
    template_id = created.json()["id"]
    activation = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": {
                "client_name": "Representative Client",
                "approved": "yes",
                "notes": "Representative narrative",
            },
            "preview_purpose": "activation",
        },
    )
    assert activation.status_code == 200, activation.text
    activated = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert activated.status_code == 200, activated.text
    matter = Matter(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        user_id=test_user.id,
        slug=slug,
        matter_name=f"{slug} Matter",
        matter_type="general",
    )
    db_session.add(matter)
    await db_session.commit()
    values = {
        "client_name": "Ada Lovelace",
        "approved": "yes",
        "notes": "Reviewed narrative",
    }
    generation = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": values,
            "matter_id": str(matter.id),
            "preview_purpose": "generation",
        },
    )
    assert generation.status_code == 200, generation.text
    return template_id, matter, values, generation.headers["x-clarity-preview-id"]


def test_render_template_preserves_unknown_variables():
    rendered = document_templates.render_template(
        "Dear {{ client_name }}, matter {{case_number}} remains {{unknown}}.",
        {"client_name": "Ada Lovelace", "case_number": "2026-CV-100"},
    )

    assert rendered == "Dear Ada Lovelace, matter 2026-CV-100 remains {{unknown}}."


def test_extract_template_variables_preserves_first_seen_order():
    variables = document_templates.extract_template_variables(
        "{{client_name}} {{ case_number }} {{client_name}} {{ attorney_email }}"
    )

    assert variables == ["client_name", "case_number", "attorney_email"]


def test_pdf_preview_value_evidence_is_keyed_and_contains_no_raw_values(monkeypatch):
    variables = {"client_name": "Ada", "approved": "false"}
    first = document_templates._pdf_values_hmac_sha256(
        variables=variables,
        flatten_pdf=True,
        matter_id=None,
    )
    unkeyed = document_templates._canonical_sha256(
        {"variables": variables, "flatten_pdf": True, "matter_id": None}
    )
    monkeypatch.setattr(
        document_templates.settings,
        "SECRET_KEY",
        "different-preview-evidence-test-secret",
    )
    second = document_templates._pdf_values_hmac_sha256(
        variables=variables,
        flatten_pdf=True,
        matter_id=None,
    )

    assert len(first) == 64
    assert "Ada" not in first
    assert first != unkeyed
    assert second != first


@pytest.mark.asyncio
async def test_pdf_preview_trimming_never_deletes_consumed_evidence(
    db_session, test_tenant, test_user
):
    from sqlalchemy import func, select

    from app.models.document_template import DocumentTemplate
    from app.models.document_template_preview import DocumentTemplatePreview

    template = DocumentTemplate(
        tenant_id=test_tenant.id,
        title="Retention form",
        body="",
        format="pdf",
        is_active=True,
    )
    db_session.add(template)
    await db_session.flush()
    now = datetime.now(timezone.utc)
    consumed = DocumentTemplatePreview(
        tenant_id=test_tenant.id,
        template_id=template.id,
        previewed_by_user_id=test_user.id,
        purpose="activation",
        contract_sha256="a" * 64,
        values_hmac_sha256="b" * 64,
        output_sha256="c" * 64,
        renderer_version="retention-test",
        flatten_pdf=True,
        reviewed_field_count=0,
        nonblank_field_count=0,
        reviewed_field_names=[],
        created_at=now - timedelta(days=100),
        expires_at=now - timedelta(days=99),
        consumed_at=now - timedelta(days=98),
    )
    attempts = [
        DocumentTemplatePreview(
            tenant_id=test_tenant.id,
            template_id=template.id,
            previewed_by_user_id=test_user.id,
            purpose="activation",
            contract_sha256=f"{index:064x}",
            values_hmac_sha256="d" * 64,
            output_sha256="e" * 64,
            renderer_version="retention-test",
            flatten_pdf=True,
            reviewed_field_count=0,
            nonblank_field_count=0,
            reviewed_field_names=[],
            created_at=now - timedelta(minutes=55 - index),
            expires_at=(
                now - timedelta(minutes=1) if index < 30 else now + timedelta(hours=1)
            ),
        )
        for index in range(55)
    ]
    db_session.add_all([consumed, *attempts])
    await db_session.commit()

    await document_templates._trim_preview_evidence(
        db_session,
        tenant_id=test_tenant.id,
        template_id=template.id,
        user_id=test_user.id,
        purpose="activation",
    )
    await db_session.commit()

    assert await db_session.get(DocumentTemplatePreview, consumed.id) is not None
    unconsumed_count = await db_session.scalar(
        select(func.count())
        .select_from(DocumentTemplatePreview)
        .where(
            DocumentTemplatePreview.template_id == template.id,
            DocumentTemplatePreview.consumed_at.is_(None),
        )
    )
    assert unconsumed_count == 50


@pytest.mark.asyncio
async def test_generation_trim_preserves_recent_inflight_window_and_removes_old_expiry(
    db_session, test_tenant, test_user
):
    from sqlalchemy import select

    from app.models.document_template import DocumentTemplate
    from app.models.document_template_preview import DocumentTemplatePreview

    template = DocumentTemplate(
        tenant_id=test_tenant.id,
        title="Generation trim grace",
        body="",
        format="pdf",
    )
    db_session.add(template)
    await db_session.flush()
    now = datetime.now(timezone.utc)

    def attempt(marker: str, expires_at: datetime):
        return DocumentTemplatePreview(
            tenant_id=test_tenant.id,
            template_id=template.id,
            previewed_by_user_id=test_user.id,
            purpose="generation",
            contract_sha256=marker * 64,
            values_hmac_sha256="b" * 64,
            output_sha256="c" * 64,
            renderer_version="generation-trim-test",
            flatten_pdf=True,
            reviewed_field_count=0,
            nonblank_field_count=0,
            reviewed_field_names=[],
            expires_at=expires_at,
        )

    recent = [
        attempt(f"{index % 10}", now - timedelta(minutes=1)) for index in range(55)
    ]
    old = attempt("f", now - timedelta(hours=2))
    db_session.add_all([*recent, old])
    await db_session.commit()

    await document_templates._trim_preview_evidence(
        db_session,
        tenant_id=test_tenant.id,
        template_id=template.id,
        user_id=test_user.id,
        purpose="generation",
    )
    await db_session.commit()

    remaining_ids = set(
        await db_session.scalars(
            select(DocumentTemplatePreview.id).where(
                DocumentTemplatePreview.template_id == template.id
            )
        )
    )
    assert old.id not in remaining_ids
    assert {row.id for row in recent}.issubset(remaining_ids)


@pytest.mark.asyncio
async def test_pdf_preview_cannot_be_consumed_and_unresolved_reconciliation(
    db_session, test_tenant, test_user
):
    from sqlalchemy.exc import IntegrityError

    from app.models.document_template import DocumentTemplate
    from app.models.document_template_preview import DocumentTemplatePreview

    template = DocumentTemplate(
        tenant_id=test_tenant.id,
        title="Terminal state constraint",
        body="",
        format="pdf",
    )
    db_session.add(template)
    await db_session.flush()
    now = datetime.now(timezone.utc)
    db_session.add(
        DocumentTemplatePreview(
            tenant_id=test_tenant.id,
            template_id=template.id,
            previewed_by_user_id=test_user.id,
            purpose="generation",
            contract_sha256="a" * 64,
            values_hmac_sha256="b" * 64,
            output_sha256="c" * 64,
            renderer_version="terminal-state-test",
            flatten_pdf=True,
            reviewed_field_count=0,
            nonblank_field_count=0,
            reviewed_field_names=[],
            expires_at=now + timedelta(hours=1),
            consumed_at=now,
            reconciliation_required_at=now,
            reconciliation_reason="cleanup_failed",
        )
    )

    with pytest.raises(IntegrityError):
        await db_session.commit()
    await db_session.rollback()


@pytest.mark.asyncio
async def test_ambiguous_commit_absence_uses_two_second_observation_window(monkeypatch):
    sessions = []
    sleeps = []

    class VerificationSession:
        async def scalar(self, _statement):
            return None

    class VerificationContext:
        async def __aenter__(self):
            session = VerificationSession()
            sessions.append(session)
            return session

        async def __aexit__(self, *_args):
            return None

    async def set_context(_db, _tenant_id):
        return None

    async def record_sleep(seconds):
        sleeps.append(seconds)

    monkeypatch.setattr(document_templates, "async_session_maker", VerificationContext)
    monkeypatch.setattr(document_templates, "set_tenant_context", set_context)
    monkeypatch.setattr(document_templates.asyncio, "sleep", record_sleep)

    outcome = await document_templates._matter_document_commit_outcome(
        tenant_id=uuid.uuid4(), document_id=uuid.uuid4()
    )

    assert outcome is False
    assert len(sessions) == 5
    assert sleeps == [0.1, 0.3, 0.6, 1.0]
    assert sum(sleeps) == pytest.approx(2.0)


@pytest.mark.asyncio
async def test_ambiguous_commit_verification_error_is_unknown(monkeypatch):
    class FailingSession:
        async def scalar(self, _statement):
            raise RuntimeError("verification database unavailable")

    class FailingContext:
        async def __aenter__(self):
            return FailingSession()

        async def __aexit__(self, *_args):
            return None

    async def set_context(_db, _tenant_id):
        return None

    monkeypatch.setattr(document_templates, "async_session_maker", FailingContext)
    monkeypatch.setattr(document_templates, "set_tenant_context", set_context)

    outcome = await document_templates._matter_document_commit_outcome(
        tenant_id=uuid.uuid4(), document_id=uuid.uuid4()
    )

    assert outcome is None


@pytest.mark.asyncio
async def test_template_deletion_preserves_terminal_preview_audit_rows(
    client, db_session, test_tenant, test_user
):
    from sqlalchemy import select

    from app.models.document_template import DocumentTemplate
    from app.models.document_template_preview import DocumentTemplatePreview

    await _grant_manage_documents(db_session, test_tenant, test_user)
    template = DocumentTemplate(
        tenant_id=test_tenant.id,
        title="Retired audit template",
        body="",
        format="pdf",
    )
    db_session.add(template)
    await db_session.flush()
    now = datetime.now(timezone.utc)

    def evidence(marker: str, *, purpose: str = "generation", **terminal_fields):
        return DocumentTemplatePreview(
            tenant_id=test_tenant.id,
            template_id=template.id,
            previewed_by_user_id=test_user.id,
            purpose=purpose,
            contract_sha256=marker * 64,
            values_hmac_sha256="b" * 64,
            output_sha256="c" * 64,
            renderer_version="retirement-test",
            flatten_pdf=True,
            reviewed_field_count=0,
            nonblank_field_count=0,
            reviewed_field_names=[],
            expires_at=now + timedelta(hours=1),
            **terminal_fields,
        )

    consumed = evidence("1", consumed_at=now)
    reconciliation = evidence(
        "2",
        reconciliation_required_at=now,
        reconciliation_reason="cleanup_failed",
    )
    pending_generation = evidence("3")
    ordinary = evidence("4", purpose="draft")
    db_session.add_all([consumed, reconciliation, pending_generation, ordinary])
    await db_session.commit()
    terminal_ids = {consumed.id, reconciliation.id, pending_generation.id}
    ordinary_id = ordinary.id

    response = await client.delete(f"/api/templates/{template.id}")

    assert response.status_code == 204, response.text
    terminal_rows = list(
        await db_session.scalars(
            select(DocumentTemplatePreview)
            .where(DocumentTemplatePreview.id.in_(terminal_ids))
            .execution_options(populate_existing=True)
        )
    )
    assert {row.id for row in terminal_rows} == terminal_ids
    assert all(row.template_id is None for row in terminal_rows)
    assert await db_session.get(DocumentTemplatePreview, ordinary_id) is None


@pytest.mark.asyncio
async def test_template_delete_stays_successful_when_post_commit_source_cleanup_fails(
    client,
    db_session,
    test_tenant,
    test_user,
    tmp_path,
    monkeypatch,
    caplog,
):
    import logging
    from pathlib import Path

    from sqlalchemy import select

    from app.models.document_template import DocumentTemplate

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)
    template = DocumentTemplate(
        tenant_id=test_tenant.id,
        title="Cleanup failure template",
        body="",
        format="pdf",
    )
    db_session.add(template)
    await db_session.flush()
    source_dir = Path(
        document_templates._template_source_dir(str(test_tenant.id), template.id)
    )
    source_dir.mkdir(parents=True)
    source_path = source_dir / "source.pdf"
    source_path.write_bytes(b"%PDF-1.7 source")
    template.source_storage_path = str(source_path)
    await db_session.commit()
    original_unlink = Path.unlink

    def fail_source_unlink(path, *args, **kwargs):
        if path.resolve() == source_path.resolve():
            raise OSError("injected source cleanup failure")
        return original_unlink(path, *args, **kwargs)

    caplog.set_level(logging.ERROR, logger="app.routers.document_templates")
    with monkeypatch.context() as cleanup_failure:
        cleanup_failure.setattr(Path, "unlink", fail_source_unlink)
        response = await client.delete(f"/api/templates/{template.id}")

    assert response.status_code == 204, response.text
    assert (
        await db_session.scalar(
            select(DocumentTemplate.id).where(DocumentTemplate.id == template.id)
        )
        is None
    )
    assert source_path.is_file()
    assert "Template source cleanup failed after committed delete" in caplog.text


def test_upload_analysis_detects_fields_and_letterhead_from_text():
    sample = b"""Painter Legal PLLC
123 Main Street
Fargo, ND 58102
(701) 555-0100

July 8, 2026

Dear Ada Lovelace,

Re: Estate Administration

Case No. PB-2026-10

Fee: $2,500.00
"""

    analysis = analyze_template_upload(
        file_bytes=sample,
        filename="fee-agreement.txt",
        content_type="text/plain",
    )

    field_names = {field["name"] for field in analysis.variable_schema["fields"]}
    assert "client_name" in field_names
    assert "matter_name" in field_names
    assert "case_number" in field_names
    assert "fee_amount" in field_names
    assert analysis.branding_profile["letterhead_detected"] is True
    assert "{{client_name}}" in analysis.body


def test_upload_analysis_detects_common_application_label_values():
    analysis = analyze_template_upload(
        file_bytes=(
            b"Applicant Name: Ada Lovelace\n"
            b"Phone: (701) 555-0100\n"
            b"Address: 123 Main Street\n"
            b"City: Fargo\nState: ND\nZip Code: 58102\n"
        ),
        filename="client-application.txt",
        content_type="text/plain",
    )

    fields = {field["name"]: field for field in analysis.variable_schema["fields"]}
    assert {
        "client_name",
        "client_phone",
        "client_street",
        "client_city",
        "client_state",
        "client_zip",
    } <= set(fields)
    assert fields["client_name"]["source_text"] == "Ada Lovelace"
    assert "Applicant Name: {{client_name}}" in analysis.body


def test_upload_analysis_turns_custom_labeled_values_into_reviewable_fields():
    analysis = analyze_template_upload(
        file_bytes=(
            b"Emergency Contact: Grace Hopper\n"
            b"Preferred Pronouns: she/her\n"
            b"Internal Reference: NORTH-42\n"
        ),
        filename="custom-intake.txt",
        content_type="text/plain",
    )

    fields = {field["name"]: field for field in analysis.variable_schema["fields"]}
    assert {"emergency_contact", "preferred_pronouns", "internal_reference"} <= set(
        fields
    )
    assert fields["emergency_contact"]["source_text"] == "Grace Hopper"
    assert "Emergency Contact: {{emergency_contact}}" in analysis.body


def test_pdf_analysis_discovers_acroform_fields_without_losing_mapping():
    analysis = analyze_template_upload(
        file_bytes=_fillable_pdf(),
        filename="client-intake.pdf",
        content_type="application/pdf",
    )
    fields = {field["name"]: field for field in analysis.variable_schema["fields"]}
    assert fields["client_name"]["pdf_field_name"] == "Client Name"
    assert fields["approved"]["field_type"] == "checkbox"
    assert fields["notes"]["multiline"] is True
    assert "{{client_name}}" in analysis.body
    assert analysis.format == "pdf"


def test_static_application_pdf_discovers_and_renders_reviewed_overlay_fields():
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from app.utils.text_processing import extract_text_from_pdf

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 720, "Applicant Name:")
    pdf.drawString(72, 690, "Dear Ada Lovelace,")
    pdf.drawString(72, 660, "Case No. CV-2026-42")
    pdf.save()

    analysis = analyze_template_upload(
        file_bytes=output.getvalue(),
        filename="client-application.pdf",
        content_type="application/pdf",
    )
    fields = {field["name"]: field for field in analysis.variable_schema["fields"]}

    assert analysis.variable_schema["source"] == "pdf_text_overlay"
    assert {"client_name", "case_number"} <= set(fields)
    assert all(field.get("pdf_overlay") for field in fields.values())
    assert all(field.get("pdf_source_key") for field in fields.values())
    assert len(fields["client_name"]["pdf_overlays"]) == 2

    rendered = fill_pdf_template(
        output.getvalue(),
        variable_schema=analysis.variable_schema,
        variables={
            "client_name": "Grace Hopper",
            "case_number": "CV-2027-9",
        },
    )
    rendered_text = extract_text_from_pdf(rendered)
    assert "Grace Hopper" in rendered_text
    assert "CV-2027-9" in rendered_text
    assert "Ada Lovelace" not in rendered_text
    assert "CV-2026-42" not in rendered_text


def test_pdf_intake_reuses_one_validated_reader(monkeypatch):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    import app.services.pdf_templates as pdf_template_service

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 720, "Applicant Name:")
    pdf.drawString(72, 690, "Dear Ada Lovelace,")
    pdf.save()

    open_count = 0
    original_open = pdf_template_service._open_pdf

    def counted_open(content):
        nonlocal open_count
        open_count += 1
        return original_open(content)

    monkeypatch.setattr(pdf_template_service, "_open_pdf", counted_open)

    analysis = analyze_template_upload(
        file_bytes=output.getvalue(),
        filename="application.pdf",
        content_type="application/pdf",
    )

    assert analysis.variable_schema["source"] == "pdf_text_overlay"
    assert open_count == 1


def test_image_only_pdf_uses_ocr_coordinates_and_renders_flattened_fields(
    monkeypatch,
):
    from PIL import Image, ImageDraw
    from pypdf import PdfReader
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas
    import app.services.template_intake as template_intake
    from app.utils.text_processing import extract_text_from_pdf

    image = Image.new("RGB", (1224, 1584), "white")
    drawing = ImageDraw.Draw(image)
    drawing.text((140, 180), "Applicant Name: Ada Lovelace", fill="black")
    drawing.text((140, 260), "Case Number: CV-2026-42", fill="black")
    source = BytesIO()
    pdf = canvas.Canvas(source, pagesize=(612, 792))
    pdf.drawImage(ImageReader(image), 0, 0, width=612, height=792)
    pdf.save()

    monkeypatch.setattr(
        template_intake,
        "ocr_pdf",
        lambda _content: PdfOcrResult(
            text="Applicant Name: Ada Lovelace\nCase Number: CV-2026-42",
            lines=(
                OcrLine(0, "Applicant Name: Ada Lovelace", 0.97, (70, 680, 310, 704)),
                OcrLine(0, "Case Number: CV-2026-42", 0.95, (70, 638, 300, 662)),
            ),
            pages_analyzed=1,
            pages_total=1,
            average_confidence=0.96,
            truncated=False,
        ),
    )

    analysis = analyze_template_upload(
        file_bytes=source.getvalue(),
        filename="scanned-application.pdf",
        content_type="application/pdf",
    )
    fields = {field["name"]: field for field in analysis.variable_schema["fields"]}

    assert analysis.variable_schema["source"] == "pdf_ocr_overlay"
    assert analysis.variable_schema["detection"]["method"] == "ocr"
    assert {"client_name", "case_number"} <= set(fields)
    assert fields["client_name"]["pdf_overlay"]["source_kind"] == "ocr"

    rendered = fill_pdf_template(
        source.getvalue(),
        variable_schema=analysis.variable_schema,
        variables={"client_name": "Grace Hopper", "case_number": "CV-2027-9"},
    )
    rendered_text = extract_text_from_pdf(rendered)
    assert "Grace Hopper" in rendered_text
    assert "CV-2027-9" in rendered_text
    assert len(PdfReader(BytesIO(rendered)).pages[0].images) == 1


def test_reviewed_static_pdf_schema_keeps_server_discovered_placements():
    import copy
    import json
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 720, "Applicant Name:")
    pdf.drawString(72, 690, "Dear Ada Lovelace,")
    pdf.save()
    analysis = analyze_template_upload(
        file_bytes=output.getvalue(),
        filename="application.pdf",
        content_type="application/pdf",
    )
    discovered = analysis.variable_schema
    reviewed = copy.deepcopy(discovered)
    reviewed["fields"][0]["name"] = "party_name"
    reviewed["fields"][0]["pdf_overlays"][0]["rect"] = [72, 700, 220, 716]
    assert discovered["pages"]

    validated = document_templates._reviewed_variable_schema(
        json.dumps(reviewed), discovered
    )

    assert validated["fields"][0]["name"] == "party_name"
    assert validated["fields"][0]["pdf_overlays"][0]["rect"] == [72.0, 700.0, 220.0, 716.0]
    assert validated["fields"][0]["pdf_overlays"][0]["source_rect"] == discovered["fields"][0]["pdf_overlays"][0]["source_rect"]
    assert len(validated["fields"][0]["pdf_overlays"]) == 2


def test_docx_source_render_preserves_structure_and_replaces_split_run_values():
    from docx import Document

    document = Document()
    paragraph = document.add_paragraph("Dear ")
    paragraph.add_run("Ada ").bold = True
    paragraph.add_run("Lovelace,")
    document.add_paragraph("Case No. CV-2026-42")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Client: Ada Lovelace"
    source = BytesIO()
    document.save(source)

    analysis = analyze_template_upload(
        file_bytes=source.getvalue(),
        filename="engagement-letter.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert analysis.variable_schema["source"] == "docx_source"

    rendered = fill_docx_template(
        source.getvalue(),
        variable_schema=analysis.variable_schema,
        variables={"client_name": "Ada Lovelace Jr.", "case_number": "CV-2027-9"},
    )
    reopened = Document(BytesIO(rendered))

    assert reopened.paragraphs[0].text == "Dear Ada Lovelace Jr.,"
    assert reopened.paragraphs[1].text == "Case No. CV-2027-9"
    assert reopened.tables[0].cell(0, 0).text == "Client: Ada Lovelace Jr."
    assert len(reopened.tables) == 1


def test_docx_intake_recognizes_bracket_placeholders_and_ignores_static_brackets():
    from docx import Document

    document = Document()
    document.add_paragraph("Plaintiff: [PLAINTIFF'S FULL NAME]")
    document.add_paragraph("Defendant: [DEFENDANT'S FULL NAME]")
    document.add_paragraph("[THIS SPACE INTENTIONALLY LEFT BLANK]")
    source = BytesIO()
    document.save(source)

    analysis = analyze_template_upload(
        file_bytes=source.getvalue(),
        filename="stipulation.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    fields = {field["name"]: field for field in analysis.variable_schema["fields"]}

    assert set(fields) == {"plaintiff_name", "defendant_name"}
    assert fields["plaintiff_name"]["source_text"] == "[PLAINTIFF'S FULL NAME]"

    rendered = fill_docx_template(
        source.getvalue(),
        variable_schema=analysis.variable_schema,
        variables={"plaintiff_name": "Alex Plaintiff", "defendant_name": "Dana Defendant"},
    )
    reopened = Document(BytesIO(rendered))
    assert reopened.paragraphs[0].text == "Plaintiff: Alex Plaintiff"
    assert reopened.paragraphs[1].text == "Defendant: Dana Defendant"
    assert reopened.paragraphs[2].text == "[THIS SPACE INTENTIONALLY LEFT BLANK]"


def test_docx_intake_anchors_identical_underscore_blanks_independently():
    from docx import Document

    document = Document()
    document.add_paragraph("Husband's date of birth: __________")
    document.add_paragraph("Wife's date of birth: __________")
    document.add_paragraph("Section A: Basic Information")
    source = BytesIO()
    document.save(source)

    analysis = analyze_template_upload(
        file_bytes=source.getvalue(),
        filename="questionnaire.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    fields = {field["name"]: field for field in analysis.variable_schema["fields"]}

    assert set(fields) == {"husband_s_date_of_birth", "wife_s_date_of_birth"}
    assert all(field["docx_anchor"] for field in fields.values())
    assert fields["husband_s_date_of_birth"]["source_text"] == "__________"

    rendered = fill_docx_template(
        source.getvalue(),
        variable_schema=analysis.variable_schema,
        variables={
            "husband_s_date_of_birth": "01/02/1980",
            "wife_s_date_of_birth": "03/04/1982",
        },
    )
    reopened = Document(BytesIO(rendered))
    assert reopened.paragraphs[0].text.endswith("01/02/1980")
    assert reopened.paragraphs[1].text.endswith("03/04/1982")
    assert reopened.paragraphs[2].text == "Section A: Basic Information"


def test_docx_intake_finds_late_placeholders_without_truncating_schema():
    from docx import Document

    document = Document()
    document.add_paragraph("A" * 20_050)
    document.add_paragraph("Signature date: [DATE]")
    source = BytesIO()
    document.save(source)

    analysis = analyze_template_upload(
        file_bytes=source.getvalue(),
        filename="long-template.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert len(analysis.body) == 20_000
    assert len(analysis.extracted_text) == 20_000
    assert [field["name"] for field in analysis.variable_schema["fields"]] == ["date"]
    assert any("full Word source" in warning for warning in analysis.warnings)


def test_docx_intake_rejects_tracked_changes_before_field_detection():
    from docx import Document
    import zipfile

    document = Document()
    document.add_paragraph("Client: Ada Lovelace")
    source = BytesIO()
    document.save(source)

    revised = BytesIO()
    with zipfile.ZipFile(source, "r") as source_zip, zipfile.ZipFile(revised, "w") as revised_zip:
        for item in source_zip.infolist():
            payload = source_zip.read(item.filename)
            if item.filename == "word/document.xml":
                payload = payload.replace(
                    b"<w:body>",
                    b'<w:body><w:ins w:id="1" w:author="test" w:date="2026-01-01T00:00:00Z">',
                    1,
                ).replace(b"</w:body>", b"</w:ins></w:body>", 1)
            revised_zip.writestr(item, payload)

    with pytest.raises(TemplateDocxError, match="tracked changes"):
        analyze_template_upload(
            file_bytes=revised.getvalue(),
            filename="revised.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


def test_docx_render_recurses_through_deep_tables():
    from docx import Document

    document = Document()
    outer = document.add_table(rows=1, cols=1)
    middle = outer.cell(0, 0).add_table(rows=1, cols=1)
    inner = middle.cell(0, 0).add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "Deep value: {{deep_value}}"
    source = BytesIO()
    document.save(source)

    rendered = fill_docx_template(
        source.getvalue(),
        variable_schema={"fields": [{"name": "deep_value"}]},
        variables={"deep_value": "Found"},
    )
    reopened = Document(BytesIO(rendered))
    reopened_middle = reopened.tables[0].cell(0, 0).tables[0]
    reopened_inner = reopened_middle.cell(0, 0).tables[0]

    assert reopened_inner.cell(0, 0).text == "Deep value: Found"


def test_docx_render_never_reinterprets_inserted_values_as_source_text():
    from docx import Document

    document = Document()
    document.add_paragraph("{{first_value}} / ORIGINAL SECOND")
    source = BytesIO()
    document.save(source)

    rendered = fill_docx_template(
        source.getvalue(),
        variable_schema={
            "fields": [
                {"name": "first_value"},
                {"name": "second_value", "source_text": "ORIGINAL SECOND"},
            ]
        },
        variables={
            "first_value": "ORIGINAL SECOND",
            "second_value": "FINAL SECOND",
        },
    )
    reopened = Document(BytesIO(rendered))

    assert reopened.paragraphs[0].text == "ORIGINAL SECOND / FINAL SECOND"


def test_docx_analysis_rejects_damaged_upload_with_actionable_error():
    with pytest.raises(TemplateDocxError, match="damaged or could not be parsed"):
        analyze_template_upload(
            file_bytes=b"this is not an office package",
            filename="damaged.docx",
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )


@pytest.mark.asyncio
async def test_docx_intake_retains_source_and_binary_preview(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from docx import Document

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)
    document = Document()
    paragraph = document.add_paragraph("Dear ")
    paragraph.add_run("Ada ")
    paragraph.add_run("Lovelace,")
    source = BytesIO()
    document.save(source)

    created = await client.post(
        "/api/templates/intake/create",
        files={
            "file": (
                "engagement.docx",
                source.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert created.status_code == 201, created.text
    payload = created.json()
    assert payload["format"] == "docx"
    assert payload["source_filename"] == "engagement.docx"
    assert payload["source_sha256"]
    assert payload["variable_schema"]["source"] == "docx_source"

    preview = await client.post(
        f"/api/templates/{payload['id']}/render-file",
        json={"variables": {"client_name": "Grace Hopper"}},
    )
    assert preview.status_code == 200, preview.text
    assert preview.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    rendered = Document(BytesIO(preview.content))
    assert rendered.paragraphs[0].text == "Dear Grace Hopper,"


@pytest.mark.asyncio
async def test_static_pdf_intake_creates_reviewable_overlay_template(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)
    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 720, "Applicant Name:")
    pdf.drawString(72, 690, "Dear Ada Lovelace,")
    pdf.save()

    created = await client.post(
        "/api/templates/intake/create",
        files={"file": ("application.pdf", output.getvalue(), "application/pdf")},
    )
    assert created.status_code == 201, created.text
    payload = created.json()
    fields = payload["variable_schema"]["fields"]
    assert {field["name"] for field in fields} == {"client_name"}
    assert all(field.get("pdf_overlay") for field in fields)
    assert len(fields[0]["pdf_overlays"]) == 2

    preview = await client.post(
        f"/api/templates/{payload['id']}/render-file",
        json={
            "variables": {"client_name": "Grace Hopper"},
            "preview_purpose": "activation",
        },
    )
    assert preview.status_code == 200, preview.text
    assert preview.headers["content-type"].startswith("application/pdf")
    assert preview.headers["x-clarity-preview-purpose"] == "activation"


def test_pdf_discovery_and_fill_support_radio_choice_pairs_and_hierarchy():
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import (
        ArrayObject,
        DictionaryObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    source = _special_fields_pdf()
    fields = {field["name"]: field for field in discover_pdf_fields(source)}
    assert fields["filing_type"]["field_type"] == "radio"
    assert fields["filing_type"]["options"] == ["Complaint", "Motion"]
    assert fields["forum"]["options"] == [
        {"value": "State", "label": "State Court"},
        {"value": "Federal", "label": "Federal Court"},
    ]

    schema = {"fields": list(fields.values())}
    flattened = fill_pdf_template(
        source,
        variable_schema=schema,
        variables={"forum": "Federal", "filing_type": "Motion"},
        flatten=True,
    )
    flattened_reader = PdfReader(BytesIO(flattened))
    assert "Federal Court" in (flattened_reader.pages[0].extract_text() or "")
    assert flattened_reader.get_fields() is None

    editable = fill_pdf_template(
        source,
        variable_schema=schema,
        variables={"forum": "Federal", "filing_type": "Motion"},
        flatten=False,
    )
    editable_fields = PdfReader(BytesIO(editable)).get_fields()
    assert str(editable_fields["forum"]["/V"]) == "Federal"
    assert str(editable_fields["filing_type"]["/V"]) == "/Motion"

    base = BytesIO()
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(base, pagesize=letter)
    pdf.drawString(72, 740, "Party")
    pdf.acroForm.textfield(name="name", x=150, y=705, width=240, height=24)
    pdf.save()
    writer = PdfWriter()
    writer.clone_document_from_reader(PdfReader(BytesIO(base.getvalue())))
    acroform = writer._root_object["/AcroForm"].get_object()
    child_ref = acroform["/Fields"][0]
    child = child_ref.get_object()
    parent = DictionaryObject(
        {
            NameObject("/T"): TextStringObject("party"),
            NameObject("/Kids"): ArrayObject([child_ref]),
        }
    )
    parent_ref = writer._add_object(parent)
    child[NameObject("/Parent")] = parent_ref
    acroform[NameObject("/Fields")] = ArrayObject([parent_ref])
    hierarchical = BytesIO()
    writer.write(hierarchical)
    hierarchical_fields = discover_pdf_fields(hierarchical.getvalue())
    assert [field["pdf_field_name"] for field in hierarchical_fields] == ["party.name"]

    push_writer = PdfWriter()
    push_writer.clone_document_from_reader(PdfReader(BytesIO(_fillable_pdf())))
    push_field = next(
        annotation.get_object()
        for page in push_writer.pages
        for annotation in (page.get("/Annots") or [])
        if str(annotation.get_object().get("/T")) == "Approved"
    )
    push_field[NameObject("/Ff")] = NumberObject(
        int(push_field.get("/Ff", 0) or 0) | (1 << 16)
    )
    push_source = BytesIO()
    push_writer.write(push_source)
    push_names = {
        field["pdf_field_name"] for field in discover_pdf_fields(push_source.getvalue())
    }
    assert "Approved" not in push_names


def test_pdf_fill_and_flatten_preserves_page_and_removes_form_widgets():
    from pypdf import PdfReader

    source = _fillable_pdf()
    schema = {"fields": discover_pdf_fields(source)}
    output = fill_pdf_template(
        source,
        variable_schema=schema,
        variables={"client_name": "Ada Lovelace", "approved": "yes"},
        flatten=True,
    )
    reader = PdfReader(BytesIO(output))
    assert len(reader.pages) == 1
    assert reader.get_fields() is None
    text = reader.pages[0].extract_text()
    for source_label in ("Client:", "Approved:", "Notes:"):
        assert source_label in text
    assert "Ada Lovelace" in text
    assert "X" in text


def test_pdf_parser_rejects_non_pdf_and_encrypted_pdf():
    import pytest
    from pypdf import PdfWriter

    with pytest.raises(TemplatePdfError, match="not a valid PDF"):
        discover_pdf_fields(b"not-pdf")

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt("secret")
    encrypted = BytesIO()
    writer.write(encrypted)
    with pytest.raises(TemplatePdfError, match="Password-protected"):
        discover_pdf_fields(encrypted.getvalue())


def test_pdf_parser_rejects_javascript_active_content():
    from pypdf import PdfReader, PdfWriter

    writer = PdfWriter()
    writer.clone_document_from_reader(PdfReader(BytesIO(_fillable_pdf())))
    writer.add_js("app.alert('unsafe')")
    active = BytesIO()
    writer.write(active)

    with pytest.raises(TemplatePdfError, match="Active PDF content"):
        discover_pdf_fields(active.getvalue())


def test_pdf_render_fails_instead_of_corrupting_unsupported_glyphs():
    source = _fillable_pdf()
    schema = {"fields": discover_pdf_fields(source)}
    with pytest.raises(TemplatePdfError, match="cannot safely display"):
        fill_pdf_template(
            source,
            variable_schema=schema,
            variables={"client_name": "Client \U0001f4bc"},
            flatten=True,
        )

    for unsupported_text in ("مرحبا", "नमस्ते", "שלום"):
        with pytest.raises(TemplatePdfError, match="cannot safely shape"):
            fill_pdf_template(
                source,
                variable_schema=schema,
                variables={"client_name": unsupported_text},
                flatten=True,
            )


def test_pdf_final_render_requires_explicit_review_and_true_required_checkbox():
    source = _required_pdf()
    schema = {"fields": discover_pdf_fields(source)}

    # Preview is intentionally permissive so the user can inspect incomplete work.
    fill_pdf_template(
        source,
        variable_schema=schema,
        variables={},
        flatten=True,
        enforce_required=False,
    )
    with pytest.raises(TemplatePdfError, match="Review every PDF field"):
        fill_pdf_template(
            source,
            variable_schema=schema,
            variables={"client_name": "Ada"},
            flatten=True,
            enforce_required=True,
        )
    with pytest.raises(TemplatePdfError, match="empty or unchecked"):
        fill_pdf_template(
            source,
            variable_schema=schema,
            variables={"client_name": "Ada", "accepted": "false"},
            flatten=True,
            enforce_required=True,
        )
    fill_pdf_template(
        source,
        variable_schema=schema,
        variables={"client_name": "Ada", "accepted": "true"},
        flatten=True,
        enforce_required=True,
    )
    with pytest.raises(TemplatePdfError, match="cannot safely display"):
        fill_pdf_template(
            source,
            variable_schema=schema,
            variables={"client_name": "Client \U0001f4bc"},
            flatten=False,
        )


def test_pdf_activation_preview_requires_representative_non_signature_values():
    source = _fillable_pdf()
    schema = {"fields": discover_pdf_fields(source)}

    with pytest.raises(TemplatePdfError, match="representative values") as exc_info:
        validate_representative_pdf_variables(
            schema,
            {"client_name": "", "approved": "false", "notes": ""},
        )
    assert "client_name" in str(exc_info.value)
    assert "notes" in str(exc_info.value)

    validate_representative_pdf_variables(
        schema,
        {
            "client_name": "Representative Client",
            "approved": "false",
            "notes": "Representative multiline narrative",
        },
    )


def test_pdf_final_render_clears_sample_values_and_leaves_signature_blank():
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import NameObject, NumberObject

    source = _fillable_pdf()
    writer = PdfWriter()
    writer.clone_document_from_reader(PdfReader(BytesIO(source)))
    for page in writer.pages:
        writer.update_page_form_field_values(
            page,
            {
                "Client Name": "SAMPLE CLIENT SECRET",
                "Approved": "/Yes",
                "Notes": "sample privileged narrative",
            },
            auto_regenerate=True,
        )
    populated = BytesIO()
    writer.write(populated)
    populated_source = populated.getvalue()
    populated_schema = {"fields": discover_pdf_fields(populated_source)}
    with pytest.raises(TemplatePdfError, match="Review every PDF field"):
        fill_pdf_template(
            populated_source,
            variable_schema=populated_schema,
            variables={"client_name": "Final Client", "approved": "false"},
            flatten=True,
            enforce_required=True,
        )
    cleared = fill_pdf_template(
        populated_source,
        variable_schema=populated_schema,
        variables={
            "client_name": "Final Client",
            "approved": "false",
            "notes": "",
        },
        flatten=True,
        enforce_required=True,
    )
    cleared_text = PdfReader(BytesIO(cleared)).pages[0].extract_text() or ""
    assert "Final Client" in cleared_text
    assert "SAMPLE CLIENT SECRET" not in cleared_text
    assert "sample privileged narrative" not in cleared_text

    signature_writer = PdfWriter()
    signature_writer.clone_document_from_reader(PdfReader(BytesIO(source)))
    signature_field = next(
        annotation.get_object()
        for page in signature_writer.pages
        for annotation in (page.get("/Annots") or [])
        if str(annotation.get_object().get("/T")) == "Client Name"
    )
    signature_field[NameObject("/FT")] = NameObject("/Sig")
    signature_field[NameObject("/Ff")] = NumberObject(2)
    signature_source = BytesIO()
    signature_writer.write(signature_source)
    signature_schema = {"fields": discover_pdf_fields(signature_source.getvalue())}
    signature_meta = {field["name"]: field for field in signature_schema["fields"]}
    assert signature_meta["client_name"]["field_type"] == "signature"
    signed_later = fill_pdf_template(
        signature_source.getvalue(),
        variable_schema=signature_schema,
        variables={"approved": "false", "notes": ""},
        flatten=True,
        enforce_required=True,
    )
    signed_later_reader = PdfReader(BytesIO(signed_later))
    assert signed_later_reader.get_fields() is None
    assert "SAMPLE CLIENT SECRET" not in (
        signed_later_reader.pages[0].extract_text() or ""
    )


def test_pdf_text_extraction_honors_page_and_character_caps():
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    from app.utils.text_processing import extract_text_from_pdf

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    for page_number in range(1, 4):
        pdf.drawString(72, 720, f"PAGE-{page_number} " + "x" * 50)
        pdf.showPage()
    pdf.save()

    first_two = extract_text_from_pdf(output.getvalue(), max_pages=2)
    assert "PAGE-1" in first_two
    assert "PAGE-2" in first_two
    assert "PAGE-3" not in first_two
    assert len(extract_text_from_pdf(output.getvalue(), max_chars=25)) <= 25


def test_pdf_intake_analysis_does_not_extract_past_page_cap():
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    for page_number in range(1, 56):
        pdf.drawString(72, 720, f"INTAKE-PAGE-{page_number:02d}")
        pdf.showPage()
    pdf.save()

    analysis = analyze_template_upload(
        file_bytes=output.getvalue(),
        filename="large.pdf",
        content_type="application/pdf",
    )
    assert "INTAKE-PAGE-50" in analysis.extracted_text
    assert "INTAKE-PAGE-51" not in analysis.extracted_text


def test_extract_schema_variables_when_body_has_no_placeholders():
    template = SimpleNamespace(
        variable_schema={
            "fields": [
                {"name": "client_name"},
                {"name": "case_number"},
                {"name": "client_name"},
            ]
        }
    )

    variables = document_templates.extract_schema_variables(template)

    assert variables == ["client_name", "case_number"]


@pytest.mark.asyncio
async def test_build_variable_suggestions_from_matter_client_and_current_user(
    monkeypatch,
):
    client_id = uuid.uuid4()
    attorney_id = uuid.uuid4()
    matter_id = uuid.uuid4()
    user_id = uuid.uuid4()

    matter = SimpleNamespace(
        id=matter_id,
        matter_name="Estate of Lovelace",
        matter_type="probate",
        description=None,
        status="open",
        stage="intake",
        jurisdiction="North Dakota",
        case_number="PB-2026-10",
        court="Cass County District Court",
        judge=None,
        billing_method="hourly",
        billing_cycle="monthly",
        hourly_rate=Decimal("250.00"),
        budget_amount=None,
        counterparty=None,
        client=SimpleNamespace(
            id=client_id,
            display_name="Ada Lovelace",
            email="ada@example.com",
            phone="555-0100",
            address={"city": "Fargo", "state": "ND", "zip": "58102"},
        ),
        attorney_of_record=SimpleNamespace(
            id=attorney_id,
            full_name="Grace Hopper",
            email="grace@example.com",
        ),
    )

    async def fake_load_matter_context(**kwargs):
        return matter

    monkeypatch.setattr(
        document_templates, "_load_matter_context", fake_load_matter_context
    )

    template = SimpleNamespace(
        id=uuid.uuid4(),
        body=(
            "{{client_name}} {{case_number}} {{attorney_email}} "
            "{{current_user_email}} {{missing_fact}}"
        ),
    )
    current_user = SimpleNamespace(
        id=user_id,
        full_name="Test Attorney",
        email="test@example.com",
    )

    _, suggestions = await document_templates.build_variable_suggestions(
        template=template,
        requested_variables=None,
        matter_id=str(matter_id),
        tenant_id=uuid.uuid4(),
        current_user=current_user,
        db=SimpleNamespace(),
    )

    by_variable = {item.variable: item for item in suggestions}

    assert by_variable["client_name"].suggested_value == "Ada Lovelace"
    assert by_variable["client_name"].source_type == "contact"
    assert by_variable["case_number"].suggested_value == "PB-2026-10"
    assert by_variable["attorney_email"].suggested_value == "grace@example.com"
    assert by_variable["current_user_email"].suggested_value == "test@example.com"
    assert by_variable["missing_fact"].suggested_value is None
    assert by_variable["missing_fact"].review_required is True


@pytest.mark.asyncio
async def test_template_upload_lost_commit_ack_preserves_committed_source(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from pathlib import Path

    from sqlalchemy import select

    from app.models.document_template import DocumentTemplate

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)
    original_commit = db_session.commit

    async def commit_then_lose_ack():
        await original_commit()
        raise RuntimeError("injected lost template commit acknowledgement")

    with monkeypatch.context() as lost_ack:
        lost_ack.setattr(db_session, "commit", commit_then_lose_ack)
        response = await client.post(
            "/api/templates/intake/create",
            files={"file": ("lost-ack.pdf", _fillable_pdf(), "application/pdf")},
            data={"title": "Lost ACK template", "category": "other"},
        )

    assert response.status_code == 201, response.text
    template_id = uuid.UUID(response.json()["id"])
    template = await db_session.scalar(
        select(DocumentTemplate).where(DocumentTemplate.id == template_id)
    )
    assert template is not None
    assert Path(template.source_storage_path).is_file()


@pytest.mark.asyncio
async def test_template_upload_confirmed_rollback_removes_staged_source(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from pathlib import Path

    from sqlalchemy import func, select

    from app.models.document_template import DocumentTemplate

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)

    async def fail_commit():
        raise RuntimeError("injected template commit failure")

    async def confirmed_rollback(**_kwargs):
        return False, None

    with monkeypatch.context() as failure:
        failure.setattr(db_session, "commit", fail_commit)
        failure.setattr(
            document_templates, "_template_commit_outcome", confirmed_rollback
        )
        response = await client.post(
            "/api/templates/intake/create",
            files={"file": ("rolled-back.pdf", _fillable_pdf(), "application/pdf")},
            data={"title": "Rolled back template", "category": "other"},
        )

    assert response.status_code == 500, response.text
    assert "staged source was removed" in response.json()["detail"]
    assert list(Path(tmp_path).rglob("*.pdf")) == []
    assert (
        await db_session.scalar(select(func.count()).select_from(DocumentTemplate)) == 0
    )


@pytest.mark.asyncio
async def test_template_upload_confirmed_rollback_cleanup_failure_preserves_source(
    client,
    db_session,
    test_tenant,
    test_user,
    tmp_path,
    monkeypatch,
    caplog,
):
    import logging
    from pathlib import Path

    from sqlalchemy import func, select

    from app.models.document_template import DocumentTemplate

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)

    async def fail_commit():
        raise RuntimeError("injected template commit failure")

    async def confirmed_rollback(**_kwargs):
        return False, None

    def fail_unlink(_path, *_args, **_kwargs):
        raise OSError("injected scoped source cleanup failure")

    caplog.set_level(logging.CRITICAL, logger="app.routers.document_templates")
    with monkeypatch.context() as failure:
        failure.setattr(db_session, "commit", fail_commit)
        failure.setattr(
            document_templates, "_template_commit_outcome", confirmed_rollback
        )
        failure.setattr(Path, "unlink", fail_unlink)
        response = await client.post(
            "/api/templates/intake/create",
            files={"file": ("cleanup-failed.pdf", _fillable_pdf(), "application/pdf")},
            data={"title": "Cleanup failed template", "category": "other"},
        )

    assert response.status_code == 500, response.text
    assert "retained-source cleanup failed" in response.json()["detail"]
    preserved = list(Path(tmp_path).rglob("*.pdf"))
    assert len(preserved) == 1
    assert preserved[0].is_file()
    assert "uncommitted template-source cleanup failed" in caplog.text
    assert (
        await db_session.scalar(select(func.count()).select_from(DocumentTemplate)) == 0
    )


@pytest.mark.asyncio
async def test_template_upload_unknown_commit_preserves_source_for_reconciliation(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from pathlib import Path

    from sqlalchemy import func, select

    from app.models.document_template import DocumentTemplate

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)

    async def fail_commit():
        raise RuntimeError("injected ambiguous template commit")

    async def unknown_outcome(**_kwargs):
        return None, None

    with monkeypatch.context() as failure:
        failure.setattr(db_session, "commit", fail_commit)
        failure.setattr(document_templates, "_template_commit_outcome", unknown_outcome)
        response = await client.post(
            "/api/templates/intake/create",
            files={"file": ("unknown-commit.pdf", _fillable_pdf(), "application/pdf")},
            data={"title": "Unknown commit template", "category": "other"},
        )

    assert response.status_code == 500, response.text
    assert "outcome could not be verified" in response.json()["detail"]
    preserved = list(Path(tmp_path).rglob("*.pdf"))
    assert len(preserved) == 1
    assert preserved[0].read_bytes().startswith(b"%PDF")
    assert (
        await db_session.scalar(select(func.count()).select_from(DocumentTemplate)) == 0
    )


@pytest.mark.asyncio
async def test_pdf_upload_preview_and_matter_render_are_binary_and_unique(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from pathlib import Path

    from pypdf import PdfReader
    from sqlalchemy import func, select

    from app.models.document_template import DocumentTemplate
    from app.models.document_template_preview import DocumentTemplatePreview
    from app.models.matter_document import MatterDocument
    from app.models.plugin import Matter, MatterEvent
    from app.services import matter_file_store as matter_store_module

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    monkeypatch.setattr(matter_store_module.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)

    response = await client.post(
        "/api/templates/intake/create",
        files={"file": ("client-intake.pdf", _fillable_pdf(), "application/pdf")},
        data={"title": "../../Client Intake", "category": "other"},
    )
    assert response.status_code == 201, response.text
    template_id = response.json()["id"]
    assert response.json()["format"] == "pdf"
    assert response.json()["source_filename"] == "client-intake.pdf"
    assert "source_storage_path" not in response.json()

    template = await db_session.scalar(
        select(DocumentTemplate).where(DocumentTemplate.id == uuid.UUID(template_id))
    )
    source_path = Path(template.source_storage_path).resolve()
    assert source_path.is_relative_to(tmp_path.resolve())
    import hashlib

    assert (
        hashlib.sha256(source_path.read_bytes()).hexdigest()
        == response.json()["source_sha256"]
    )
    source_response = await client.get(f"/api/templates/{template_id}/source")
    assert source_response.status_code == 200
    assert source_response.headers["cache-control"] == "private, no-store"
    assert hashlib.sha256(source_response.content).hexdigest() == template.source_sha256

    preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": {
                "client_name": "Representative Client",
                "approved": "yes",
                "notes": "Representative multiline narrative",
            },
            "preview_purpose": "activation",
        },
    )
    assert preview.status_code == 200, preview.text
    assert preview.headers["content-type"].startswith("application/pdf")
    preview_reader = PdfReader(BytesIO(preview.content))
    assert preview_reader.get_fields() is None
    preview_text = preview_reader.pages[0].extract_text()
    assert "Client:" in preview_text
    assert "Representative Client" in preview_text
    assert preview.headers["x-clarity-preview-purpose"] == "activation"
    assert preview.headers["x-clarity-preview-id"]

    matter = Matter(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        user_id=test_user.id,
        slug="pdf-incident-matter",
        matter_name="PDF Incident Matter",
        matter_type="general",
    )
    db_session.add(matter)
    await db_session.commit()

    inactive_save = await client.post(
        f"/api/templates/{template_id}/render",
        json={
            "variables": {
                "client_name": "Ada Lovelace",
                "approved": "yes",
                "notes": "",
            },
            "matter_id": str(matter.id),
        },
    )
    assert inactive_save.status_code == 409

    activated = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert activated.status_code == 200, activated.text

    generation_values = {
        "client_name": "Ada Lovelace",
        "approved": "yes",
        "notes": "",
    }
    missing_preview = await client.post(
        f"/api/templates/{template_id}/render",
        json={"variables": generation_values, "matter_id": str(matter.id)},
    )
    assert missing_preview.status_code == 409
    assert "Preview the exact current PDF values" in missing_preview.json()["detail"]

    generation_preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": generation_values,
            "matter_id": str(matter.id),
            "preview_purpose": "generation",
        },
    )
    assert generation_preview.status_code == 200, generation_preview.text
    generation_preview_id = generation_preview.headers["x-clarity-preview-id"]

    changed_after_preview = await client.post(
        f"/api/templates/{template_id}/render",
        json={
            "variables": {**generation_values, "client_name": "Grace Hopper"},
            "matter_id": str(matter.id),
            "preview_id": generation_preview_id,
        },
    )
    assert changed_after_preview.status_code == 409
    assert "field values changed" in changed_after_preview.json()["detail"]

    first_save = await client.post(
        f"/api/templates/{template_id}/render",
        json={
            "variables": generation_values,
            "matter_id": str(matter.id),
            "preview_id": generation_preview_id,
        },
    )
    assert first_save.status_code == 200, first_save.text
    first_payload = first_save.json()
    assert first_payload["output_format"] == "pdf"
    assert first_payload["output_filename"].endswith(".pdf")
    assert first_payload["download_url"].endswith("/download")
    assert first_payload["storage_backend"] == "local"
    assert first_payload.get("storage_warning") is None

    replay = await client.post(
        f"/api/templates/{template_id}/render",
        json={
            "variables": generation_values,
            "matter_id": str(matter.id),
            "preview_id": generation_preview_id,
        },
    )
    assert replay.status_code == 200, replay.text
    assert replay.json()["matter_document_id"] == first_payload["matter_document_id"]
    assert replay.json()["output_filename"] == first_payload["output_filename"]
    assert "original request" in replay.json()["rendered"]
    consumed_wrong_values = await client.post(
        f"/api/templates/{template_id}/render",
        json={
            "variables": {**generation_values, "client_name": "Different Client"},
            "matter_id": str(matter.id),
            "preview_id": generation_preview_id,
        },
    )
    assert consumed_wrong_values.status_code == 409
    assert "field values changed" in consumed_wrong_values.json()["detail"]
    assert "consumed" not in consumed_wrong_values.json()["detail"]
    assert (
        await db_session.scalar(
            select(func.count())
            .select_from(MatterDocument)
            .where(MatterDocument.matter_id == matter.id)
        )
        == 1
    )
    assert (
        await db_session.scalar(
            select(func.count())
            .select_from(MatterEvent)
            .where(
                MatterEvent.matter_id == matter.id,
                MatterEvent.event_type == "document_generated",
            )
        )
        == 1
    )

    second_preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": generation_values,
            "matter_id": str(matter.id),
            "preview_purpose": "generation",
        },
    )
    assert second_preview.status_code == 200, second_preview.text
    second_preview_id = second_preview.headers["x-clarity-preview-id"]
    second_save = await client.post(
        f"/api/templates/{template_id}/render",
        json={
            "variables": generation_values,
            "matter_id": str(matter.id),
            "preview_id": second_preview_id,
        },
    )
    assert second_save.status_code == 200, second_save.text
    filenames = [
        first_payload["output_filename"],
        second_save.json()["output_filename"],
    ]
    assert filenames[0] != filenames[1]

    documents = list(
        (
            await db_session.scalars(
                select(MatterDocument).where(MatterDocument.matter_id == matter.id)
            )
        ).all()
    )
    assert len(documents) == 2
    assert documents[0].storage_path != documents[1].storage_path
    assert all(Path(document.storage_path).is_file() for document in documents)
    events = list(
        (
            await db_session.scalars(
                select(MatterEvent).where(
                    MatterEvent.matter_id == matter.id,
                    MatterEvent.event_type == "document_generated",
                )
            )
        ).all()
    )
    assert len(events) == 2
    first_event = next(
        event
        for event in events
        if event.metadata_json["preview_evidence_id"] == generation_preview_id
    )
    assert first_event.metadata_json["template_source_sha256"] == template.source_sha256
    assert first_event.metadata_json["filled_variables"] == ["approved", "client_name"]
    assert "Ada Lovelace" not in str(first_event.metadata_json)
    evidence = await db_session.scalar(
        select(DocumentTemplatePreview).where(
            DocumentTemplatePreview.id == uuid.UUID(generation_preview_id)
        )
    )
    assert evidence.values_hmac_sha256
    assert "Ada Lovelace" not in str(evidence.__dict__)
    assert evidence.expires_at > evidence.created_at
    assert evidence.consumed_at is not None
    assert str(evidence.consumed_by_document_id) == first_payload["matter_document_id"]

    expiry_preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": generation_values,
            "matter_id": str(matter.id),
            "preview_purpose": "generation",
        },
    )
    assert expiry_preview.status_code == 200, expiry_preview.text
    expiry_preview_id = expiry_preview.headers["x-clarity-preview-id"]
    expiry_evidence = await db_session.scalar(
        select(DocumentTemplatePreview).where(
            DocumentTemplatePreview.id == uuid.UUID(expiry_preview_id)
        )
    )
    assert expiry_evidence.consumed_at is None
    expiry_evidence.expires_at = datetime.now(timezone.utc) - timedelta(seconds=1)
    await db_session.commit()
    expired_preview = await client.post(
        f"/api/templates/{template_id}/render",
        json={
            "variables": generation_values,
            "matter_id": str(matter.id),
            "preview_id": expiry_preview_id,
        },
    )
    assert expired_preview.status_code == 409
    assert "preview expired" in expired_preview.json()["detail"]


@pytest.mark.asyncio
async def test_pdf_save_rejects_output_that_differs_from_reviewed_preview_before_write(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from pathlib import Path

    from sqlalchemy import func, select

    from app.models.matter_document import MatterDocument
    from app.models.plugin import Matter, MatterEvent
    from app.services import matter_file_store as matter_store_module

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    monkeypatch.setattr(matter_store_module.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)

    created = await client.post(
        "/api/templates/intake/create",
        files={"file": ("integrity-form.pdf", _fillable_pdf(), "application/pdf")},
        data={"title": "Integrity Form", "category": "other"},
    )
    assert created.status_code == 201, created.text
    template_id = created.json()["id"]
    representative_values = {
        "client_name": "Representative Client",
        "approved": "yes",
        "notes": "Representative narrative",
    }
    activation_preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": representative_values,
            "preview_purpose": "activation",
        },
    )
    assert activation_preview.status_code == 200, activation_preview.text
    activated = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert activated.status_code == 200, activated.text

    matter = Matter(
        id=uuid.uuid4(),
        tenant_id=test_tenant.id,
        user_id=test_user.id,
        slug="pdf-integrity-matter",
        matter_name="PDF Integrity Matter",
        matter_type="general",
    )
    db_session.add(matter)
    await db_session.commit()
    generation_values = {
        "client_name": "Ada Lovelace",
        "approved": "yes",
        "notes": "Reviewed output",
    }
    generation_preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": generation_values,
            "matter_id": str(matter.id),
            "preview_purpose": "generation",
        },
    )
    assert generation_preview.status_code == 200, generation_preview.text
    preview_id = generation_preview.headers["x-clarity-preview-id"]
    files_before = {
        path.resolve() for path in Path(tmp_path).rglob("*") if path.is_file()
    }

    with monkeypatch.context() as mismatch:
        mismatch.setattr(
            document_templates,
            "fill_pdf_template",
            lambda *_args, **_kwargs: b"%PDF-1.7\n% deterministic mismatch\n%%EOF\n",
        )
        rejected = await client.post(
            f"/api/templates/{template_id}/render",
            json={
                "variables": generation_values,
                "matter_id": str(matter.id),
                "preview_id": preview_id,
            },
        )

    assert rejected.status_code == 409, rejected.text
    assert "does not match the reviewed preview" in rejected.json()["detail"]
    assert (
        await db_session.scalar(
            select(func.count())
            .select_from(MatterDocument)
            .where(MatterDocument.matter_id == matter.id)
        )
        == 0
    )
    assert (
        await db_session.scalar(
            select(func.count())
            .select_from(MatterEvent)
            .where(
                MatterEvent.matter_id == matter.id,
                MatterEvent.event_type == "document_generated",
            )
        )
        == 0
    )
    files_after = {
        path.resolve() for path in Path(tmp_path).rglob("*") if path.is_file()
    }
    assert files_after == files_before


@pytest.mark.asyncio
async def test_failed_commit_with_confirmed_absence_removes_local_staged_file(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from pathlib import Path

    from sqlalchemy import func, select

    from app.models.document_template_preview import DocumentTemplatePreview
    from app.models.matter_document import MatterDocument

    template_id, matter, values, preview_id = await _prepare_active_pdf_generation(
        client=client,
        db_session=db_session,
        test_tenant=test_tenant,
        test_user=test_user,
        tmp_path=tmp_path,
        monkeypatch=monkeypatch,
        slug="commit-failure-local",
    )
    files_before = {
        path.resolve() for path in Path(tmp_path).rglob("*") if path.is_file()
    }

    async def fail_commit():
        raise RuntimeError("injected pre-commit failure")

    async def confirmed_absent(**_kwargs):
        return False

    with monkeypatch.context() as failure:
        failure.setattr(db_session, "commit", fail_commit)
        failure.setattr(
            document_templates,
            "_matter_document_commit_outcome",
            confirmed_absent,
        )
        response = await client.post(
            f"/api/templates/{template_id}/render",
            json={
                "variables": values,
                "matter_id": str(matter.id),
                "preview_id": preview_id,
            },
        )

    assert response.status_code == 500, response.text
    assert "staged storage was removed" in response.json()["detail"]
    files_after = {
        path.resolve() for path in Path(tmp_path).rglob("*") if path.is_file()
    }
    assert files_after == files_before
    assert (
        await db_session.scalar(select(func.count()).select_from(MatterDocument)) == 0
    )
    evidence = await db_session.scalar(
        select(DocumentTemplatePreview)
        .where(DocumentTemplatePreview.id == uuid.UUID(preview_id))
        .execution_options(populate_existing=True)
    )
    assert evidence.consumed_at is None
    assert evidence.reconciliation_required_at is None


@pytest.mark.asyncio
async def test_cleanup_failure_blocks_preview_and_retry_performs_no_storage_write(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from sqlalchemy import select

    from app.models.document_template_preview import DocumentTemplatePreview
    from app.services.matter_file_store import StorageResult

    template_id, matter, values, preview_id = await _prepare_active_pdf_generation(
        client=client,
        db_session=db_session,
        test_tenant=test_tenant,
        test_user=test_user,
        tmp_path=tmp_path,
        monkeypatch=monkeypatch,
        slug="cleanup-reconciliation",
    )
    staged = StorageResult(
        provider="google",
        backend="google_drive",
        storage_path="https://attacker.invalid/display-only",
        provider_item_id="provider-item-123",
        drive_id="provider-drive-456",
    )
    store_calls = []
    cleanup_calls = []
    matter_id = str(matter.id)

    async def fake_store(**kwargs):
        store_calls.append(kwargs)
        return staged

    async def failed_cleanup(**kwargs):
        cleanup_calls.append(kwargs)
        raise RuntimeError("injected provider cleanup failure")

    async def fail_commit():
        raise RuntimeError("injected commit failure")

    async def confirmed_absent(**_kwargs):
        return False

    with monkeypatch.context() as failure:
        failure.setattr(
            document_templates.matter_file_store,
            "store_matter_file_result",
            fake_store,
        )
        failure.setattr(
            document_templates.matter_file_store,
            "delete_stored_result",
            failed_cleanup,
        )
        failure.setattr(db_session, "commit", fail_commit)
        failure.setattr(
            document_templates,
            "_matter_document_commit_outcome",
            confirmed_absent,
        )
        failed = await client.post(
            f"/api/templates/{template_id}/render",
            json={
                "variables": values,
                "matter_id": matter_id,
                "preview_id": preview_id,
            },
        )

    assert failed.status_code == 500, failed.text
    assert "cleanup failed" in failed.json()["detail"]
    assert len(store_calls) == 1
    assert len(cleanup_calls) == 1
    assert cleanup_calls[0]["result"] is staged
    evidence = await db_session.scalar(
        select(DocumentTemplatePreview)
        .where(DocumentTemplatePreview.id == uuid.UUID(preview_id))
        .execution_options(populate_existing=True)
    )
    assert evidence.consumed_at is None
    assert evidence.reconciliation_required_at is not None
    assert evidence.reconciliation_reason == "cleanup_failed"
    assert evidence.reconciliation_storage_backend == "google_drive"
    assert evidence.reconciliation_provider_item_id == "provider-item-123"
    assert evidence.reconciliation_provider_drive_id == "provider-drive-456"
    assert evidence.reconciliation_local_path is None

    async def unexpected_store(**_kwargs):
        raise AssertionError("blocked evidence must fail before another storage write")

    with monkeypatch.context() as retry_guard:
        retry_guard.setattr(
            document_templates.matter_file_store,
            "store_matter_file_result",
            unexpected_store,
        )
        retry = await client.post(
            f"/api/templates/{template_id}/render",
            json={
                "variables": values,
                "matter_id": matter_id,
                "preview_id": preview_id,
            },
        )
    assert retry.status_code == 409, retry.text
    assert "blocked pending storage reconciliation" in retry.json()["detail"]
    assert len(store_calls) == 1


@pytest.mark.asyncio
async def test_unknown_commit_outcome_preserves_storage_and_blocks_preview(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from sqlalchemy import select

    from app.models.document_template_preview import DocumentTemplatePreview
    from app.services.matter_file_store import StorageResult

    template_id, matter, values, preview_id = await _prepare_active_pdf_generation(
        client=client,
        db_session=db_session,
        test_tenant=test_tenant,
        test_user=test_user,
        tmp_path=tmp_path,
        monkeypatch=monkeypatch,
        slug="unknown-commit",
    )
    staged = StorageResult(
        provider="microsoft",
        backend="onedrive",
        storage_path="https://attacker.invalid/display-only",
        provider_item_id="unknown-item",
    )
    cleanup_calls = []
    matter_id = str(matter.id)

    async def fake_store(**_kwargs):
        return staged

    async def unexpected_cleanup(**kwargs):
        cleanup_calls.append(kwargs)

    async def fail_commit():
        raise RuntimeError("injected ambiguous commit")

    async def unknown_outcome(**_kwargs):
        return None

    with monkeypatch.context() as failure:
        failure.setattr(
            document_templates.matter_file_store,
            "store_matter_file_result",
            fake_store,
        )
        failure.setattr(
            document_templates.matter_file_store,
            "delete_stored_result",
            unexpected_cleanup,
        )
        failure.setattr(db_session, "commit", fail_commit)
        failure.setattr(
            document_templates,
            "_matter_document_commit_outcome",
            unknown_outcome,
        )
        response = await client.post(
            f"/api/templates/{template_id}/render",
            json={
                "variables": values,
                "matter_id": matter_id,
                "preview_id": preview_id,
            },
        )

    assert response.status_code == 500, response.text
    assert "outcome could not be verified" in response.json()["detail"]
    assert cleanup_calls == []
    evidence = await db_session.scalar(
        select(DocumentTemplatePreview)
        .where(DocumentTemplatePreview.id == uuid.UUID(preview_id))
        .execution_options(populate_existing=True)
    )
    assert evidence.reconciliation_reason == "commit_outcome_unknown"
    assert evidence.reconciliation_provider_item_id == "unknown-item"


@pytest.mark.asyncio
async def test_lost_commit_ack_preserves_committed_storage_and_replays_idempotently(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from pathlib import Path

    from sqlalchemy import func, select

    from app.models.matter_document import MatterDocument

    template_id, matter, values, preview_id = await _prepare_active_pdf_generation(
        client=client,
        db_session=db_session,
        test_tenant=test_tenant,
        test_user=test_user,
        tmp_path=tmp_path,
        monkeypatch=monkeypatch,
        slug="lost-commit-ack",
    )
    original_commit = db_session.commit
    matter_id = str(matter.id)
    files_before = {
        path.resolve() for path in Path(tmp_path).rglob("*") if path.is_file()
    }

    async def commit_then_lose_ack():
        await original_commit()
        raise RuntimeError("injected lost commit acknowledgement")

    with monkeypatch.context() as lost_ack:
        lost_ack.setattr(db_session, "commit", commit_then_lose_ack)
        first = await client.post(
            f"/api/templates/{template_id}/render",
            json={
                "variables": values,
                "matter_id": matter_id,
                "preview_id": preview_id,
            },
        )

    assert first.status_code == 200, first.text
    generated_id = first.json()["matter_document_id"]
    files_after = {
        path.resolve() for path in Path(tmp_path).rglob("*") if path.is_file()
    }
    assert len(files_after - files_before) == 1
    assert (
        await db_session.scalar(select(func.count()).select_from(MatterDocument)) == 1
    )

    replay = await client.post(
        f"/api/templates/{template_id}/render",
        json={
            "variables": values,
            "matter_id": matter_id,
            "preview_id": preview_id,
        },
    )
    assert replay.status_code == 200, replay.text
    assert replay.json()["matter_document_id"] == generated_id
    assert (
        await db_session.scalar(select(func.count()).select_from(MatterDocument)) == 1
    )


@pytest.mark.asyncio
async def test_pdf_intake_rejects_empty_non_form_and_active_documents(
    client, db_session, test_tenant, test_user
):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from pypdf import PdfReader, PdfWriter

    await _grant_manage_documents(db_session, test_tenant, test_user)
    empty = await client.post(
        "/api/templates/intake/analyze",
        files={"file": ("empty.pdf", b"", "application/pdf")},
    )
    assert empty.status_code == 400

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 720, "Static agreement without fields")
    pdf.save()
    non_form = output.getvalue()
    analysis = await client.post(
        "/api/templates/intake/analyze",
        files={"file": ("static.pdf", non_form, "application/pdf")},
    )
    assert analysis.status_code == 200
    assert any(
        "No reusable field locations" in warning
        for warning in analysis.json()["warnings"]
    )
    create = await client.post(
        "/api/templates/intake/create",
        files={"file": ("static.pdf", non_form, "application/pdf")},
    )
    assert create.status_code == 422

    writer = PdfWriter()
    writer.clone_document_from_reader(PdfReader(BytesIO(_fillable_pdf())))
    writer.add_js("app.alert('unsafe')")
    active_output = BytesIO()
    writer.write(active_output)
    active = await client.post(
        "/api/templates/intake/analyze",
        files={"file": ("active.pdf", active_output.getvalue(), "application/pdf")},
    )
    assert active.status_code == 422
    assert "Active PDF content" in active.json()["detail"]


@pytest.mark.asyncio
async def test_pdf_patch_revalidates_field_map_source_and_activation(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    import copy

    from sqlalchemy import select

    from app.models.document_template import DocumentTemplate

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)
    created = await client.post(
        "/api/templates/intake/create",
        files={"file": ("mapped.pdf", _fillable_pdf(), "application/pdf")},
    )
    assert created.status_code == 201, created.text
    template_id = created.json()["id"]
    valid_schema = created.json()["variable_schema"]

    duplicate_mapping = copy.deepcopy(valid_schema)
    duplicate_mapping["fields"][1]["pdf_field_name"] = duplicate_mapping["fields"][0][
        "pdf_field_name"
    ]
    duplicate = await client.patch(
        f"/api/templates/{template_id}",
        json={"variable_schema": duplicate_mapping},
    )
    assert duplicate.status_code == 422
    assert "Duplicate PDF field mapping" in duplicate.json()["detail"]

    phantom = copy.deepcopy(valid_schema)
    phantom["fields"].append({"name": "phantom", "label": "Phantom"})
    unmapped = await client.patch(
        f"/api/templates/{template_id}", json={"variable_schema": phantom}
    )
    assert unmapped.status_code == 422
    assert "missing pdf_field_name" in unmapped.json()["detail"]

    phantom_body = await client.patch(
        f"/api/templates/{template_id}", json={"body": "{{phantom}}"}
    )
    assert phantom_body.status_code == 422
    assert "without source mappings" in phantom_body.json()["detail"]

    blocked_activation = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert blocked_activation.status_code == 409
    assert "representative flattened PDF preview" in blocked_activation.json()["detail"]

    draft_preview = await client.post(
        f"/api/templates/{template_id}/render-file", json={"variables": {}}
    )
    assert draft_preview.status_code == 200, draft_preview.text
    still_blocked = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert still_blocked.status_code == 409

    blank_activation_preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={"variables": {}, "preview_purpose": "activation"},
    )
    assert blank_activation_preview.status_code == 422
    assert "representative values" in blank_activation_preview.json()["detail"]

    activation_values = {
        "client_name": "Representative Client",
        "approved": "true",
        "notes": "Representative narrative",
    }
    preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": activation_values,
            "preview_purpose": "activation",
        },
    )
    assert preview.status_code == 200, preview.text
    assert preview.headers["x-clarity-preview-purpose"] == "activation"

    activated = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert activated.status_code == 200, activated.text
    assert activated.json()["is_active"] is True
    assert activated.json()["last_test_rendered_at"] is not None
    assert activated.json()["approved_at"] is not None
    assert activated.json()["approved_by_user_id"] == str(test_user.id)

    combined_edit_activation = await client.patch(
        f"/api/templates/{template_id}",
        json={"body": "{{client_name}}", "is_active": True},
    )
    assert combined_edit_activation.status_code == 409
    assert (
        "Preview the updated PDF successfully"
        in combined_edit_activation.json()["detail"]
    )

    metadata_only = await client.patch(
        f"/api/templates/{template_id}", json={"title": "Mapped PDF v2"}
    )
    assert metadata_only.status_code == 200, metadata_only.text
    assert metadata_only.json()["is_active"] is True
    assert metadata_only.json()["approved_at"] == activated.json()["approved_at"]

    edited_contract = await client.patch(
        f"/api/templates/{template_id}", json={"body": "{{client_name}}"}
    )
    assert edited_contract.status_code == 200, edited_contract.text
    assert edited_contract.json()["is_active"] is False
    assert edited_contract.json()["last_test_rendered_at"] is None
    assert edited_contract.json()["approved_at"] is None
    assert edited_contract.json()["approved_by_user_id"] is None

    edited_activation = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert edited_activation.status_code == 409
    assert "representative flattened PDF preview" in edited_activation.json()["detail"]

    second_preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": activation_values,
            "preview_purpose": "activation",
        },
    )
    assert second_preview.status_code == 200, second_preview.text
    reactivated = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert reactivated.status_code == 200, reactivated.text
    assert reactivated.json()["is_active"] is True

    template = await db_session.scalar(
        select(DocumentTemplate).where(DocumentTemplate.id == uuid.UUID(template_id))
    )
    template.source_sha256 = None
    await db_session.commit()
    missing_integrity = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert missing_integrity.status_code == 409
    assert "integrity check" in missing_integrity.json()["detail"]


@pytest.mark.asyncio
async def test_pdf_creation_rejects_unmapped_reviewed_body_and_json_shortcut(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)
    phantom_body = await client.post(
        "/api/templates/intake/create",
        files={"file": ("mapped.pdf", _fillable_pdf(), "application/pdf")},
        data={"reviewed_body": "Matter: {{phantom}}"},
    )
    assert phantom_body.status_code == 422
    assert "without source mappings" in phantom_body.json()["detail"]

    analysis = await client.post(
        "/api/templates/intake/analyze",
        files={"file": ("excluded.pdf", _fillable_pdf(), "application/pdf")},
    )
    assert analysis.status_code == 200, analysis.text
    analyzed = analysis.json()
    reviewed_schema = analyzed["suggested_variable_schema"]
    excluded_name = reviewed_schema["fields"][0]["name"]
    reviewed_schema["fields"][0]["included"] = False
    excluded_placeholder = await client.post(
        "/api/templates/intake/create",
        files={"file": ("excluded.pdf", _fillable_pdf(), "application/pdf")},
        data={
            "analysis_token": analyzed["analysis_token"],
            "variable_schema": json.dumps(reviewed_schema),
            "reviewed_body": f"Excluded: {{{{{excluded_name}}}}}",
        },
    )
    assert excluded_placeholder.status_code == 422
    assert "without source mappings" in excluded_placeholder.json()["detail"]

    shortcut = await client.post(
        "/api/templates",
        json={
            "title": "Unsafe PDF shortcut",
            "body": "{{client_name}}",
            "format": "pdf",
            "variable_schema": {
                "fields": [{"name": "client_name", "pdf_field_name": "client_name"}]
            },
        },
    )
    assert shortcut.status_code == 422
    assert "multipart" in shortcut.json()["detail"]


@pytest.mark.asyncio
async def test_manual_only_pdf_can_be_previewed_and_activated(
    client, db_session, test_tenant, test_user, tmp_path, monkeypatch
):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    from app.services.template_intake import TemplateAnalysis

    monkeypatch.setattr(document_templates.settings, "UPLOAD_DIR", str(tmp_path))
    await _grant_manage_documents(db_session, test_tenant, test_user)
    source_buffer = BytesIO()
    pdf = canvas.Canvas(source_buffer, pagesize=letter)
    pdf.showPage()
    pdf.save()
    source = source_buffer.getvalue()
    discovered = TemplateAnalysis(
        title="Manual only",
        format="pdf",
        body="",
        body_preview="",
        extracted_text="",
        variable_schema={
            "version": 1,
            "source": "pdf_ocr_overlay",
            "pages": [{"page": 1, "width": 612, "height": 792, "rotation": 0}],
            "fields": [],
            "detection": {"method": "none", "label": "Manual review"},
        },
        branding_profile={},
        warnings=[],
    )
    monkeypatch.setattr(
        document_templates,
        "analyze_template_upload",
        lambda **_kwargs: discovered,
    )
    manual_key = f"manual:{uuid.uuid4()}"
    reviewed_schema = {
        **discovered.variable_schema,
        "fields": [{
            "name": "manual_name",
            "label": "Name",
            "pdf_source_key": manual_key,
            "pdf_overlay": {"page": 1, "rect": [72, 700, 260, 724]},
            "field_type": "text",
            "required": True,
            "multiline": False,
            "included": True,
        }],
    }
    created = await client.post(
        "/api/templates/intake/create",
        files={"file": ("manual-only.pdf", source, "application/pdf")},
        data={
            "reviewed_body": "{{manual_name}}",
            "variable_schema": json.dumps(reviewed_schema),
        },
    )
    assert created.status_code == 201, created.text
    template_id = created.json()["id"]

    preview = await client.post(
        f"/api/templates/{template_id}/render-file",
        json={
            "variables": {"manual_name": "Representative Name"},
            "preview_purpose": "activation",
        },
    )
    assert preview.status_code == 200, preview.text
    activated = await client.patch(
        f"/api/templates/{template_id}", json={"is_active": True}
    )
    assert activated.status_code == 200, activated.text
    assert activated.json()["is_active"] is True


@pytest.mark.asyncio
async def test_document_template_sensitive_routes_require_manage_documents(
    client, db_session, test_tenant, test_user
):
    from app.models.document_template import DocumentTemplate
    from app.models.rbac import Role, UserRole

    limited_role = Role(
        tenant_id=test_tenant.id,
        name="Matter-only staff",
        capabilities=["manage_matters"],
    )
    db_session.add(limited_role)
    await db_session.flush()
    db_session.add(
        UserRole(
            user_id=test_user.id,
            role_id=limited_role.id,
            tenant_id=test_tenant.id,
            source="manual",
        )
    )
    template = DocumentTemplate(
        tenant_id=test_tenant.id,
        title="Restricted template",
        body="Hello {{client_name}}",
        category="other",
        format="markdown",
        is_active=True,
    )
    db_session.add(template)
    await db_session.commit()

    listed = await client.get("/api/templates")
    assert listed.status_code == 200
    mutation = await client.post(
        "/api/templates",
        json={"title": "Denied", "body": "Body", "format": "markdown"},
    )
    source = await client.get(f"/api/templates/{template.id}/source")
    rendered = await client.post(
        f"/api/templates/{template.id}/render",
        json={"variables": {"client_name": "Ada"}},
    )
    assert mutation.status_code == 403
    assert source.status_code == 403
    assert rendered.status_code == 403
    assert all(
        response.json()["detail"] == "Missing capability: manage_documents"
        for response in (mutation, source, rendered)
    )


@pytest.mark.asyncio
async def test_pdf_mime_normalization_blocks_active_extensionless_uploads(client):
    from pypdf import PdfReader, PdfWriter

    benign = await client.post(
        "/api/templates/intake/analyze",
        files={"file": ("upload", _fillable_pdf(), "Application/PDF; charset=binary")},
    )
    assert benign.status_code == 200, benign.text
    assert benign.json()["format"] == "pdf"

    writer = PdfWriter()
    writer.clone_document_from_reader(PdfReader(BytesIO(_fillable_pdf())))
    writer.add_js("app.alert('unsafe')")
    active = BytesIO()
    writer.write(active)
    rejected = await client.post(
        "/api/templates/intake/analyze",
        files={
            "file": (
                "upload",
                active.getvalue(),
                "Application/PDF; charset=binary",
            )
        },
    )
    assert rejected.status_code == 422
    assert "Active PDF content" in rejected.json()["detail"]

    disguised = await client.post(
        "/api/templates/intake/analyze",
        files={"file": ("upload.txt", active.getvalue(), "text/plain")},
    )
    assert disguised.status_code == 422
    assert "Active PDF content" in disguised.json()["detail"]


def _document_image_bytes(*, image_format: str = "PNG") -> bytes:
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (1200, 400), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 80), "Client name:", fill="black")
    output = BytesIO()
    image.save(output, format=image_format, dpi=(150, 150))
    image.close()
    return output.getvalue()


def test_pdf_page_preview_reports_authoritative_geometry_and_png():
    from PIL import Image
    from reportlab.pdfgen import canvas

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=(612, 792))
    pdf.drawString(72, 720, "Page one")
    pdf.showPage()
    pdf.drawString(72, 720, "Page two")
    pdf.save()

    image_bytes, metadata = render_pdf_page_preview(output.getvalue(), 2)

    with Image.open(BytesIO(image_bytes)) as image:
        assert image.format == "PNG"
        assert image.width * image.height <= 4_100_000
    assert metadata == {
        "page": 2,
        "width": 612.0,
        "height": 792.0,
        "left": 0.0,
        "bottom": 0.0,
        "right": 612.0,
        "top": 792.0,
        "rotation": 0,
        "page_count": 2,
    }


def test_standalone_image_is_normalized_to_bounded_pdf():
    normalized = image_to_pdf(_document_image_bytes())

    assert normalized.content.startswith(b"%PDF-")
    assert normalized.pages == 1
    assert normalized.image_format == "PNG"
    metadata = pdf_page_metadata(normalized.content)
    assert len(metadata) == 1
    assert metadata[0]["width"] == pytest.approx(576, abs=1)
    assert metadata[0]["height"] == pytest.approx(192, abs=1)


def test_image_normalization_uses_exif_corrected_page_geometry():
    from PIL import Image
    from pypdf import PdfReader

    source = BytesIO()
    image = Image.new("RGB", (120, 60), "white")
    exif = Image.Exif()
    exif[274] = 6
    image.save(source, format="JPEG", exif=exif)
    image.close()

    normalized = image_to_pdf(source.getvalue())
    page = PdfReader(BytesIO(normalized.content)).pages[0]

    assert float(page.mediabox.height) > float(page.mediabox.width)


def test_unsupported_gif_image_is_rejected():
    with pytest.raises(TemplateOcrError, match="Unsupported image format"):
        image_to_pdf(_document_image_bytes(image_format="GIF"))


@pytest.mark.asyncio
async def test_image_upload_uses_ocr_and_returns_pdf_mapping(monkeypatch):
    from app.services import template_intake
    from starlette.datastructures import Headers, UploadFile

    monkeypatch.setattr(
        template_intake,
        "ocr_pdf",
        lambda _content, **_kwargs: PdfOcrResult(
            text="Client name:",
            lines=(
                OcrLine(
                    page_index=0,
                    text="Client name:",
                    score=0.93,
                    rect=(30, 90, 150, 115),
                ),
            ),
            pages_analyzed=1,
            pages_total=1,
            average_confidence=0.93,
            truncated=False,
            page_indexes=(0,),
        ),
    )

    sample = await document_templates._read_template_sample(
        UploadFile(
            file=BytesIO(_document_image_bytes()),
            filename="intake.png",
            headers=Headers({"content-type": "image/png"}),
        )
    )
    analysis = analyze_template_upload(
        file_bytes=sample.content,
        filename=sample.filename,
        content_type=sample.content_type,
    )
    analysis.warnings.extend(sample.warnings)
    payload = analysis.as_dict()
    assert payload["format"] == "pdf"
    assert payload["suggested_variable_schema"]["detection"]["method"] == "ocr"
    assert payload["suggested_variable_schema"]["detection"]["ocr_pages"] == [1]
    assert payload["suggested_variable_schema"]["fields"][0]["name"] == "client_name"
    assert any("converted to a safe 1-page PDF" in warning for warning in payload["warnings"])


def test_mixed_pdf_merges_native_text_and_only_sparse_page_ocr(monkeypatch):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    from app.services import template_intake

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    pdf.drawString(72, 720, "Client: Alice Example")
    pdf.showPage()
    pdf.rect(72, 650, 300, 40)
    pdf.showPage()
    pdf.save()

    scanned_pages = []

    def fake_ocr(_content, *, page_indexes, **_kwargs):
        scanned_pages.extend(page_indexes)
        return PdfOcrResult(
            text="Case Number: BETA-77",
            lines=(
                OcrLine(
                    page_index=1,
                    text="Case Number: BETA-77",
                    score=0.91,
                    rect=(72, 650, 300, 675),
                ),
            ),
            pages_analyzed=1,
            pages_total=2,
            average_confidence=0.91,
            truncated=False,
            page_indexes=(1,),
        )

    monkeypatch.setattr(template_intake, "ocr_pdf", fake_ocr)
    analysis = analyze_template_upload(
        file_bytes=output.getvalue(),
        filename="mixed.pdf",
        content_type="application/pdf",
    )

    assert scanned_pages == [1]
    fields = {
        field["name"]: field
        for field in analysis.variable_schema["fields"]
    }
    assert {"client_name", "case_number"} <= set(fields)
    assert fields["client_name"]["pdf_overlay"]["source_kind"] == "text"
    assert fields["case_number"]["pdf_overlay"]["source_kind"] == "ocr"
    assert analysis.variable_schema["detection"]["ocr_pages"] == [2]


def test_docx_package_rejects_embedded_payloads():
    import zipfile

    from docx import Document

    source = BytesIO()
    document = Document()
    document.add_paragraph("Clean source")
    document.save(source)
    with zipfile.ZipFile(source, mode="a") as package:
        package.writestr("word/embeddings/embedded.bin", b"not allowed")

    with pytest.raises(TemplateDocxError, match="embedded files"):
        validate_docx_package(source.getvalue())


def test_docx_package_rejects_duplicate_internal_parts():
    import warnings
    import zipfile

    from docx import Document

    source = BytesIO()
    document = Document()
    document.add_paragraph("Clean source")
    document.save(source)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        with zipfile.ZipFile(source, mode="a") as package:
            package.writestr(
                "word/document.xml",
                package.read("word/document.xml"),
            )

    with pytest.raises(TemplateDocxError, match="duplicate internal parts"):
        validate_docx_package(source.getvalue())


def test_docx_content_control_placeholder_is_detected_and_rendered():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    document.add_paragraph("Visible paragraph")
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "{{wrapped_client_name}}"
    run.append(text)
    paragraph.append(run)
    content.append(paragraph)
    sdt.append(content)
    body = document.element.body
    body.insert(len(body) - 1, sdt)
    source = BytesIO()
    document.save(source)

    analysis = analyze_template_upload(
        file_bytes=source.getvalue(),
        filename="wrapped.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    assert [
        field["name"] for field in analysis.variable_schema["fields"]
    ] == ["wrapped_client_name"]

    rendered = fill_docx_template(
        source.getvalue(),
        variable_schema=analysis.variable_schema,
        variables={"wrapped_client_name": "Donna Price"},
    )
    rendered_document = Document(BytesIO(rendered))
    all_text = " ".join(
        node.text or ""
        for node in rendered_document.element.body.iter(qn("w:t"))
    )
    assert "Donna Price" in all_text
    assert "wrapped_client_name" not in all_text


def test_ai_docx_proposals_require_one_exact_source_location():
    from docx import Document

    document = Document()
    document.add_paragraph("Internal reference ALPHA-123")
    document.add_paragraph("Repeated token")
    document.add_paragraph("Repeated token")
    source = BytesIO()
    document.save(source)
    analysis = analyze_template_upload(
        file_bytes=source.getvalue(),
        filename="uncurated.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    mapped, unmapped = reconcile_ai_template_fields(
        analysis=analysis,
        file_bytes=source.getvalue(),
        proposals=[
            AiFieldProposal(
                name="internal_reference",
                label="Internal reference",
                source_text="ALPHA-123",
                confidence=0.96,
                reason="A reusable identifier.",
            ),
            AiFieldProposal(
                name="ambiguous_value",
                label="Ambiguous value",
                source_text="Repeated token",
                confidence=0.9,
                reason="Appears to repeat.",
            ),
        ],
    )

    assert len(mapped) == 1
    assert mapped[0]["name"] == "internal_reference"
    assert mapped[0]["confidence"] == 0.75
    assert mapped[0]["review_required"] is True
    assert mapped[0]["docx_anchor"]["paragraph_ordinal"] >= 0
    assert len(unmapped) == 1
    assert "more than once" in unmapped[0]["unmapped_reason"]


def test_ai_can_improve_existing_docx_field_without_moving_its_anchor():
    from docx import Document

    document = Document()
    document.add_paragraph("Yes: __________")
    source = BytesIO()
    document.save(source)
    content = source.getvalue()
    analysis = analyze_template_upload(
        file_bytes=content,
        filename="questionnaire.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    original = analysis.variable_schema["fields"][0]
    original_anchor = dict(original["docx_anchor"])

    mapped, unmapped = reconcile_ai_template_fields(
        analysis=analysis,
        file_bytes=content,
        proposals=[
            AiFieldProposal(
                existing_name=original["name"],
                name="client_consents",
                label="Client consents",
                source_text=original["source_text"],
                field_type="checkbox",
                confidence=0.93,
                reason="The label describes a yes/no response.",
            )
        ],
    )

    assert unmapped == []
    assert len(mapped) == 1
    assert len(analysis.variable_schema["fields"]) == 1
    updated = analysis.variable_schema["fields"][0]
    assert updated["name"] == "client_consents"
    assert updated["label"] == "Client consents"
    assert updated["field_type"] == "checkbox"
    assert updated["docx_anchor"] == original_anchor
    assert updated["ai_update_kind"] == "updated"


def test_submitted_ai_update_is_reconciled_against_fresh_docx_analysis():
    from docx import Document

    document = Document()
    document.add_paragraph("Yes: __________")
    source = BytesIO()
    document.save(source)
    content = source.getvalue()
    proposed = analyze_template_upload(
        file_bytes=content,
        filename="questionnaire.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    current_name = proposed.variable_schema["fields"][0]["name"]
    reconcile_ai_template_fields(
        analysis=proposed,
        file_bytes=content,
        proposals=[
            AiFieldProposal(
                existing_name=current_name,
                name="client_consents",
                label="Client consents",
                source_text="__________",
                field_type="checkbox",
                confidence=0.9,
                reason="A clearer reusable-field name.",
            )
        ],
    )
    raw_schema = json.dumps(proposed.variable_schema)
    fresh = analyze_template_upload(
        file_bytes=content,
        filename="questionnaire.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    document_templates._reconcile_submitted_ai_fields(
        raw_schema,
        analysis=fresh,
        file_bytes=content,
    )

    fields = fresh.variable_schema["fields"]
    assert len(fields) == 1
    assert fields[0]["name"] == "client_consents"
    assert fields[0]["ai_update_kind"] == "updated"
    assert fresh.variable_schema["detection"]["ai_added_count"] == 0
    assert fresh.variable_schema["detection"]["ai_updated_count"] == 1


def test_docx_renderer_rejects_a_tampered_reviewed_anchor():
    from docx import Document

    document = Document()
    document.add_paragraph("Client: __________")
    source = BytesIO()
    document.save(source)
    content = source.getvalue()
    analysis = analyze_template_upload(
        file_bytes=content,
        filename="agreement.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    field = analysis.variable_schema["fields"][0]
    field["docx_anchor"]["start"] += 1
    field["docx_anchor"]["end"] += 1

    with pytest.raises(TemplateDocxError, match="reviewed location"):
        fill_docx_template(
            content,
            variable_schema=analysis.variable_schema,
            variables={},
        )


def test_reviewed_docx_schema_restores_server_authoritative_identical_blank_locations():
    from docx import Document

    document = Document()
    document.add_paragraph("Client name: __________")
    document.add_paragraph("Matter name: __________")
    source = BytesIO()
    document.save(source)
    analysis = analyze_template_upload(
        file_bytes=source.getvalue(),
        filename="identical-blanks.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    original_fields = analysis.variable_schema["fields"]
    submitted = json.loads(json.dumps(analysis.variable_schema))
    submitted["fields"][0]["docx_anchor"] = dict(
        submitted["fields"][1]["docx_anchor"]
    )

    reviewed = document_templates._reviewed_variable_schema(
        json.dumps(submitted),
        analysis.variable_schema,
    )

    assert reviewed["fields"][0]["docx_source_key"] == original_fields[0][
        "docx_source_key"
    ]
    assert reviewed["fields"][0]["docx_anchor"] == original_fields[0][
        "docx_anchor"
    ]
    assert reviewed["fields"][0]["docx_anchor"] != original_fields[1][
        "docx_anchor"
    ]


def test_premium_ai_evidence_redacts_obvious_identifiers():
    redacted = _redact_evidence(
        "SSN 123-45-6789 email client@example.com phone (701) 555-0100 account 123456789012"
    )

    assert "123-45-6789" not in redacted
    assert "client@example.com" not in redacted
    assert "701" not in redacted
    assert "123456789012" not in redacted
    assert "[REDACTED_SSN]" in redacted
    assert "[REDACTED_ACCOUNT_NUMBER]" in redacted


@pytest.mark.asyncio
async def test_premium_ai_requires_explicit_consent_before_any_model_work():
    analysis = analyze_template_upload(
        file_bytes=b"Client: Example Person",
        filename="sample.txt",
        content_type="text/plain",
    )

    with pytest.raises(TemplateAiAssistError, match="Confirm"):
        await run_template_ai_mapping(
            db=object(),
            user=object(),
            analysis=analysis,
            file_bytes=b"Client: Example Person",
            consent_to_external_ai=False,
        )


@pytest.mark.asyncio
async def test_premium_ai_proposal_is_audited_and_reconciled_locally(monkeypatch):
    from app.services import template_ai_service

    class FakeDatabase:
        def __init__(self):
            self.added = []
            self.commits = 0

        def add(self, value):
            self.added.append(value)

        async def commit(self):
            self.commits += 1

    class FakeLlm:
        async def complete(self, **_kwargs):
            return (
                '{"document_type":"agreement","fields":['
                '{"name":"internal_reference","label":"Internal reference",'
                '"source_text":"ALPHA-123","field_type":"text",'
                '"confidence":0.94,"reason":"Reusable identifier"}],'
                '"warnings":[]}',
                120,
                35,
            )

    async def allow_budget(_db, _user):
        return None

    async def premium_route(_db, _tenant_id, *, use_premium):
        assert use_premium is True
        return SimpleNamespace(
            model="premium-test",
            provider="litellm",
            customer_api_key="tenant-key",
            customer_provider="openai",
            customer_endpoint=None,
            requested_route="premium",
            resolved_route="customer",
            gateway_provider="customer",
            gateway_alias="premium-test",
        )

    monkeypatch.setattr(template_ai_service, "check_token_budget", allow_budget)
    monkeypatch.setattr(template_ai_service, "resolve_llm_route", premium_route)
    database = FakeDatabase()
    user = SimpleNamespace(
        id=uuid.uuid4(),
        tenant_id=uuid.uuid4(),
        tenant=SimpleNamespace(name="Test firm", billing_tier="payg"),
    )
    analysis = analyze_template_upload(
        file_bytes=b"Internal reference ALPHA-123",
        filename="sample.txt",
        content_type="text/plain",
    )

    result = await run_template_ai_mapping(
        db=database,
        user=user,
        analysis=analysis,
        file_bytes=b"Internal reference ALPHA-123",
        consent_to_external_ai=True,
        llm=FakeLlm(),
    )

    fields = result.variable_schema["fields"]
    assert fields[0]["name"] == "internal_reference"
    assert fields[0]["ai_suggested"] is True
    assert fields[0]["confidence"] == 0.75
    assert result.variable_schema["detection"]["ai_added_count"] == 1
    assert database.added[0].operation_type == "template_ai_map"
    assert database.commits == 1
