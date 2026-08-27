from __future__ import annotations

import io
import hashlib
from datetime import datetime, timezone
from types import SimpleNamespace
from uuid import uuid4

import pytest
from docx import Document
from pydantic import ValidationError

from app.schemas.workspace_mcp import ProposeDocumentFromTemplateArgs
from app.services.automation_capabilities import CapabilityContext, CapabilityError
from app.services import document_template_workspace as workspace
from app.services.template_intake import analyze_template_upload


class _DB:
    def __init__(self, values):
        self.values = list(values)

    async def scalar(self, _statement):
        return self.values.pop(0)


def _context(db, tenant_id):
    return CapabilityContext(
        db=db,
        user=SimpleNamespace(id=uuid4(), tenant_id=tenant_id),
        channel="workspace_mcp",
        granted_scopes=frozenset(
            {"matters:read", "templates:read", "documents:propose"}
        ),
    )


def _matter(tenant_id):
    return SimpleNamespace(
        id=uuid4(),
        tenant_id=tenant_id,
        jurisdiction="North Dakota",
        stage="discovery",
        primary_plugin="civil",
    )


def _template(tenant_id, *, body="Dear {{client_name}}"):
    return SimpleNamespace(
        id=uuid4(),
        tenant_id=tenant_id,
        title="Client Letter",
        category="correspondence",
        kind="letter",
        description="Approved client letter",
        body=body,
        format="markdown",
        status="approved",
        approved_at=datetime(2026, 1, 1, tzinfo=timezone.utc),
        is_active=True,
        source_storage_path=None,
        source_sha256=None,
        source_filename=None,
        variable_schema={
            "properties": {"client_name": {"type": "string"}},
            "required": ["client_name"],
        },
        jurisdiction="North Dakota",
        stage="discovery",
        module="civil",
    )


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("A deterministic review document")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _variable_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Dear {{client_name}}")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_canonical_docx_bytes_remove_zip_timestamp_variance():
    content = _docx_bytes()

    first = workspace._canonical_docx_bytes(content)
    second = workspace._canonical_docx_bytes(content)

    assert first == second
    assert workspace.inspect_cloud_docx_snapshot(first).review_text == (
        "A deterministic review document"
    )


@pytest.mark.asyncio
async def test_markdown_template_render_binds_variables_and_provenance():
    tenant_id = uuid4()
    matter = _matter(tenant_id)
    template = _template(tenant_id)

    rendered = await workspace.render_workspace_template(
        _context(_DB([matter, template]), tenant_id),
        matter_id=matter.id,
        template_id=template.id,
        variables={"client_name": "Avery Client"},
        title=None,
    )

    assert rendered.review_text == "Dear Avery Client"
    assert rendered.template is template
    assert rendered.template_format == "markdown"
    assert rendered.source_docx_bytes is None
    assert len(rendered.template_sha256) == 64
    assert rendered.variable_snapshot == {"client_name": "Avery Client"}


@pytest.mark.asyncio
async def test_docx_template_verifies_source_and_returns_exact_rendered_bytes(
    monkeypatch, tmp_path
):
    tenant_id = uuid4()
    matter = _matter(tenant_id)
    template = _template(tenant_id)
    template.format = "docx"
    source = _variable_docx_bytes()
    template.variable_schema = analyze_template_upload(
        file_bytes=source,
        filename="client-letter.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument." "wordprocessingml.document"
        ),
    ).variable_schema
    source_path = (
        tmp_path / str(tenant_id) / "templates" / str(template.id) / "source.docx"
    )
    source_path.parent.mkdir(parents=True)
    source_path.write_bytes(source)
    template.source_storage_path = str(source_path)
    template.source_sha256 = hashlib.sha256(source).hexdigest()
    monkeypatch.setattr(workspace.settings, "UPLOAD_DIR", str(tmp_path))

    rendered = await workspace.render_workspace_template(
        _context(_DB([matter, template]), tenant_id),
        matter_id=matter.id,
        template_id=template.id,
        variables={"client_name": "Avery Client"},
        title="Client Letter",
    )

    snapshot = workspace.inspect_cloud_docx_snapshot(rendered.source_docx_bytes)
    assert rendered.template_format == "docx"
    assert rendered.template_sha256 == template.source_sha256
    assert snapshot.review_text == "Dear Avery Client"


@pytest.mark.asyncio
async def test_template_render_rejects_missing_and_unknown_variables():
    tenant_id = uuid4()
    matter = _matter(tenant_id)
    template = _template(tenant_id)

    with pytest.raises(CapabilityError) as missing:
        await workspace.render_workspace_template(
            _context(_DB([matter, template]), tenant_id),
            matter_id=matter.id,
            template_id=template.id,
            variables={},
            title=None,
        )
    assert missing.value.code == "required_template_variable_missing"

    with pytest.raises(CapabilityError) as unknown:
        await workspace.render_workspace_template(
            _context(_DB([matter, template]), tenant_id),
            matter_id=matter.id,
            template_id=template.id,
            variables={"client_name": "Avery", "send_to": "outside@example.test"},
            title=None,
        )
    assert unknown.value.code == "unknown_template_variable"


@pytest.mark.asyncio
async def test_template_must_match_matter_workflow_and_jurisdiction():
    tenant_id = uuid4()
    matter = _matter(tenant_id)
    template = _template(tenant_id)
    template.jurisdiction = "Montana"

    with pytest.raises(CapabilityError) as exc_info:
        await workspace.require_workspace_template(
            _context(_DB([matter, template]), tenant_id),
            matter_id=matter.id,
            template_id=template.id,
        )

    assert exc_info.value.code == "template_incompatible"


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("template_format", "body", "title", "code"),
    [
        ("pdf", "PDF text", None, "template_preview_required"),
        ("rtf", "RTF text", None, "unsupported_template_format"),
        ("markdown", "Plain text", "unsafe/path", "invalid_document_title"),
        ("markdown", "", None, "template_not_found"),
    ],
)
async def test_template_render_fails_closed_for_unsafe_or_unreviewable_inputs(
    template_format, body, title, code
):
    tenant_id = uuid4()
    matter = _matter(tenant_id)
    template = _template(tenant_id, body=body)
    template.format = template_format
    if "{{" not in body:
        template.variable_schema = {}

    with pytest.raises(CapabilityError) as exc_info:
        await workspace.render_workspace_template(
            _context(_DB([matter, template]), tenant_id),
            matter_id=matter.id,
            template_id=template.id,
            variables={"client_name": "Avery Client"} if "{{" in body else {},
            title=title,
        )

    assert exc_info.value.code == code


@pytest.mark.asyncio
async def test_template_source_rejects_escape_and_hash_mismatch(monkeypatch, tmp_path):
    tenant_id = uuid4()
    template = _template(tenant_id)
    outside = tmp_path / "outside.docx"
    outside.write_bytes(b"outside")
    template.source_storage_path = str(outside)
    template.source_sha256 = hashlib.sha256(b"outside").hexdigest()
    monkeypatch.setattr(workspace.settings, "UPLOAD_DIR", str(tmp_path))

    with pytest.raises(CapabilityError) as escaped:
        await workspace.verified_template_source(template)
    assert escaped.value.code == "template_source_unavailable"

    inside = tmp_path / str(tenant_id) / "templates" / str(template.id) / "source.docx"
    inside.parent.mkdir(parents=True)
    inside.write_bytes(b"tampered")
    template.source_storage_path = str(inside)
    with pytest.raises(CapabilityError) as mismatch:
        await workspace.verified_template_source(template)
    assert mismatch.value.code == "template_integrity_failed"


def test_template_proposal_variables_are_strictly_bounded():
    common = {"matter_id": uuid4(), "template_id": uuid4()}

    with pytest.raises(ValidationError):
        ProposeDocumentFromTemplateArgs(
            **common,
            variables={"client_name": "x" * 10_001},
        )
    with pytest.raises(ValidationError):
        ProposeDocumentFromTemplateArgs(
            **common,
            variables={"a": "x" * 7_000, "b": "y" * 7_000, "c": "z" * 7_000},
        )
    with pytest.raises(ValidationError):
        ProposeDocumentFromTemplateArgs(**common, variables={" ": "value"})
