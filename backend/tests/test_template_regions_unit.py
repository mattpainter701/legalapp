"""Regions stored as paragraph ranges, and how they reach the renderer."""

import io

import pytest
from docx import Document

from app.services.docx_templates import (
    TemplateDocxError,
    _open_docx,
    fill_docx_template,
    inject_stored_regions,
    iter_docx_paragraphs,
)
from app.services.template_regions import (
    MAX_REGION_NESTING,
    MAX_REGIONS,
    TemplateRegion,
    TemplateRegionError,
    parse_regions,
    stored_regions,
)


def _docx(lines):
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _text(content):
    return [
        paragraph.text
        for paragraph in iter_docx_paragraphs(_open_docx(content))
        if paragraph.text.strip()
    ]


def _raw(kind, name, start, end):
    return {"kind": kind, "name": name, "from_ordinal": start, "to_ordinal": end}


class TestParsing:
    def test_a_valid_region_round_trips(self):
        [region] = parse_regions([_raw("if", "entity", 2, 4)], known_fields={"entity"})
        assert region == TemplateRegion("if", "entity", 2, 4)
        assert region.as_dict() == _raw("if", "entity", 2, 4)

    def test_regions_are_returned_in_document_order_outermost_first(self):
        regions = parse_regions(
            [_raw("if", "b", 3, 4), _raw("if", "a", 0, 9)], known_fields={"a", "b"}
        )
        assert [(r.name, r.from_ordinal) for r in regions] == [("a", 0), ("b", 3)]

    @pytest.mark.parametrize(
        "raw",
        [
            "not-a-list",
            [7],
            [_raw("switch", "a", 0, 1)],
            [_raw("if", "", 0, 1)],
            [_raw("if", "9bad", 0, 1)],
            [_raw("if", "a", -1, 1)],
            [_raw("if", "a", 4, 2)],
            [{"kind": "if", "name": "a", "from_ordinal": "x", "to_ordinal": 1}],
            [{"kind": "if", "name": "a", "from_ordinal": 0}],
        ],
    )
    def test_malformed_regions_are_rejected(self, raw):
        with pytest.raises(TemplateRegionError):
            parse_regions(raw, known_fields={"a"})

    def test_a_condition_on_an_undeclared_field_is_rejected(self):
        with pytest.raises(TemplateRegionError):
            parse_regions([_raw("if", "ghost", 0, 1)], known_fields={"other"})

    def test_a_repeat_must_name_a_known_collection(self):
        parse_regions([_raw("each", "parties", 0, 1)], known_fields=set())
        with pytest.raises(TemplateRegionError):
            parse_regions([_raw("each", "invoices", 0, 1)], known_fields=set())

    def test_nesting_is_allowed_but_straddling_is_not(self):
        parse_regions(
            [_raw("if", "a", 0, 9), _raw("if", "b", 2, 4)], known_fields={"a", "b"}
        )
        with pytest.raises(TemplateRegionError, match="overlaps"):
            parse_regions(
                [_raw("if", "a", 0, 5), _raw("if", "b", 3, 9)], known_fields={"a", "b"}
            )

    def test_a_region_past_the_end_of_the_document_is_rejected(self):
        with pytest.raises(TemplateRegionError, match="past the end"):
            parse_regions(
                [_raw("if", "a", 0, 40)], known_fields={"a"}, paragraph_count=10
            )

    def test_counts_and_nesting_are_bounded(self):
        with pytest.raises(TemplateRegionError):
            parse_regions(
                [_raw("if", "a", 0, 1)] * (MAX_REGIONS + 1), known_fields={"a"}
            )
        deep = [
            _raw("if", "a", index, 400 - index)
            for index in range(MAX_REGION_NESTING + 2)
        ]
        with pytest.raises(TemplateRegionError, match="nested deeper"):
            parse_regions(deep, known_fields={"a"})

    def test_stored_regions_tolerate_a_schema_saved_before_they_existed(self):
        assert stored_regions(None) == []
        assert stored_regions({"fields": []}) == []
        assert stored_regions({"regions": "broken"}) == []
        assert stored_regions({"regions": [_raw("if", "a", 0, 1)]})[0].name == "a"


class TestInjection:
    SCHEMA = {"fields": [{"name": "client", "source_text": "NAME"}, {"name": "entity"}]}

    def test_a_stored_condition_includes_or_drops_the_range(self):
        source = _docx(["NAME", "Authority clause.", "Tail"])
        regions = [TemplateRegion("if", "entity", 1, 1)]

        kept = fill_docx_template(
            source,
            variable_schema=self.SCHEMA,
            variables={"client": "Acme", "entity": "yes"},
            regions=regions,
        )
        assert _text(kept) == ["Acme", "Authority clause.", "Tail"]

        dropped = fill_docx_template(
            source,
            variable_schema=self.SCHEMA,
            variables={"client": "Acme", "entity": ""},
            regions=regions,
        )
        assert _text(dropped) == ["Acme", "Tail"]

    def test_a_stored_repeat_renders_once_per_item_with_item_values(self):
        # A field bound to an item resolves per clone, which is why item
        # bindings are held out of the main replacement pass.
        source = _docx(["Signed:", "PARTY, as ROLE"])
        schema = {
            "fields": [
                {"name": "p_name", "source_text": "PARTY", "binding": "item.party_name"},
                {"name": "p_role", "source_text": "ROLE", "binding": "item.party_role"},
            ]
        }
        output = fill_docx_template(
            source,
            variable_schema=schema,
            variables={},
            collections={
                "parties": [
                    {"party_name": "Acme LLC", "party_role": "plaintiff"},
                    {"party_name": "Bo Li", "party_role": "defendant"},
                ]
            },
            regions=[TemplateRegion("each", "parties", 1, 1)],
        )
        assert _text(output) == [
            "Signed:",
            "Acme LLC, as plaintiff",
            "Bo Li, as defendant",
        ]

    def test_the_retained_source_is_never_modified(self):
        # Retained bytes are the integrity contract every fill re-checks, so a
        # region must exist only in the rendered output.
        source = _docx(["NAME", "Clause"])
        before = bytes(source)
        fill_docx_template(
            source,
            variable_schema=self.SCHEMA,
            variables={"client": "Acme", "entity": ""},
            regions=[TemplateRegion("if", "entity", 1, 1)],
        )
        assert source == before

    def test_stored_and_in_document_regions_resolve_together(self):
        source = _docx(["NAME", "{{#if entity}}", "Typed clause.", "{{/if}}", "Marked clause."])
        output = fill_docx_template(
            source,
            variable_schema=self.SCHEMA,
            variables={"client": "Acme", "entity": ""},
            regions=[TemplateRegion("if", "entity", 4, 4)],
        )
        assert _text(output) == ["Acme"]

    def test_nested_stored_regions_resolve_outermost_first(self):
        source = _docx(["NAME", "Outer", "Inner", "Outer end"])
        schema = {
            "fields": [
                {"name": "client", "source_text": "NAME"},
                {"name": "a"},
                {"name": "b"},
            ]
        }
        regions = [TemplateRegion("if", "a", 1, 3), TemplateRegion("if", "b", 2, 2)]
        both = fill_docx_template(
            source, variable_schema=schema,
            variables={"client": "Acme", "a": "1", "b": "1"}, regions=regions,
        )
        assert _text(both) == ["Acme", "Outer", "Inner", "Outer end"]

        inner_off = fill_docx_template(
            source, variable_schema=schema,
            variables={"client": "Acme", "a": "1", "b": ""}, regions=regions,
        )
        assert _text(inner_off) == ["Acme", "Outer", "Outer end"]

        outer_off = fill_docx_template(
            source, variable_schema=schema,
            variables={"client": "Acme", "a": "", "b": "1"}, regions=regions,
        )
        assert _text(outer_off) == ["Acme"]

    def test_anchored_fields_still_resolve_around_a_region(self):
        # Anchors address ordinals in the original document, so markers must
        # not be injected until replacement has finished.
        source = _docx(["Dropped", "Client: NAME here"])
        schema = {
            "fields": [
                {
                    "name": "client",
                    "source_text": "NAME",
                    "docx_anchor": {"paragraph_ordinal": 1, "start": 8, "end": 12},
                },
                {"name": "entity"},
            ]
        }
        output = fill_docx_template(
            source,
            variable_schema=schema,
            variables={"client": "Acme", "entity": ""},
            regions=[TemplateRegion("if", "entity", 0, 0)],
        )
        assert _text(output) == ["Client: Acme here"]

    def test_a_region_past_the_end_of_the_retained_document_is_rejected(self):
        with pytest.raises(TemplateDocxError, match="past the end"):
            inject_stored_regions(
                _open_docx(_docx(["one", "two"])),
                [TemplateRegion("if", "a", 0, 40)],
            )

    def test_a_region_spanning_two_containers_is_rejected(self):
        # It cannot be removed or repeated as one unit.
        document = Document()
        document.add_paragraph("Body")
        document.add_table(rows=1, cols=1).cell(0, 0).text = "Cell"
        buffer = io.BytesIO()
        document.save(buffer)
        with pytest.raises(TemplateDocxError, match="same part"):
            inject_stored_regions(
                _open_docx(buffer.getvalue()), [TemplateRegion("if", "a", 0, 1)]
            )

    def test_no_regions_leaves_the_document_alone(self):
        document = _open_docx(_docx(["one", "two"]))
        before = len(list(iter_docx_paragraphs(document)))
        assert inject_stored_regions(document, ()) == 0
        # No marker paragraphs added, so every ordinal still addresses the same
        # paragraph it did before.
        assert len(list(iter_docx_paragraphs(document))) == before
        assert [
            p.text for p in iter_docx_paragraphs(document) if p.text.strip()
        ] == ["one", "two"]
