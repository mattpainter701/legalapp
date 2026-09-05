"""Template logic: conditions, markdown blocks, and Word regions."""

import io

import pytest
from docx import Document

from app.services.docx_templates import (
    TemplateDocxError,
    _open_docx,
    apply_docx_logic,
    fill_docx_template,
    iter_docx_paragraphs,
)
from app.services.template_logic import (
    MAX_BLOCK_DEPTH,
    TemplateLogicError,
    expand_markdown_logic,
    field_conditions,
    parse_condition,
    suppressed_fields,
    validate_condition,
)


def _docx(lines, bold_index=None):
    document = Document()
    for index, line in enumerate(lines):
        run = document.add_paragraph().add_run(line)
        if bold_index is not None and index == bold_index:
            run.bold = True
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _text(content):
    return [
        paragraph.text
        for paragraph in iter_docx_paragraphs(_open_docx(content))
        if paragraph.text.strip()
    ]


class TestConditions:
    @pytest.mark.parametrize(
        "operator,value,variables,expected",
        [
            ("present", None, {"x": "a"}, True),
            ("present", None, {"x": "   "}, False),
            ("absent", None, {}, True),
            ("truthy", None, {"x": "Yes"}, True),
            ("truthy", None, {"x": "no"}, False),
            ("falsy", None, {"x": "off"}, True),
            ("equals", "entity", {"x": "Entity"}, True),
            ("not_equals", "entity", {"x": "person"}, True),
            ("in", ["a", "b"], {"x": "B"}, True),
            ("not_in", ["a", "b"], {"x": "c"}, True),
        ],
    )
    def test_operators_evaluate_case_insensitively(
        self, operator, value, variables, expected
    ):
        raw = {"field": "x", "operator": operator}
        if value is not None:
            raw["value"] = value
        assert parse_condition(raw).evaluate(variables) is expected

    def test_missing_value_is_treated_as_empty(self):
        condition = parse_condition({"field": "x", "operator": "equals", "value": "a"})
        assert condition.evaluate({}) is False

    @pytest.mark.parametrize(
        "raw",
        [
            "not-an-object",
            {"field": "x", "operator": "regex", "value": "a"},
            {"field": "", "operator": "present"},
            {"field": "9bad", "operator": "present"},
            {"field": "x", "operator": "present", "value": "unexpected"},
            {"field": "x", "operator": "equals"},
            {"field": "x", "operator": "equals", "value": ["a"]},
            {"field": "x", "operator": "in", "value": "a"},
            {"field": "x", "operator": "in", "value": []},
            {"field": "x", "operator": "equals", "value": "a" * 501},
            {"field": "x", "operator": "in", "value": ["a"] * 51},
        ],
    )
    def test_malformed_conditions_are_rejected(self, raw):
        with pytest.raises(TemplateLogicError):
            parse_condition(raw)

    def test_condition_must_reference_a_declared_field(self):
        raw = {"field": "ghost", "operator": "present"}
        validate_condition(raw, known_fields={"ghost"})
        with pytest.raises(TemplateLogicError):
            validate_condition(raw, known_fields={"other"})

    def test_unparseable_stored_logic_is_skipped_on_read(self):
        schema = {
            "fields": [
                {"name": "good", "logic": {"field": "a", "operator": "present"}},
                {"name": "bad", "logic": {"operator": "nope"}},
                {"name": "none"},
            ]
        }
        assert set(field_conditions(schema)) == {"good"}

    def test_suppressed_fields_track_false_conditions(self):
        schema = {
            "fields": [
                {"name": "spouse", "logic": {"field": "married", "operator": "truthy"}}
            ]
        }
        assert suppressed_fields(schema, {"married": "yes"}) == set()
        assert suppressed_fields(schema, {"married": "no"}) == {"spouse"}


class TestMarkdownLogic:
    def test_if_and_unless_are_complementary(self):
        body = "A{{#if x}}B{{/if}}{{#unless x}}C{{/unless}}D"
        assert expand_markdown_logic(body, {"x": "1"})[0] == "ABD"
        assert expand_markdown_logic(body, {"x": ""})[0] == "ACD"

    def test_each_scopes_item_values_to_their_own_placeholders(self):
        body, extras = expand_markdown_logic(
            "{{#each parties}}[{{name}}]{{/each}}",
            {},
            collections={"parties": [{"name": "X"}, {"name": "Y"}]},
        )
        assert body == "[{{name.__each0}}][{{name.__each1}}]"
        assert extras == {"name.__each0": "X", "name.__each1": "Y"}

    def test_empty_collection_drops_the_region(self):
        assert expand_markdown_logic("a{{#each p}}X{{/each}}b", {})[0] == "ab"

    def test_item_values_are_never_written_into_the_body(self):
        # A value that looks like a marker must stay a value: inlining it would
        # let a customer's data rewrite the template's own structure.
        body, extras = expand_markdown_logic(
            "{{#each p}}{{v}}{{/each}}",
            {},
            collections={"p": [{"v": "{{#if x}}INJECTED{{/if}}"}]},
        )
        assert "INJECTED" not in body
        assert extras["v.__each0"] == "{{#if x}}INJECTED{{/if}}"

    def test_nested_blocks_resolve_inside_out(self):
        body = "{{#if a}}O{{#if b}}I{{/if}}{{/if}}"
        assert expand_markdown_logic(body, {"a": "1", "b": "1"})[0] == "OI"
        assert expand_markdown_logic(body, {"a": "1", "b": ""})[0] == "O"
        assert expand_markdown_logic(body, {"a": "", "b": "1"})[0] == ""

    def test_repeat_item_count_is_bounded(self):
        with pytest.raises(TemplateLogicError):
            expand_markdown_logic(
                "{{#each p}}x{{/each}}",
                {},
                collections={"p": [{"v": "1"}] * 201},
            )

    def test_nesting_depth_is_bounded(self):
        depth = MAX_BLOCK_DEPTH + 2
        body = "{{#if a}}" * depth + "x" + "{{/if}}" * depth
        with pytest.raises(TemplateLogicError):
            expand_markdown_logic(body, {"a": "1"})


class TestDocxLogic:
    SCHEMA = {"fields": [{"name": "client", "source_text": "NAME"}, {"name": "entity"}]}

    def test_conditional_clause_is_included_or_dropped(self):
        source = _docx(
            [
                "NAME",
                "{{#if entity}}",
                "Authority clause.",
                "{{/if}}",
                "{{#unless entity}}",
                "Individual clause.",
                "{{/unless}}",
            ]
        )
        included = fill_docx_template(
            source,
            variable_schema=self.SCHEMA,
            variables={"client": "Acme", "entity": "yes"},
        )
        assert _text(included) == ["Acme", "Authority clause."]
        dropped = fill_docx_template(
            source,
            variable_schema=self.SCHEMA,
            variables={"client": "Acme", "entity": ""},
        )
        assert _text(dropped) == ["Acme", "Individual clause."]

    def test_repeat_emits_one_region_per_item_and_keeps_formatting(self):
        source = _docx(
            ["{{#each parties}}", "{{party_name}} ({{party_role}})", "{{/each}}"],
            bold_index=1,
        )
        output = fill_docx_template(
            source,
            variable_schema={"fields": []},
            variables={},
            collections={
                "parties": [
                    {"party_name": "Acme LLC", "party_role": "plaintiff"},
                    {"party_name": "Bo Li", "party_role": "defendant"},
                ]
            },
        )
        document = _open_docx(output)
        rendered = [p for p in iter_docx_paragraphs(document) if p.text.strip()]
        assert [p.text for p in rendered] == [
            "Acme LLC (plaintiff)",
            "Bo Li (defendant)",
        ]
        # Cloned regions must carry the source run formatting, or a repeated
        # signature block silently loses its styling.
        assert all(run.bold for paragraph in rendered for run in paragraph.runs)

    def test_regions_resolve_inside_a_table_cell(self):
        document = Document()
        table = document.add_table(rows=2, cols=1)
        table.cell(0, 0).paragraphs[0].text = "{{#if entity}}"
        table.cell(0, 0).add_paragraph("Cell clause.")
        table.cell(0, 0).add_paragraph("{{/if}}")
        table.cell(1, 0).paragraphs[0].text = "NAME"
        buffer = io.BytesIO()
        document.save(buffer)
        source = buffer.getvalue()

        kept = fill_docx_template(
            source,
            variable_schema=self.SCHEMA,
            variables={"client": "Acme", "entity": "y"},
        )
        assert "Cell clause." in _text(kept)
        removed = fill_docx_template(
            source,
            variable_schema=self.SCHEMA,
            variables={"client": "Acme", "entity": ""},
        )
        assert "Cell clause." not in _text(removed)

    def test_a_marker_sharing_a_paragraph_with_prose_stays_prose(self):
        # Treating it as a region would delete the sentence around it.
        output = fill_docx_template(
            _docx(["NAME", "Clause {{#if entity}} inline", "tail"]),
            variable_schema=self.SCHEMA,
            variables={"client": "Acme"},
        )
        assert _text(output) == ["Acme", "Clause {{#if entity}} inline", "tail"]

    @pytest.mark.parametrize(
        "lines",
        [
            ["{{#if a}}", "x"],
            ["x", "{{/if}}"],
            ["{{#if a}}", "{{/each}}"],
        ],
    )
    def test_unbalanced_markers_are_rejected(self, lines):
        with pytest.raises(TemplateDocxError):
            apply_docx_logic(_open_docx(_docx(lines)), {"a": "1"})

    def test_repeat_item_count_is_bounded(self):
        document = _open_docx(_docx(["{{#each p}}", "{{v}}", "{{/each}}"]))
        with pytest.raises(TemplateDocxError):
            apply_docx_logic(document, {}, {"p": [{"v": "1"}] * 201})

    def test_anchored_fields_survive_region_changes(self):
        # Anchors address paragraphs by ordinal in the original document, so
        # regions must not be resolved until replacement has finished.
        source = _docx(["{{#if entity}}", "dropped", "{{/if}}", "NAME here"])
        schema = {
            "fields": [
                {
                    "name": "client",
                    "source_text": "NAME",
                    "docx_anchor": {"paragraph_ordinal": 3, "start": 0, "end": 4},
                },
                {"name": "entity"},
            ]
        }
        output = fill_docx_template(
            source, variable_schema=schema, variables={"client": "Acme", "entity": ""}
        )
        assert _text(output) == ["Acme here"]

    def test_a_value_that_only_drives_a_condition_is_not_a_broken_field_map(self):
        output = fill_docx_template(
            _docx(["{{#if entity}}", "clause", "{{/if}}"]),
            variable_schema={"fields": [{"name": "entity"}]},
            variables={"entity": "yes"},
        )
        assert _text(output) == ["clause"]
