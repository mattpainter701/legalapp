import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

from app.services.plugins.manifest import valid_plugin_names


ROOT = Path(__file__).resolve().parents[2]
PACK = ROOT / "demo" / "cybersafeadvisor-corporate-pack"
DISCLAIMER = "SYNTHETIC DEMO - NOT LEGAL ADVICE"


def test_demo_manifest_covers_every_shipped_practice_module_with_rich_sources():
    manifest = json.loads((PACK / "manifest.json").read_text(encoding="utf-8"))

    assert manifest["tenant_domain"] == "cybersafeadvisor.com"
    assert manifest["synthetic"] is True
    assert manifest["schema_version"] == 2
    assert manifest["pack_version"] == "demo-scenario-library-v1"
    matters = manifest["matters"]
    assert set(valid_plugin_names()) <= {
        matter["primary_plugin"] for matter in matters
    }
    assert all(
        {
            "external_key",
            "primary_plugin",
            "matter_type",
            "jurisdiction",
            "name",
            "practice_area",
            "status",
            "client",
            "client_profile",
            "description",
            "documents",
            "demo_prompt",
            "suggested_tasks",
        }
        <= matter.keys()
        for matter in matters
    )
    assert all(len(matter["suggested_tasks"]) >= 3 for matter in matters)
    assert all(
        matter["client_profile"]["address"]
        and matter["client_profile"]["primary_contact"]["email"].endswith(".invalid")
        for matter in matters
    )
    assert all(
        matter["client_profile"]["opposing_party"]["organization"].endswith(" (fictional)")
        and "[" not in matter["client_profile"]["opposing_party"]["organization"]
        and "]" not in matter["client_profile"]["opposing_party"]["organization"]
        for matter in matters
    )
    documents = [name for matter in manifest["matters"] for name in matter["documents"]]
    assert len(documents) >= len(matters) * 3
    assert len(set(documents)) == len(documents)
    assert all((PACK / name).is_file() for name in documents)


def test_demo_documents_are_labeled_structured_and_metadata_scrubbed():
    for path in sorted(PACK.glob("*.docx")):
        document = Document(path)
        visible_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                paragraph.text
                for section in document.sections
                for paragraph in section.header.paragraphs
            ]
        )
        assert DISCLAIMER in visible_text
        assert "Research basis" in visible_text
        assert document.core_properties.author in (None, "")
        assert document.core_properties.last_modified_by in (None, "")

        numbered_headings = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name == "Heading 1"
            and paragraph.text
            not in {"Research basis", "Acknowledged by the synthetic signatories"}
        ]
        assert numbered_headings
        assert all(
            paragraph._p.pPr is not None
            and paragraph._p.pPr.find(qn("w:numPr")) is not None
            for paragraph in numbered_headings
        )

        with ZipFile(path) as package:
            assert "docProps/custom.xml" not in package.namelist()
            for name in package.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    assert b"w:rsid" not in package.read(name)


def test_demo_support_intake_documents_match_manifest_party_links():
    manifest = json.loads((PACK / "manifest.json").read_text(encoding="utf-8"))

    for matter in manifest["matters"]:
        intake = next(
            PACK / filename
            for filename in matter["documents"]
            if filename.startswith("support-") and filename.endswith("-intake-and-contact-profile.docx")
        )
        text = "\n".join(paragraph.text for paragraph in Document(intake).paragraphs)
        opposing = matter["client_profile"]["opposing_party"]
        assert matter["client"] in text
        assert opposing["organization"] in text
        assert opposing["email"] in text
