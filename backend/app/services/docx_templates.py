"""Source-preserving DOCX template rendering."""

from __future__ import annotations

import copy
import io
import hashlib
import json
import re
import zipfile
from bisect import bisect_left, bisect_right
from pathlib import PurePosixPath
from typing import Any, Iterable, Sequence
from urllib.parse import urlparse
from xml.etree import ElementTree

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.services.template_bindings import is_item_binding, item_key


class TemplateDocxError(ValueError):
    """A customer-actionable DOCX template error."""


_VARIABLE_PATTERN = re.compile(r"\{\{\s*([a-zA-Z][a-zA-Z0-9_.-]*)\s*\}\}")
_MAX_DOCX_PACKAGE_FILES = 1_000
_MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
_MAX_DOCX_PART_BYTES = 50 * 1024 * 1024
_MAX_DOCX_COMPRESSION_RATIO = 1_000
_UNSAFE_DOCX_PART_PREFIXES = ("word/activeX/", "word/embeddings/")
_UNSAFE_DOCX_PARTS = {"word/vbaProject.bin"}
_MAX_DOCX_REPEAT_ITEMS = 200


def _docx_xml(package: zipfile.ZipFile, part_name: str) -> ElementTree.Element:
    content = package.read(part_name)
    lowered = content.lower()
    if b"<!doctype" in lowered or b"<!entity" in lowered:
        raise TemplateDocxError("The DOCX contains unsafe XML declarations.")
    try:
        return ElementTree.fromstring(content)
    except ElementTree.ParseError as exc:
        raise TemplateDocxError("The DOCX contains malformed XML.") from exc


def _docx_local_name(tag: str) -> str:
    return str(tag).rsplit("}", 1)[-1]


def _safe_external_docx_hyperlink(target: str) -> bool:
    if len(target) > 2_048 or any(ord(character) < 0x20 for character in target):
        return False
    parsed = urlparse(target)
    scheme = parsed.scheme.casefold()
    return bool(
        (scheme in {"http", "https"} and parsed.netloc)
        or (scheme == "mailto" and parsed.path)
    )


def _validate_docx_relationships_and_altchunks(
    package: zipfile.ZipFile, names: set[str]
) -> None:
    for part_name in names:
        lowered_name = part_name.casefold()
        if lowered_name.endswith(".rels"):
            root = _docx_xml(package, part_name)
            for relation in root.iter():
                if _docx_local_name(relation.tag) != "Relationship":
                    continue
                relation_type = str(relation.attrib.get("Type") or "").casefold()
                target_mode = str(relation.attrib.get("TargetMode") or "").casefold()
                if (
                    "vbaproject" in relation_type
                    or relation_type.endswith("/oleobject")
                    or relation_type.endswith("/package")
                    or relation_type.endswith("/control")
                    or relation_type.endswith("/attachedtemplate")
                    or relation_type.endswith("/afchunk")
                ):
                    raise TemplateDocxError(
                        "The DOCX contains an unsafe active or embedded relationship."
                    )
                if target_mode == "external":
                    target = str(relation.attrib.get("Target") or "")
                    if not relation_type.endswith(
                        "/hyperlink"
                    ) or not _safe_external_docx_hyperlink(target):
                        raise TemplateDocxError(
                            "Only ordinary web and email hyperlinks may be external."
                        )
        if lowered_name.startswith("word/") and lowered_name.endswith(".xml"):
            root = _docx_xml(package, part_name)
            if any(
                _docx_local_name(element.tag).casefold() == "altchunk"
                for element in root.iter()
            ):
                raise TemplateDocxError(
                    "Word documents containing imported altChunk content are not supported."
                )


def docx_source_key(source_text: str, anchor: dict) -> str:
    """Return a stable, opaque identity for one reviewed Word source span."""

    payload = {
        "paragraph_ordinal": int(anchor["paragraph_ordinal"]),
        "start": int(anchor["start"]),
        "end": int(anchor["end"]),
        "source_text": str(source_text),
    }
    encoded = json.dumps(
        payload,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return "docx:" + hashlib.sha256(encoded).hexdigest()[:24]


def validate_docx_package(content: bytes) -> None:
    """Reject spoofed, encrypted, active, or explosively compressed DOCX files."""

    if not content.startswith(b"PK"):
        raise TemplateDocxError("The DOCX is damaged or could not be parsed.")
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as package:
            infos = package.infolist()
            names = {info.filename for info in infos}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise TemplateDocxError("The DOCX is damaged or could not be parsed.")
            if len(infos) > _MAX_DOCX_PACKAGE_FILES:
                raise TemplateDocxError(
                    "This Word file contains too many internal parts to process safely."
                )
            total_uncompressed = 0
            seen_package_names: set[str] = set()
            for info in infos:
                normalized = info.filename.replace("\\", "/")
                normalized_key = normalized.casefold()
                if normalized_key in seen_package_names:
                    raise TemplateDocxError(
                        "This Word file contains duplicate internal parts and cannot be processed safely."
                    )
                seen_package_names.add(normalized_key)
                path = PurePosixPath(normalized)
                if path.is_absolute() or ".." in path.parts:
                    raise TemplateDocxError(
                        "This Word file contains an unsafe internal path."
                    )
                if info.flag_bits & 0x1:
                    raise TemplateDocxError(
                        "Password-protected Word files are not supported. Upload an unlocked clean master."
                    )
                if normalized_key in {
                    part.casefold() for part in _UNSAFE_DOCX_PARTS
                } or normalized_key.startswith(
                    tuple(prefix.casefold() for prefix in _UNSAFE_DOCX_PART_PREFIXES)
                ):
                    raise TemplateDocxError(
                        "Word templates with macros, ActiveX controls, or embedded files are not supported. Remove them and upload a clean master."
                    )
                if info.file_size > _MAX_DOCX_PART_BYTES:
                    raise TemplateDocxError(
                        "This Word file contains an internal part that is too large to process safely."
                    )
                total_uncompressed += int(info.file_size)
                if total_uncompressed > _MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise TemplateDocxError(
                        "This Word file expands beyond the safe processing limit."
                    )
                if (
                    info.file_size > 1_000_000
                    and info.file_size
                    > max(1, int(info.compress_size)) * _MAX_DOCX_COMPRESSION_RATIO
                ):
                    raise TemplateDocxError(
                        "This Word file uses an unsafe compression ratio."
                    )
            _validate_docx_relationships_and_altchunks(package, names)
    except TemplateDocxError:
        raise
    except (OSError, zipfile.BadZipFile) as exc:
        raise TemplateDocxError("The DOCX is damaged or could not be parsed.") from exc


def iter_docx_paragraphs(document: Document) -> Iterable[Any]:
    """Yield body, table, header, and footer paragraphs exactly once.

    The order is part of the source-template contract.  Intake stores an
    ordinal from this iterator for fields whose displayed source text is not
    unique (for example, several underscore blanks in one questionnaire).
    """

    seen: set[Any] = set()

    def emit(paragraphs):
        for paragraph in paragraphs:
            marker = paragraph._p
            if marker not in seen:
                seen.add(marker)
                yield paragraph

    def emit_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from emit(cell.paragraphs)
                    yield from emit_tables(cell.tables)

    def emit_missing_xml(root, parent):
        """Include paragraphs hidden in content controls and text boxes."""

        for paragraph_element in root.iter(qn("w:p")):
            marker = paragraph_element
            if marker in seen:
                continue
            seen.add(marker)
            yield Paragraph(paragraph_element, parent)

    yield from emit(document.paragraphs)
    yield from emit_tables(document.tables)
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from emit(container.paragraphs)
            yield from emit_tables(container.tables)

    # python-docx omits paragraphs wrapped in some structured document tags
    # and DrawingML/VML text boxes from its high-level collections. Append
    # those missing paragraphs after the legacy traversal so existing stored
    # paragraph ordinals remain stable.
    yield from emit_missing_xml(document.element.body, document)
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from emit_missing_xml(container._element, container)


def iter_docx_paragraphs_with_anchors(document: Document) -> Iterable[tuple[int, Any]]:
    """Yield paragraphs with stable source-template ordinals."""

    yield from enumerate(iter_docx_paragraphs(document))


def _compile_replacements(
    replacements: list[tuple[str, str]],
) -> tuple[re.Pattern[str] | None, dict[str, str]]:
    """Compile ordered, literal replacements once for the whole document."""

    replacement_by_source: dict[str, str] = {}
    for source, replacement in replacements:
        if source and source not in replacement_by_source:
            replacement_by_source[source] = replacement
    if not replacement_by_source:
        return None, replacement_by_source
    pattern = re.compile(
        "|".join(re.escape(source) for source in replacement_by_source)
    )
    return pattern, replacement_by_source


def _replace_in_paragraph(
    paragraph: Any,
    pattern: re.Pattern[str] | None,
    replacement_by_source: dict[str, str],
) -> int:
    """Replace original paragraph text, including matches split across Word runs.

    Matches are collected before any mutation, so inserted values can never be
    interpreted as another field's source text.  Run boundaries are indexed
    once and edits are applied right-to-left to keep original offsets stable.
    """

    runs = list(paragraph.runs)
    if not runs or pattern is None:
        return 0
    combined = "".join(run.text for run in runs)
    matches = list(pattern.finditer(combined))
    if not matches:
        return 0

    run_starts: list[int] = []
    run_ends: list[int] = []
    offset = 0
    for run in runs:
        run_starts.append(offset)
        offset += len(run.text)
        run_ends.append(offset)

    applied = 0
    for match in reversed(matches):
        start, end = match.span()
        first = bisect_right(run_ends, start)
        last = bisect_left(run_starts, end) - 1
        if first >= len(runs) or last < first:
            continue
        replacement = replacement_by_source[match.group(0)]
        prefix = runs[first].text[: start - run_starts[first]]
        suffix = runs[last].text[end - run_starts[last] :]
        runs[first].text = prefix + replacement + (suffix if first == last else "")
        for index in range(first + 1, last):
            runs[index].text = ""
        if last != first:
            runs[last].text = suffix
        applied += 1
    return applied


def _replace_at_span(paragraph: Any, start: int, end: int, replacement: str) -> bool:
    """Replace one exact character span while preserving surrounding runs."""

    runs = list(paragraph.runs)
    if not runs or start < 0 or end <= start:
        return False
    combined = "".join(run.text for run in runs)
    if end > len(combined):
        return False

    run_starts: list[int] = []
    run_ends: list[int] = []
    offset = 0
    for run in runs:
        run_starts.append(offset)
        offset += len(run.text)
        run_ends.append(offset)

    first = bisect_right(run_ends, start)
    last = bisect_left(run_starts, end) - 1
    if first >= len(runs) or last < first:
        return False

    prefix = runs[first].text[: start - run_starts[first]]
    suffix = runs[last].text[end - run_starts[last] :]
    runs[first].text = prefix + replacement + (suffix if first == last else "")
    for index in range(first + 1, last):
        runs[index].text = ""
    if last != first:
        runs[last].text = suffix
    return True


_LOGIC_MARKER = re.compile(
    r"^\{\{\s*(?:"
    r"(?P<open>#(?:if|unless|each))\s+(?P<name>[A-Za-z][A-Za-z0-9_.-]*)"
    r"|(?P<close>/(?:if|unless|each))"
    r")\s*\}\}$"
)

#: A region must open and close inside the same container — the document body,
#: or one table cell. Spanning containers would mean deleting half a table.
_MAX_LOGIC_DEPTH = 8


def _marker(paragraph: Any) -> tuple[str, str] | None:
    """Return ``(keyword, name)`` when a paragraph is exactly one logic marker.

    The whole paragraph must be the marker.  A marker sharing a paragraph with
    real prose would force us to delete that prose along with the region, so it
    is treated as ordinary text instead.
    """

    match = _LOGIC_MARKER.match((paragraph.text or "").strip())
    if not match:
        return None
    if match.group("close"):
        return (match.group("close"), "")
    return (match.group("open"), match.group("name"))


def _paragraph_containers(document: Document) -> Iterable[Any]:
    """Yield every element that may directly contain logic regions."""

    yield document.element.body

    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield cell._element
                    yield from walk_tables(cell.tables)

    yield from walk_tables(document.tables)
    for section in document.sections:
        for container in (section.header, section.footer):
            yield container._element
            yield from walk_tables(container.tables)


def _region_bounds(
    children: list[Any], parent: Any
) -> list[tuple[int, int, str, str]]:
    """Return the outermost logic regions in one container.

    Each entry is ``(open index, close index, keyword, name)`` over ``children``.
    Only outermost regions are returned; nested ones are resolved recursively
    after their parent region's fate is decided, so a dropped outer region
    never pays to evaluate what is inside it.
    """

    from docx.text.paragraph import Paragraph

    regions: list[tuple[int, int, str, str]] = []
    stack: list[tuple[int, str, str]] = []
    for index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        found = _marker(Paragraph(child, parent))
        if not found:
            continue
        keyword, name = found
        if keyword.startswith("#"):
            if len(stack) >= _MAX_LOGIC_DEPTH:
                raise TemplateDocxError(
                    f"Word template logic is nested deeper than {_MAX_LOGIC_DEPTH} levels."
                )
            stack.append((index, keyword, name))
            continue
        if not stack:
            raise TemplateDocxError(
                "The Word template closes a logic block that was never opened."
            )
        open_index, open_keyword, open_name = stack.pop()
        expected = "/" + open_keyword[1:]
        if keyword != expected:
            raise TemplateDocxError(
                f"The Word template closes {keyword} where {expected} was expected."
            )
        if not stack:
            regions.append((open_index, index, open_keyword, open_name))
    if stack:
        raise TemplateDocxError(
            "The Word template opens a logic block that is never closed."
        )
    return regions


def _apply_container_logic(
    parent: Any,
    variables: dict[str, str],
    collections: dict[str, Sequence[dict[str, str]]],
    depth: int = 0,
    item_fields: Sequence[tuple[str, str]] = (),
) -> int:
    """Resolve conditional and repeating regions inside one container.

    Returns how many regions were resolved.
    """

    if depth > _MAX_LOGIC_DEPTH:
        raise TemplateDocxError(
            f"Word template logic is nested deeper than {_MAX_LOGIC_DEPTH} levels."
        )
    children = list(parent)
    regions = _region_bounds(children, parent)
    if not regions:
        return 0
    resolved = len(regions)

    # Rewrite right-to-left so earlier indexes stay valid as regions are
    # removed or expanded.
    for open_index, close_index, keyword, name in reversed(regions):
        open_element = children[open_index]
        close_element = children[close_index]
        inner = children[open_index + 1 : close_index]

        if keyword == "#each":
            items = list(collections.get(name) or ())
            if len(items) > _MAX_DOCX_REPEAT_ITEMS:
                raise TemplateDocxError(
                    f"Repeating section {name!r} exceeds the "
                    f"{_MAX_DOCX_REPEAT_ITEMS}-item limit."
                )
            for item in items:
                for element in inner:
                    clone = copy.deepcopy(element)
                    open_element.addprevious(clone)
                    _substitute_item(clone, parent, item, item_fields)
            _drop(parent, [open_element, close_element, *inner])
            continue

        present = bool(str(variables.get(name) or "").strip())
        keep = present if keyword == "#if" else not present
        if keep:
            _drop(parent, [open_element, close_element])
        else:
            _drop(parent, [open_element, close_element, *inner])

    # Nested regions became outermost once their parent resolved.
    if _region_bounds(list(parent), parent):
        resolved += _apply_container_logic(
            parent, variables, collections, depth + 1, item_fields=item_fields
        )
    return resolved


def _drop(parent: Any, elements: Sequence[Any]) -> None:
    for element in elements:
        if element.getparent() is parent:
            parent.remove(element)


def _substitute_item(
    element: Any,
    parent: Any,
    item: dict[str, str],
    item_fields: Sequence[tuple[str, str]] = (),
) -> None:
    """Fill one repeat item's values into a cloned region element.

    Two things are replaced. A ``{{party_name}}`` placeholder, for an author
    who typed one in Word; and the reviewed source text of any field bound to
    that item key, for a field marked in the editor against text the document
    already contained. The second is why item-bound fields are held back from
    the main replacement pass — each clone needs a different value, and by the
    time cloning happens a single shared value would already be in place.
    """

    from docx.text.paragraph import Paragraph

    replacements = [
        (f"{{{{{key}}}}}", "" if value is None else str(value))
        for key, value in item.items()
    ]
    for source_text, key in item_fields:
        if source_text and key in item:
            value = item.get(key)
            replacements.append((source_text, "" if value is None else str(value)))
    pattern, by_source = _compile_replacements(replacements)
    if pattern is None:
        return
    targets = (
        [element] if element.tag == qn("w:p") else list(element.iter(qn("w:p")))
    )
    for paragraph_element in targets:
        _replace_in_paragraph(Paragraph(paragraph_element, parent), pattern, by_source)


def _marker_paragraph(text: str, template: Any) -> Any:
    """Build a paragraph holding exactly one logic marker.

    Cloned from a real paragraph in the document so it inherits a valid
    namespace and section context, then emptied and refilled: constructing one
    from scratch risks a subtly different element than the rest of the body.
    """

    element = copy.deepcopy(template)
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    node.set(qn("xml:space"), "preserve")
    run.append(node)
    element.append(run)
    return element


def inject_stored_regions(document: Document, regions: Sequence[Any]) -> int:
    """Write stored paragraph ranges into the in-memory tree as markers.

    The Studio editor may not rewrite a template's retained bytes — their
    SHA-256 is the integrity contract every fill re-checks, and inserting a
    paragraph would shift every anchor after it. So a region marked in the
    editor is stored as a range of ordinals and materialised here, at render
    time, into the same markers a Word author would have typed. The tested
    marker engine then resolves both kinds identically.

    Every element is resolved by ordinal *before* anything is inserted, and
    insertion runs from the last region backwards, so an earlier region's
    bounds cannot be shifted by a later one's markers.
    """

    if not regions:
        return 0
    paragraphs = [paragraph._p for paragraph in iter_docx_paragraphs(document)]
    ordered = sorted(
        regions,
        key=lambda region: (region.from_ordinal, -region.to_ordinal),
    )
    for region in reversed(ordered):
        if region.to_ordinal >= len(paragraphs):
            raise TemplateDocxError(
                f"A {region.kind} region points past the end of the retained "
                "Word document. Re-upload the source and mark it again."
            )
        first = paragraphs[region.from_ordinal]
        last = paragraphs[region.to_ordinal]
        if first.getparent() is None or last.getparent() is None:
            raise TemplateDocxError(
                f"A {region.kind} region no longer resolves in the retained "
                "Word document. Re-upload the source and mark it again."
            )
        if first.getparent() is not last.getparent():
            # A region that opens in the body and closes in a table cell cannot
            # be removed or repeated as one unit.
            raise TemplateDocxError(
                f"A {region.kind} region must start and end in the same part of "
                "the document."
            )
        first.addprevious(
            _marker_paragraph(f"{{{{#{region.kind} {region.name}}}}}", first)
        )
        last.addnext(_marker_paragraph(f"{{{{/{region.kind}}}}}", last))
    return len(ordered)


def apply_docx_logic(
    document: Document,
    variables: dict[str, str],
    collections: dict[str, Sequence[dict[str, str]]] | None = None,
    item_fields: Sequence[tuple[str, str]] = (),
) -> int:
    """Resolve ``{{#if}}``, ``{{#unless}}``, and ``{{#each}}`` regions in place.

    Markers are authored in Word itself, each alone in its own paragraph, so a
    firm can express "include this clause only for entity clients" without
    leaving the editor they already draft in.

    This runs *after* field replacement: anchored fields address paragraphs by
    the ordinal `iter_docx_paragraphs` assigns to the original document, and
    deleting or cloning paragraphs first would invalidate every ordinal after
    the region.
    """

    resolved = 0
    for container in list(_paragraph_containers(document)):
        resolved += _apply_container_logic(
            container, variables, collections or {}, item_fields=item_fields
        )
    return resolved


def _open_docx(content: bytes) -> Document:
    validate_docx_package(content)
    try:
        return Document(io.BytesIO(content))
    except Exception as exc:
        raise TemplateDocxError("The DOCX is damaged or could not be parsed.") from exc


def fill_docx_template(
    content: bytes,
    *,
    variable_schema: dict | None,
    variables: dict[str, str],
    enforce_required: bool = False,
    collections: dict[str, Sequence[dict[str, str]]] | None = None,
    regions: Sequence[Any] | None = None,
) -> bytes:
    """Fill a retained DOCX without converting it to plain text."""

    document = _open_docx(content)
    fields = (variable_schema or {}).get("fields") or []
    by_name: dict[str, dict[str, Any]] = {}
    for field in fields:
        if not isinstance(field, dict):
            raise TemplateDocxError("The stored Word field mapping is invalid.")
        name = str(field.get("name") or "").strip()
        if not name or name in by_name:
            raise TemplateDocxError(
                "The stored Word field mapping contains duplicates."
            )
        by_name[name] = field

    unknown = set(variables) - set(by_name)
    if unknown:
        raise TemplateDocxError(
            "Unknown Word template variable(s): " + ", ".join(sorted(unknown)[:5])
        )
    if enforce_required:
        missing = sorted(
            name
            for name, field in by_name.items()
            if field.get("required") and not str(variables.get(name) or "").strip()
        )
        if missing:
            raise TemplateDocxError(
                "Required Word field(s) are empty: " + ", ".join(missing)
            )

    replacements: list[tuple[str, str]] = []
    anchored_replacements: dict[int, list[tuple[int, int, str, str, str]]] = {}
    # A field bound to a repeat item resolves once per clone, not once for the
    # document, so it is excluded here and applied during cloning instead.
    item_fields: list[tuple[str, str]] = []
    for name, field in by_name.items():
        binding = field.get("binding")
        if is_item_binding(str(binding or "")):
            item_fields.append(
                (
                    str(field.get("source_text") or field.get("example") or ""),
                    item_key(str(binding)),
                )
            )
            continue
        value = str(variables.get(name) or "")
        if len(value) > 10_000:
            raise TemplateDocxError(
                f"Value for Word field {name!r} exceeds the 10,000-character limit."
            )
        anchor = field.get("docx_anchor")
        if isinstance(anchor, dict):
            try:
                ordinal = int(anchor["paragraph_ordinal"])
                start = int(anchor["start"])
                end = int(anchor["end"])
            except (KeyError, TypeError, ValueError) as exc:
                raise TemplateDocxError(
                    f"The stored Word location for {name!r} is invalid. Re-upload and review the template."
                ) from exc
            source_text = str(field.get("source_text") or field.get("example") or "")
            if not source_text or end - start != len(source_text):
                raise TemplateDocxError(
                    f"The stored Word location for {name!r} no longer matches its source text. Re-upload and review the template."
                )
            anchored_replacements.setdefault(ordinal, []).append(
                (start, end, source_text, value, name)
            )
            continue

        replacements.append((f"{{{{{name}}}}}", value))
        source_text = str(field.get("source_text") or field.get("example") or "")
        if source_text and source_text != f"{{{{{name}}}}}":
            replacements.append((source_text, value))

    replacement_count = 0
    replacement_pattern, replacement_by_source = _compile_replacements(replacements)
    anchored_names: set[str] = set()
    for ordinal, paragraph in iter_docx_paragraphs_with_anchors(document):
        for start, end, source_text, value, name in sorted(
            anchored_replacements.get(ordinal, []), reverse=True
        ):
            combined = "".join(run.text for run in paragraph.runs)
            if combined[start:end] != source_text:
                raise TemplateDocxError(
                    f"The retained Word document no longer matches the reviewed location for {name!r}. Re-upload and review the template."
                )
            if _replace_at_span(paragraph, start, end, value):
                replacement_count += 1
                anchored_names.add(name)
        replacement_count += _replace_in_paragraph(
            paragraph,
            replacement_pattern,
            replacement_by_source,
        )

    missing_anchors = {
        name for entries in anchored_replacements.values() for *_, name in entries
    } - anchored_names
    if missing_anchors:
        raise TemplateDocxError(
            "The retained Word document no longer matches the reviewed field location(s): "
            + ", ".join(sorted(missing_anchors)[:5])
        )

    # Regions resolve last: field anchors address paragraphs by their ordinal
    # in the original document, so nothing may be added or removed until every
    # anchored replacement has been applied. Stored ranges become markers here,
    # for the same reason and at the same moment.
    injected = inject_stored_regions(document, regions or ())
    resolved_regions = (
        apply_docx_logic(document, variables, collections, item_fields=item_fields)
        + injected
    )

    # A value that only drove a conditional clause is still a value that did
    # something, so a logic-only template is not a broken field map.
    if any(variables.values()) and replacement_count == 0 and resolved_regions == 0:
        raise TemplateDocxError(
            "The retained Word document no longer matches its field map. Re-upload the source and review the detected fields."
        )

    output = io.BytesIO()
    try:
        document.save(output)
        _open_docx(output.getvalue())
    except TemplateDocxError:
        raise
    except Exception as exc:
        raise TemplateDocxError(
            "The generated Word document could not be finalized."
        ) from exc
    return output.getvalue()


def docx_placeholder_names(content: bytes) -> list[str]:
    """Return explicit placeholders in first-seen document order."""

    document = _open_docx(content)
    found: list[str] = []
    seen: set[str] = set()
    for paragraph in iter_docx_paragraphs(document):
        for match in _VARIABLE_PATTERN.finditer(paragraph.text or ""):
            name = match.group(1)
            if name not in seen:
                seen.add(name)
                found.append(name)
    return found
