"""Deterministic execution of an approved task's ``pending_action``.

This is the half of the chat action layer with no model in it. The assistant may
draft a client email, but a human approves it on the work board and *this* code
sends it — so an outbound message never depends on model behavior at send time.

One automatic attempt per approval is the database guarantee. A
``task_automation_runs`` row claims the work, and the unique constraint on
``(task_id, idempotency_key)`` prevents a double-click, replay, or concurrent
worker from starting another application-level attempt. Provider acceptance
followed by a timeout is inherently ambiguous, so a failed or ``sending`` run
stays as evidence and is not automatically retried. An attorney must edit and
re-approve a changed action payload to create a new key.
"""

from __future__ import annotations

import asyncio
import hashlib
import json
import logging
import os
from io import BytesIO
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Awaitable, Callable

from pydantic import ValidationError
from sqlalchemy import select
from sqlalchemy.dialects.postgresql import insert as pg_insert
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import async_session_maker, set_tenant_context
from app.models.task import Task, TaskAutomationRun
from app.models.matter_document import MatterDocument
from app.models.plugin import Matter
from app.models.tenant import TenantSettings
from app.models.contact import Contact
from app.models.document import Document
from app.models.matter_party import MatterParty
from app.schemas.chat_action import (
    EmailClientAction,
    MatterDocumentDraftAction,
    normalize_single_mailbox,
)
from app.services.matter_file_store import MatterFileStore
from app.services.connected_mail import (
    DELIVERY_CONFIRMED_SENT,
    DELIVERY_NOT_ATTEMPTED,
    DELIVERY_OUTCOME_UNKNOWN,
    send_client_email,
)
from app.services.email import email_service
from app.services.task_workflow import append_task_event

logger = logging.getLogger(__name__)
matter_file_store = MatterFileStore()

# Approval is a specific move, not merely leaving Review. Cancelling a drafted
# client email must never send it — that inverts the attorney's intent — and
# parking it in Waiting or closing the task without acting are not approvals
# either. Only accepting the draft into active work executes it.
APPROVAL_FROM_STATUS = "review"
APPROVAL_TO_STATUSES = frozenset({"in_progress"})

TASK_AUTOMATION_JOB = "task_automation"

# Delivery states an attorney may see. Only "sent" means the client was contacted.
DELIVERY_STATUSES = ("queued", "sending", "sent", "failed")
TERMINAL_DELIVERY_STATUSES = ("sent", "failed")


class ActionApprovalConflict(RuntimeError):
    """The proposed approval could create an unsafe duplicate delivery."""

    def __init__(self, detail: str):
        super().__init__(detail)
        self.detail = detail


def transition_is_approval(from_status: str | None, to_status: str | None) -> bool:
    """Whether this status change means "execute the drafted action".

    The single definition both status-changing endpoints consult. Keeping it here
    rather than in a router is the point: the rule was previously expressed at
    the call site, where the transition endpoint sent on cancellation and the
    generic PATCH approved without ever executing.
    """
    return from_status == APPROVAL_FROM_STATUS and to_status in APPROVAL_TO_STATUSES


def _canonical_action_snapshot(payload: dict[str, Any] | None) -> dict[str, Any]:
    """Deep-copy an action into the canonical JSON representation we audit."""
    encoded = json.dumps(
        payload or {},
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=str,
    )
    return json.loads(encoded)


def action_payload_sha256(payload: dict[str, Any] | None) -> str:
    snapshot = _canonical_action_snapshot(payload)
    encoded = json.dumps(
        snapshot,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def automation_idempotency_key(task: Task, from_status: str) -> str:
    """Identify one approval event.

    Keyed on the exact action payload rather than the general task version. A
    status round-trip or unrelated task edit therefore cannot resend an
    unchanged email. A terminal no-send/acknowledged-unknown retry gets a new
    approval key derived from the previous immutable run id.
    """
    return f"approve:{from_status}:sha256:{action_payload_sha256(task.pending_action)}"


async def enqueue_automation_run(
    db: AsyncSession,
    task: Task,
    *,
    from_status: str,
    actor_user_id=None,
    idempotency_key: str | None = None,
) -> TaskAutomationRun | None:
    """Record the intent to send, inside the caller's transaction.

    This is what makes delivery durable. The queued row commits atomically with
    the status change, so a process that dies between approval and send leaves
    behind a record the worker will pick up — previously that send was simply
    lost, with the attorney told it was approved.

    Does not commit: the caller owns the transaction, which is the whole point.
    """
    action_type = str((task.pending_action or {}).get("type") or "")
    if not action_type:
        return None
    action_snapshot = _canonical_action_snapshot(task.pending_action)
    action_sha256 = action_payload_sha256(action_snapshot)
    approval_key = idempotency_key or automation_idempotency_key(task, from_status)
    stmt = (
        pg_insert(TaskAutomationRun)
        .values(
            tenant_id=task.tenant_id,
            task_id=task.id,
            action_type=action_type,
            idempotency_key=approval_key,
            action_snapshot=action_snapshot,
            action_sha256=action_sha256,
            status="queued",
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
            triggered_by_user_id=actor_user_id,
        )
        # An existing row already covers this approval — including one still
        # sending or already sent. Never reset it here.
        .on_conflict_do_nothing(constraint="uq_task_automation_runs_task_key")
        .returning(TaskAutomationRun.id)
    )
    await db.execute(stmt)
    return None


async def _claim_run(
    db: AsyncSession,
    task: Task,
    *,
    action_type: str,
    idempotency_key: str,
    actor_user_id,
) -> TaskAutomationRun | None:
    """Atomically move a queued run into sending, or return None.

    A conditional UPDATE, so concurrency is resolved by Postgres rather than by
    read-then-write in application code. ``sending`` is excluded so a concurrent
    attempt cannot start a second send, and ``sent`` is excluded so nothing can
    start another automatic attempt. Losing this race is expected, not an error.

    Upserts a row when none exists so a direct call still works without a prior
    enqueue; the unique constraint keeps that safe under concurrency.
    """
    action_snapshot = _canonical_action_snapshot(task.pending_action)
    claim = (
        pg_insert(TaskAutomationRun)
        .values(
            tenant_id=task.tenant_id,
            task_id=task.id,
            action_type=action_type,
            idempotency_key=idempotency_key,
            action_snapshot=action_snapshot,
            action_sha256=action_payload_sha256(action_snapshot),
            status="sending",
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
            triggered_by_user_id=actor_user_id,
        )
        .on_conflict_do_update(
            constraint="uq_task_automation_runs_task_key",
            set_={
                "status": "sending",
                "error_message": None,
                "triggered_by_user_id": actor_user_id,
                "completed_at": None,
                "delivery_certainty": DELIVERY_NOT_ATTEMPTED,
            },
            # Failed delivery is terminal until an attorney explicitly edits and
            # re-approves a changed payload. Automatically retrying an ambiguous
            # provider failure can duplicate a message that was actually sent.
            where=TaskAutomationRun.status == "queued",
        )
        .returning(TaskAutomationRun.id)
    )
    run_id = await db.scalar(claim)
    if run_id is None:
        return None
    await db.commit()
    # set_tenant_context uses SET LOCAL. The claim commit clears it, and these
    # tables are FORCE-RLS in production, so rebind before the follow-up SELECT.
    await set_tenant_context(db, str(task.tenant_id))
    return await db.scalar(
        select(TaskAutomationRun).where(TaskAutomationRun.id == run_id)
    )


async def _record_terminal_no_send(
    db: AsyncSession,
    task: Task,
    *,
    action_type: str,
    idempotency_key: str,
    actor_user_id,
    detail: str,
) -> None:
    """Turn a queued approval into visible, terminal evidence of no delivery."""
    now = datetime.now(timezone.utc)
    action_snapshot = _canonical_action_snapshot(task.pending_action)
    run_id = await db.scalar(
        pg_insert(TaskAutomationRun)
        .values(
            tenant_id=task.tenant_id,
            task_id=task.id,
            action_type=action_type or "unknown",
            idempotency_key=idempotency_key,
            action_snapshot=action_snapshot,
            action_sha256=action_payload_sha256(action_snapshot),
            status="failed",
            error_message=detail[:500],
            delivery_detail=detail[:500],
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
            triggered_by_user_id=actor_user_id,
            completed_at=now,
        )
        .on_conflict_do_update(
            constraint="uq_task_automation_runs_task_key",
            set_={
                "status": "failed",
                "error_message": detail[:500],
                "delivery_detail": detail[:500],
                "delivery_certainty": DELIVERY_NOT_ATTEMPTED,
                "completed_at": now,
            },
            where=TaskAutomationRun.status == "queued",
        )
        .returning(TaskAutomationRun.id)
    )
    if run_id is not None:
        append_task_event(
            db,
            task,
            event_type="automation_blocked",
            actor_user_id=actor_user_id,
            note=detail[:500],
            metadata={"action_type": action_type or "unknown", "sent": False},
        )
    await db.commit()
    logger.info(
        "task_automation_not_sent task_id=%s detail=%s",
        task.id,
        detail,
    )


# ── Action handlers ─────────────────────────────────────────────────────────


@dataclass(frozen=True, slots=True)
class ActionExecutionResult:
    succeeded: bool
    detail: str
    provider: str | None = None
    provider_message_id: str | None = None
    delivery_certainty: str = DELIVERY_OUTCOME_UNKNOWN


async def _run_email_client(
    db: AsyncSession,
    task: Task,
    payload: dict[str, Any],
    actor_user_id,
) -> ActionExecutionResult:
    """Send the approved client email and retain provider audit metadata."""
    try:
        action = EmailClientAction.model_validate(payload)
    except ValidationError:
        # The payload was written by this system, so an invalid one means a bug
        # or tampering. Never improvise a partial send.
        return ActionExecutionResult(
            False,
            "The stored email draft is not a valid action payload",
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
        )

    if not await _recipient_bindings_are_current(db, task, action):
        return ActionExecutionResult(
            False,
            "Not sent: a recipient is no longer the approved party/address on this matter",
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
        )
    if not await _action_sources_are_current(db, task, action):
        return ActionExecutionResult(
            False,
            "Not sent: one or more cited local documents are no longer available",
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
        )

    # OAuth refresh may commit its session. Keep it separate from the worker's
    # task-locking transaction so approval cannot be cancelled between the
    # final status check and the external send.
    tenant_id = task.tenant_id
    async with async_session_maker() as delivery_db:
        await set_tenant_context(delivery_db, str(tenant_id))
        delivery = await send_client_email(
            delivery_db,
            tenant_id=tenant_id,
            actor_user_id=actor_user_id,
            to=list(action.to),
            subject=action.subject,
            html_body=_body_to_html(action.body),
            text_body=action.body,
            smtp_service=email_service,
        )
    return ActionExecutionResult(
        bool(delivery.result),
        delivery.detail,
        provider=delivery.provider,
        provider_message_id=delivery.provider_message_id,
        delivery_certainty=(
            delivery.delivery_certainty
            or (
                DELIVERY_CONFIRMED_SENT if delivery.result else DELIVERY_OUTCOME_UNKNOWN
            )
        ),
    )


def _word_document_bytes(title: str, body: str) -> bytes:
    """Render a straightforward, editable OOXML document from reviewed text."""
    from docx import Document

    document = Document()
    document.add_heading(title, level=0)
    for paragraph in body.split("\n\n"):
        document.add_paragraph(paragraph.strip())
    output = BytesIO()
    document.save(output)
    return output.getvalue()


async def _run_matter_document_draft(
    db: AsyncSession,
    task: Task,
    payload: dict[str, Any],
    actor_user_id,
) -> ActionExecutionResult:
    try:
        action = MatterDocumentDraftAction.model_validate(payload)
    except ValidationError:
        return ActionExecutionResult(
            False,
            "The stored document draft is not a valid action payload",
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
        )
    if task.matter_id != action.matter_id:
        return ActionExecutionResult(
            False,
            "The reviewed document no longer belongs to this matter",
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
        )
    matter = await db.scalar(
        select(Matter).where(Matter.id == action.matter_id, Matter.tenant_id == task.tenant_id)
    )
    if matter is None:
        return ActionExecutionResult(
            False,
            "The matter for this reviewed document is no longer available",
            delivery_certainty=DELIVERY_NOT_ATTEMPTED,
        )

    filename = f"{action.title}.docx"
    content = await asyncio.to_thread(_word_document_bytes, action.title, action.body)
    try:
        storage = await matter_file_store.store_matter_file_result(
            db=db,
            tenant_id=str(task.tenant_id),
            matter_slug=matter.slug,
            category="general",
            filename=filename,
            content=content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            matter_cloud_folder=matter.cloud_folder,
        )
    except Exception as exc:
        return ActionExecutionResult(False, f"Word document could not be saved: {exc}"[:500], delivery_certainty=DELIVERY_NOT_ATTEMPTED)
    if not storage.succeeded:
        return ActionExecutionResult(False, f"Word document could not be saved: {storage.error or 'storage failed'}"[:500], delivery_certainty=DELIVERY_NOT_ATTEMPTED)

    document = MatterDocument(
        tenant_id=task.tenant_id,
        matter_id=action.matter_id,
        uploaded_by_user_id=actor_user_id,
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_size=len(content),
        storage_path=storage.storage_path,
        storage_provider=storage.provider,
        storage_backend=storage.backend,
        provider_object_id=storage.provider_item_id,
        provider_drive_id=storage.drive_id,
        provider_parent_id=storage.parent_id,
        storage_error=storage.error,
        description="Assistant draft approved by attorney",
        document_category="general",
    )
    db.add(document)
    await db.flush()
    return ActionExecutionResult(
        True,
        "Approved Word document saved to Matter Documents.",
        provider="matter_document",
        provider_message_id=str(document.id),
        delivery_certainty=DELIVERY_CONFIRMED_SENT,
    )


async def _recipient_bindings_are_current(
    db: AsyncSession,
    task: Task,
    action: EmailClientAction,
) -> bool:
    """Revalidate and lock party/contact bindings immediately before send."""
    if task.matter_id != action.matter_id:
        return False
    requested_ids = [binding.party_id for binding in action.recipient_bindings]
    rows = (
        await db.execute(
            select(MatterParty.id, MatterParty.contact_id, Contact.email)
            .join(Contact, MatterParty.contact_id == Contact.id)
            .where(
                MatterParty.id.in_(requested_ids),
                MatterParty.tenant_id == task.tenant_id,
                MatterParty.matter_id == action.matter_id,
                Contact.tenant_id == task.tenant_id,
            )
            .with_for_update()
        )
    ).all()
    current = {party_id: (contact_id, email) for party_id, contact_id, email in rows}
    if len(current) != len(set(requested_ids)):
        return False
    for binding in action.recipient_bindings:
        row = current.get(binding.party_id)
        if row is None or row[0] != binding.contact_id:
            return False
        try:
            current_address = normalize_single_mailbox(row[1])
        except ValueError:
            return False
        if current_address.casefold() != binding.address.casefold():
            return False
    return True


async def _action_sources_are_current(
    db: AsyncSession,
    task: Task,
    action: EmailClientAction,
) -> bool:
    """Lock and verify every local source before approval/delivery."""
    if not action.source_document_ids:
        return True
    documents = (
        (
            await db.execute(
                select(Document)
                .where(
                    Document.id.in_(action.source_document_ids),
                    Document.tenant_id == task.tenant_id,
                )
                .with_for_update()
            )
        )
        .scalars()
        .all()
    )
    if len(documents) != len(set(action.source_document_ids)):
        return False
    bindings = {
        binding.document_id: binding.sha256
        for binding in action.source_document_bindings
    }
    for document in documents:
        if document.matter_id is not None and document.matter_id != action.matter_id:
            return False
        storage_path = str(document.storage_path or "").strip()
        expected_hash = bindings.get(document.id)
        if (
            not expected_hash
            or storage_path.startswith(("http://", "https://"))
            or not os.path.isfile(storage_path)
        ):
            return False
        actual_hash = await asyncio.to_thread(_file_sha256, storage_path)
        if actual_hash != expected_hash:
            return False
    return True


def _file_sha256(path: str) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _body_to_html(body: str) -> str:
    from html import escape

    paragraphs = [
        f"<p>{escape(block).replace(chr(10), '<br>')}</p>"
        for block in body.split("\n\n")
        if block.strip()
    ]
    return "".join(paragraphs) or f"<p>{escape(body)}</p>"


ACTION_HANDLERS: dict[
    str,
    Callable[[AsyncSession, Task, dict, Any], Awaitable[ActionExecutionResult]],
] = {
    "email_client": _run_email_client,
    "matter_document_draft": _run_matter_document_draft,
}


# ── Entry point ─────────────────────────────────────────────────────────────


async def run_task_automation(
    task_id,
    tenant_id,
    *,
    from_status: str,
    to_status: str,
    actor_user_id=None,
    approval_idempotency_key: str | None = None,
) -> None:
    """Execute an approved task's pending action, at most once.

    Runs after the transition has already been committed, on its own session,
    because the request session is gone by then. A failure here records evidence
    and never rolls the task back: the attorney's approval is a real decision
    that happened, and hiding it would be worse than a visible failed run.

    Re-checks the approval rule itself rather than trusting the caller, so a new
    call site cannot reintroduce a send-on-cancel.
    """
    if not transition_is_approval(from_status, to_status):
        return

    async with async_session_maker() as db:
        try:
            await set_tenant_context(db, str(tenant_id))
            task = await db.scalar(
                select(Task)
                .where(
                    Task.id == task_id,
                    Task.tenant_id == tenant_id,
                )
                .with_for_update()
            )
            if task is None:
                return

            approval_key = approval_idempotency_key or automation_idempotency_key(
                task, from_status
            )
            action_type = str((task.pending_action or {}).get("type") or "")

            # A reviewed task without a drafted side effect is ordinary board
            # work, not a failed automation.
            if not task.pending_action:
                return

            if task.status != to_status or to_status != "in_progress":
                await _record_terminal_no_send(
                    db,
                    task,
                    action_type=action_type,
                    idempotency_key=approval_key,
                    actor_user_id=actor_user_id,
                    detail=(
                        "Not sent: the approval was superseded before delivery "
                        f"(task is now {task.status})."
                    ),
                )
                return

            actions_enabled = await db.scalar(
                select(TenantSettings.enable_chat_actions).where(
                    TenantSettings.tenant_id == task.tenant_id
                )
            )
            if actions_enabled is not True:
                await _record_terminal_no_send(
                    db,
                    task,
                    action_type=action_type,
                    idempotency_key=approval_key,
                    actor_user_id=actor_user_id,
                    detail="Not sent: chat actions were disabled before delivery.",
                )
                return

            handler = ACTION_HANDLERS.get(action_type)
            if handler is None:
                # Fail closed on an unknown action rather than guessing.
                logger.warning(
                    "task_automation_unknown_action task_id=%s action_type=%r",
                    task_id,
                    action_type,
                )
                await _record_terminal_no_send(
                    db,
                    task,
                    action_type=action_type,
                    idempotency_key=approval_key,
                    actor_user_id=actor_user_id,
                    detail=f"Not sent: unsupported action type {action_type!r}.",
                )
                return

            run = await _claim_run(
                db,
                task,
                action_type=action_type,
                idempotency_key=approval_key,
                actor_user_id=actor_user_id,
            )
            if run is None:
                # Already claimed or terminal. This is the duplicate-attempt
                # path and is expected under double-click, not an error.
                logger.info(
                    "task_automation_already_claimed task_id=%s action_type=%s",
                    task_id,
                    action_type,
                )
                return

            # The claim commit released the task lock. Reacquire it and validate
            # the current state immediately before the external side effect. If
            # cancel/review won the race, preserve a terminal no-send result.
            task = await db.scalar(
                select(Task)
                .where(Task.id == task_id, Task.tenant_id == tenant_id)
                .with_for_update()
            )
            if (
                task is None
                or task.status != to_status
                or to_status != "in_progress"
                or not task.pending_action
                or (
                    run.action_sha256 is not None
                    and action_payload_sha256(task.pending_action) != run.action_sha256
                )
            ):
                detail = "Not sent: the approval was superseded before delivery."
                run.status = "failed"
                run.error_message = detail
                run.delivery_detail = detail
                run.delivery_certainty = DELIVERY_NOT_ATTEMPTED
                run.completed_at = datetime.now(timezone.utc)
                if task is not None:
                    append_task_event(
                        db,
                        task,
                        event_type="automation_blocked",
                        actor_user_id=actor_user_id,
                        note=detail,
                        metadata={"action_type": action_type, "sent": False},
                    )
                await db.commit()
                return

            # `_claim_run` commits, so the earlier feature-flag read is no
            # longer authoritative. Lock and recheck the tenant setting at the
            # irreversible side-effect boundary; an admin disable either wins
            # before this lock (no send) or waits until this attempt finishes.
            tenant_settings = await db.scalar(
                select(TenantSettings)
                .where(TenantSettings.tenant_id == task.tenant_id)
                .with_for_update()
            )
            if (
                tenant_settings is None
                or tenant_settings.enable_chat_actions is not True
            ):
                detail = "Not sent: chat actions were disabled before delivery."
                run.status = "failed"
                run.error_message = detail
                run.delivery_detail = detail
                run.delivery_certainty = DELIVERY_NOT_ATTEMPTED
                run.completed_at = datetime.now(timezone.utc)
                append_task_event(
                    db,
                    task,
                    event_type="automation_blocked",
                    actor_user_id=actor_user_id,
                    note=detail,
                    metadata={"action_type": action_type, "sent": False},
                )
                await db.commit()
                return

            locked_tenant_id = task.tenant_id
            # Execute the immutable payload that was claimed at approval time,
            # never a mutable task field that may have changed after enqueue.
            payload_snapshot = _canonical_action_snapshot(
                run.action_snapshot or task.pending_action
            )
            if run.action_snapshot is None:
                # Compatibility for a queued row created before migration 103.
                # All new runs already carry these fields before delivery.
                run.action_snapshot = payload_snapshot
                run.action_sha256 = action_payload_sha256(payload_snapshot)
            try:
                result = await handler(db, task, payload_snapshot, actor_user_id)
            except Exception as exc:
                result = ActionExecutionResult(
                    False, f"{type(exc).__name__}: {exc}"[:500]
                )
                logger.warning(
                    "task_automation_handler_raised task_id=%s action_type=%s",
                    task_id,
                    action_type,
                    exc_info=True,
                )

            # OAuth token refresh may commit inside the delivery handler, which
            # clears SET LOCAL. Rebind before writing the durable outcome.
            await set_tenant_context(db, str(locked_tenant_id))
            run.status = "sent" if result.succeeded else "failed"
            run.error_message = None if result.succeeded else result.detail[:500]
            run.delivery_detail = result.detail[:500]
            run.delivery_certainty = result.delivery_certainty
            run.provider = result.provider[:50] if result.provider else None
            run.provider_message_id = (
                result.provider_message_id[:500] if result.provider_message_id else None
            )
            run.completed_at = datetime.now(timezone.utc)
            audit_snapshot = run.action_snapshot or {}
            append_task_event(
                db,
                task,
                event_type=(
                    "automation_succeeded" if result.succeeded else "automation_failed"
                ),
                actor_user_id=actor_user_id,
                note=None if result.succeeded else result.detail[:500],
                metadata={
                    "action_type": action_type,
                    "detail": result.detail[:200],
                    "action_sha256": run.action_sha256,
                    "provider": result.provider,
                    "provider_message_id": result.provider_message_id,
                    "to": list(audit_snapshot.get("to") or [])[:10],
                    "subject": audit_snapshot.get("subject"),
                    "source_ids": list(audit_snapshot.get("source_ids") or [])[:10],
                },
            )
            if result.succeeded:
                # Clear the draft so a later manual transition of the same task
                # cannot re-enter the automation path at all.
                task.pending_action = None
            await db.commit()
        except Exception:
            # Infrastructure failures must escape to the durable-job worker so
            # its lease/retry machinery can run. Handler/provider failures are
            # converted to a terminal run above; swallowing a database failure
            # here would mark the durable job successful and strand the action.
            logger.warning(
                "task_automation_failed task_id=%s tenant_id=%s",
                task_id,
                tenant_id,
                exc_info=True,
            )
            await db.rollback()
            raise


async def enqueue_durable_automation(
    db: AsyncSession,
    task: Task,
    *,
    from_status: str | None,
    to_status: str | None,
    actor_user_id=None,
    acknowledge_prior_delivery_risk: bool = False,
) -> str | None:
    """Persist the send intent in the caller's transaction. Must precede commit.

    Writes both the queued run (the attorney-visible delivery state) and a
    durable job (the worker's instruction). Both are idempotent and neither
    commits, so they land atomically with the approval or not at all — there is
    no window in which a task is approved but the send is unrecorded.
    """
    if not transition_is_approval(from_status, to_status):
        return None
    if not task.pending_action:
        return None

    if str(task.pending_action.get("type") or "") == "email_client":
        try:
            pending_email = EmailClientAction.model_validate(task.pending_action)
        except ValidationError as exc:
            raise ActionApprovalConflict(
                "The stored email draft has invalid or missing recipient bindings. "
                "Create a new draft before approval."
            ) from exc
        if not await _recipient_bindings_are_current(db, task, pending_email):
            raise ActionApprovalConflict(
                "A recipient's matter-party membership or email address changed "
                "after this draft was prepared. Create and review a new draft."
            )
        if not await _action_sources_are_current(db, task, pending_email):
            raise ActionApprovalConflict(
                "A cited local document is no longer available or no longer "
                "belongs to this matter. Restore the evidence or create a new draft."
            )

    active_run = await db.scalar(
        select(TaskAutomationRun)
        .where(
            TaskAutomationRun.tenant_id == task.tenant_id,
            TaskAutomationRun.task_id == task.id,
            TaskAutomationRun.status.in_(("queued", "sending")),
        )
        .order_by(TaskAutomationRun.created_at.desc(), TaskAutomationRun.id.desc())
        .limit(1)
    )
    if active_run is not None:
        raise ActionApprovalConflict(
            "An earlier email delivery is still queued or in progress. Wait for "
            "its outcome before approving another draft."
        )

    latest_run = await db.scalar(
        select(TaskAutomationRun)
        .where(
            TaskAutomationRun.tenant_id == task.tenant_id,
            TaskAutomationRun.task_id == task.id,
        )
        .order_by(TaskAutomationRun.created_at.desc(), TaskAutomationRun.id.desc())
        .limit(1)
    )
    retry_after_run_id = None
    if latest_run is not None and latest_run.status == "failed":
        certainty = latest_run.delivery_certainty or DELIVERY_OUTCOME_UNKNOWN
        if (
            certainty == DELIVERY_OUTCOME_UNKNOWN
            and not acknowledge_prior_delivery_risk
        ):
            raise ActionApprovalConflict(
                "A prior delivery was not confirmed. Check the connected "
                "mailbox's Sent Items and explicitly acknowledge the duplicate-"
                "delivery risk before approving another attempt."
            )
        retry_after_run_id = latest_run.id

    from app.services.durable_jobs import enqueue_job

    approval_key = automation_idempotency_key(task, from_status)
    if retry_after_run_id is not None:
        # Explicit Review -> In Progress after a terminal attempt is a new
        # approval event. Include the prior immutable attempt id so an exact,
        # confirmed-no-send payload can be retried without inventing an edit.
        approval_key = f"{approval_key}:retry-after:{retry_after_run_id}"
    await enqueue_automation_run(
        db,
        task,
        from_status=from_status,
        actor_user_id=actor_user_id,
        idempotency_key=approval_key,
    )
    await enqueue_job(
        db,
        tenant_id=task.tenant_id,
        kind=TASK_AUTOMATION_JOB,
        # Same key as the run, so a retried approval reuses one job.
        idempotency_key=f"{task.id}:{approval_key}",
        payload={
            "task_id": str(task.id),
            "from_status": from_status,
            "to_status": to_status,
            "actor_user_id": str(actor_user_id) if actor_user_id else None,
            "approval_idempotency_key": approval_key,
        },
    )
    return approval_key


_INTERRUPTED_DELIVERY_DETAIL = (
    "Delivery not confirmed: outcome unknown after an interrupted worker. "
    "Check the sender's Sent Items; this action was not automatically retried."
)


async def _terminalize_interrupted_delivery(
    *,
    task_id,
    tenant_id,
    approval_idempotency_key: str,
    actor_user_id=None,
) -> bool:
    """Make an abandoned ``sending`` claim terminal without another send."""
    async with async_session_maker() as db:
        await set_tenant_context(db, str(tenant_id))
        run = await db.scalar(
            select(TaskAutomationRun)
            .where(
                TaskAutomationRun.task_id == task_id,
                TaskAutomationRun.tenant_id == tenant_id,
                TaskAutomationRun.idempotency_key == approval_idempotency_key,
            )
            .with_for_update()
        )
        if run is None or run.status != "sending":
            return False
        task = await db.scalar(
            select(Task).where(
                Task.id == task_id,
                Task.tenant_id == tenant_id,
            )
        )
        run.status = "failed"
        run.error_message = _INTERRUPTED_DELIVERY_DETAIL
        run.delivery_detail = _INTERRUPTED_DELIVERY_DETAIL
        run.delivery_certainty = DELIVERY_OUTCOME_UNKNOWN
        run.completed_at = datetime.now(timezone.utc)
        if task is not None:
            append_task_event(
                db,
                task,
                event_type="automation_failed",
                actor_user_id=actor_user_id,
                note=_INTERRUPTED_DELIVERY_DETAIL,
                metadata={
                    "action_type": run.action_type,
                    "sent": None,
                    "outcome": "unknown",
                },
            )
        run_id = run.id
        await db.commit()
        logger.warning(
            "task_automation_interrupted_outcome_unknown task_id=%s run_id=%s",
            task_id,
            run_id,
        )
        return True


async def _terminalize_legacy_delivery_jobs(*, task_id, tenant_id) -> int:
    """Fail closed for pre-audit jobs that lack an immutable approval key."""
    async with async_session_maker() as db:
        await set_tenant_context(db, str(tenant_id))
        runs = (
            (
                await db.execute(
                    select(TaskAutomationRun)
                    .where(
                        TaskAutomationRun.task_id == task_id,
                        TaskAutomationRun.tenant_id == tenant_id,
                        TaskAutomationRun.status.in_(("queued", "sending")),
                    )
                    .with_for_update()
                )
            )
            .scalars()
            .all()
        )
        if not runs:
            return 0
        detail = (
            "Delivery not confirmed for a legacy approval created before "
            "immutable action auditing. No automatic retry was attempted; "
            "check Sent Items before explicit reapproval."
        )
        now = datetime.now(timezone.utc)
        for run in runs:
            run.status = "failed"
            run.error_message = detail
            run.delivery_detail = detail
            run.delivery_certainty = DELIVERY_OUTCOME_UNKNOWN
            run.completed_at = now
        await db.commit()
        return len(runs)


async def run_task_automation_job(row) -> dict[str, Any]:
    """Durable-job entry point.

    The durable lease is the only execution path. A retry that finds the prior
    attempt stuck in ``sending`` cannot know whether the provider accepted the
    message before interruption, so it records an uncertain terminal result and
    does not resend. The tenant comes from the job row, never from its payload.
    """
    payload = row.payload or {}
    task_id = payload.get("task_id")
    if not task_id:
        return {"ignored": "missing_task_id"}
    approval_key = (
        str(payload["approval_idempotency_key"])
        if payload.get("approval_idempotency_key")
        else None
    )
    if approval_key is None:
        terminalized = await _terminalize_legacy_delivery_jobs(
            task_id=task_id,
            tenant_id=row.tenant_id,
        )
        logger.warning(
            "task_automation_legacy_job_blocked task_id=%s runs=%s",
            task_id,
            terminalized,
        )
        return {
            "task_id": task_id,
            "delivery": "legacy_outcome_unknown",
            "terminalized_runs": terminalized,
        }
    if (
        int(getattr(row, "attempts", 1) or 1) > 1
        and approval_key
        and await _terminalize_interrupted_delivery(
            task_id=task_id,
            tenant_id=row.tenant_id,
            approval_idempotency_key=approval_key,
            actor_user_id=payload.get("actor_user_id"),
        )
    ):
        return {"task_id": task_id, "delivery": "outcome_unknown"}
    await run_task_automation(
        task_id,
        row.tenant_id,
        from_status=str(payload.get("from_status") or ""),
        to_status=str(payload.get("to_status") or ""),
        actor_user_id=payload.get("actor_user_id"),
        approval_idempotency_key=approval_key,
    )
    return {"task_id": task_id}
