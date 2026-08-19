"""Convert markdown work products to downloadable PDF/DOCX bytes.

Lightweight MVP renderer: handles headings, bold/italic, lists, blockquotes,
and simple tables. Not a full markdown engine — intentionally minimal.
"""

from __future__ import annotations

import io
import re

from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
_INLINE_CODE_RE = re.compile(r"`([^`]+?)`")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBER_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
_QUOTE_RE = re.compile(r"^>\s?(.*)$")
_TABLE_DIVIDER_RE = re.compile(r"^\s*\|?[\s:|-]+\|[\s:|-]+\|?\s*$")


def _inline_html(text: str) -> str:
    """Convert simple inline markdown to reportlab-compatible markup."""
    out = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    out = _BOLD_RE.sub(r"<b>\1</b>", out)
    out = _ITALIC_RE.sub(r"<i>\1</i>", out)
    out = _INLINE_CODE_RE.sub(r'<font face="Courier">\1</font>', out)
    return out


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _markdown_blocks(content: str) -> list[tuple[str, object]]:
    """Parse markdown into (kind, payload) blocks.

    kinds: heading(level, text), para(text), bullets([items]),
           numbers([items]), quote(text), table(rows)
    """
    blocks: list[tuple[str, object]] = []
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if "|" in line and i + 1 < len(lines) and _TABLE_DIVIDER_RE.match(lines[i + 1]):
            rows = [_split_table_row(line)]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(_split_table_row(lines[i]))
                i += 1
            blocks.append(("table", rows))
            continue

        m = _HEADING_RE.match(line)
        if m:
            blocks.append(("heading", (len(m.group(1)), m.group(2).strip())))
            i += 1
            continue

        m = _QUOTE_RE.match(line)
        if m:
            quote_lines = [m.group(1)]
            i += 1
            while i < len(lines):
                qm = _QUOTE_RE.match(lines[i])
                if not qm:
                    break
                quote_lines.append(qm.group(1))
                i += 1
            blocks.append(("quote", "\n".join(quote_lines)))
            continue

        if _BULLET_RE.match(line):
            items = []
            while i < len(lines):
                bm = _BULLET_RE.match(lines[i])
                if not bm:
                    break
                items.append(bm.group(1).strip())
                i += 1
            blocks.append(("bullets", items))
            continue

        if _NUMBER_RE.match(line):
            items = []
            while i < len(lines):
                nm = _NUMBER_RE.match(lines[i])
                if not nm:
                    break
                items.append(nm.group(1).strip())
                i += 1
            blocks.append(("numbers", items))
            continue

        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            _HEADING_RE.match(lines[i])
            or _BULLET_RE.match(lines[i])
            or _NUMBER_RE.match(lines[i])
            or _QUOTE_RE.match(lines[i])
            or ("|" in lines[i] and i + 1 < len(lines) and _TABLE_DIVIDER_RE.match(lines[i + 1]))
        ):
            para_lines.append(lines[i])
            i += 1
        blocks.append(("para", " ".join(para_lines)))

    return blocks


def markdown_to_pdf_bytes(content: str, title: str | None = None) -> bytes:
    """Render markdown to a simple letter-size PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title=title or "Document",
    )
    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    body.fontSize = 10
    body.leading = 14
    quote_style = styles["Italic"]
    quote_style.leftIndent = 18
    quote_style.fontSize = 10

    heading_styles = {
        1: styles["Heading1"],
        2: styles["Heading2"],
        3: styles["Heading3"],
    }

    story = []
    if title:
        story.append(Paragraph(_inline_html(title), styles["Title"]))
        story.append(Spacer(1, 12))

    for kind, payload in _markdown_blocks(content):
        if kind == "heading":
            level, text = payload
            story.append(
                Paragraph(_inline_html(text), heading_styles.get(level, styles["Heading4"]))
            )
            story.append(Spacer(1, 6))
        elif kind == "para":
            story.append(Paragraph(_inline_html(payload), body))
            story.append(Spacer(1, 8))
        elif kind == "quote":
            story.append(Paragraph(_inline_html(payload), quote_style))
            story.append(Spacer(1, 8))
        elif kind == "bullets":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_inline_html(item), body)) for item in payload],
                    bulletType="bullet",
                    start="bulletchar",
                )
            )
            story.append(Spacer(1, 8))
        elif kind == "numbers":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_inline_html(item), body)) for item in payload],
                    bulletType="1",
                )
            )
            story.append(Spacer(1, 8))
        elif kind == "table":
            header = payload[0]
            story.append(Paragraph(_inline_html(" | ".join(header)), styles["Heading4"]))
            for row in payload[1:]:
                story.append(Paragraph(_inline_html(" | ".join(row)), body))
            story.append(Spacer(1, 8))

    doc.build(story)
    return buffer.getvalue()


def markdown_to_docx_bytes(content: str, title: str | None = None) -> bytes:
    """Render markdown to a DOCX document."""
    doc = DocxDocument()
    doc.core_properties.title = title or "Document"

    if title:
        doc.add_heading(title, level=0)

    for kind, payload in _markdown_blocks(content):
        if kind == "heading":
            level, text = payload
            doc.add_heading(text, level=min(level, 4))
        elif kind == "para":
            _add_docx_runs(doc.add_paragraph(), payload)
        elif kind == "quote":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            run = para.add_run(payload)
            run.italic = True
        elif kind == "bullets":
            for item in payload:
                _add_docx_runs(doc.add_paragraph(style="List Bullet"), item)
        elif kind == "numbers":
            for item in payload:
                _add_docx_runs(doc.add_paragraph(style="List Number"), item)
        elif kind == "table":
            rows = payload
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    if c < len(table.rows[r].cells):
                        table.rows[r].cells[c].text = cell

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_docx_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph, honoring **bold** and `code` markers."""
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)
