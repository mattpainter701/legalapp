"""Pure, bounded text revisions for structurally simple DOCX documents.

The revision engine deliberately supports a much smaller surface than Word.  It
accepts only ordinary ``.docx`` packages, exposes deterministic paragraph block
identifiers, and applies exact, block-scoped text replacements.  It never edits
the caller's source bytes and never interprets model output as OOXML.

Callers should treat :func:`inspect_docx` as the capability gate and persist the
returned source digest with any proposed operation.  A later persistence or
delivery layer must still provide optimistic concurrency, immutable releases,
visual rendering, approval, and audit.
"""

from __future__ import annotations

import hashlib
import io
import posixpath
import zipfile
from dataclasses import dataclass
from pathlib import PurePosixPath
from typing import Any, Literal, TypedDict
from xml.etree import ElementTree

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run


CAPABILITY_BOUNDED_TEXT_REVISION = "bounded_docx_text_revision"

MAX_DOCX_BYTES = 25 * 1024 * 1024
MAX_ARCHIVE_ENTRIES = 2_048
MAX_UNCOMPRESSED_BYTES = 125 * 1024 * 1024
MAX_SINGLE_PART_BYTES = 25 * 1024 * 1024
MAX_BLOCKS = 20_000
MAX_TOTAL_BLOCK_TEXT_CHARS = 2_000_000
MAX_OPERATIONS = 50
MAX_TARGET_TEXT_CHARS = 10_000
MAX_REPLACEMENT_TEXT_CHARS = 20_000
MAX_TOTAL_REPLACEMENT_TEXT_CHARS = 100_000

_DOCX_MAIN_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document.main+xml"
)
_OLE_COMPOUND_FILE_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")

_TRACKED_REVISION_ELEMENTS = {
    "cellDel",
    "cellIns",
    "cellMerge",
    "customXmlDelRangeEnd",
    "customXmlDelRangeStart",
    "customXmlInsRangeEnd",
    "customXmlInsRangeStart",
    "del",
    "delText",
    "ins",
    "moveFrom",
    "moveFromRangeEnd",
    "moveFromRangeStart",
    "moveTo",
    "moveToRangeEnd",
    "moveToRangeStart",
    "numberingChange",
    "pPrChange",
    "rPrChange",
    "sectPrChange",
    "tblGridChange",
    "tblPrChange",
    "tblPrExChange",
    "tcPrChange",
    "trPrChange",
}
_PROTECTION_ELEMENTS = {
    "documentProtection",
    "permEnd",
    "permStart",
    "writeProtection",
}
_UNSUPPORTED_STRUCTURED_ELEMENTS = {
    "customXml",
    "sdt",
    "txbxContent",
}
_UNEDITABLE_INLINE_ELEMENTS = {
    "commentReference",
    "drawing",
    "endnoteReference",
    "fldChar",
    "footnoteReference",
    "instrText",
    "object",
    "pict",
    "sym",
}


class ReplaceTextOperation(TypedDict):
    """The only operation accepted by :func:`apply_docx_revision`."""

    type: Literal["replace_text"]
    block_id: str
    target_text: str
    replacement_text: str


class DocumentRevisionError(ValueError):
    """Base error with an API-safe code and capability disposition."""

    def __init__(
        self,
        code: str,
        message: str,
        *,
        capability: str = "unsupported_for_bounded_revision",
    ) -> None:
        super().__init__(message)
        self.code = code
        self.message = message
        self.capability = capability

    def as_dict(self) -> dict[str, str]:
        return {
            "code": self.code,
            "message": self.message,
            "capability": self.capability,
        }


class DocumentCapabilityError(DocumentRevisionError):
    """The source is not safe for this bounded revision engine."""


class DocumentOperationError(DocumentRevisionError):
    """A proposed revision operation is malformed, stale, or unsafe."""

    def __init__(self, code: str, message: str) -> None:
        super().__init__(
            code,
            message,
            capability=CAPABILITY_BOUNDED_TEXT_REVISION,
        )


@dataclass(frozen=True)
class RevisionBlock:
    """One stable, paragraph-shaped unit addressable by an operation."""

    block_id: str
    kind: Literal[
        "body_paragraph",
        "table_cell_paragraph",
        "header_paragraph",
        "footer_paragraph",
    ]
    scope: str
    path: str
    text: str
    editable: bool
    restriction: str | None = None

    def as_dict(self) -> dict[str, Any]:
        return {
            "block_id": self.block_id,
            "kind": self.kind,
            "scope": self.scope,
            "path": self.path,
            "text": self.text,
            "editable": self.editable,
            "restriction": self.restriction,
        }


@dataclass(frozen=True)
class DocumentInspection:
    """Capability and immutable source evidence for a clean DOCX."""

    capability: str
    source_sha256: str
    source_size: int
    blocks: tuple[RevisionBlock, ...]

    @property
    def block_count(self) -> int:
        return len(self.blocks)


@dataclass(frozen=True)
class RevisionPreviewBlock:
    """Before/after text for a changed block, suitable for a review card."""

    block_id: str
    kind: str
    before_text: str
    after_text: str

    def as_dict(self) -> dict[str, str]:
        return {
            "block_id": self.block_id,
            "kind": self.kind,
            "before_text": self.before_text,
            "after_text": self.after_text,
        }


@dataclass(frozen=True)
class DocumentRevisionResult:
    """A newly serialized DOCX plus exact revision evidence."""

    output_bytes: bytes
    source_sha256: str
    output_sha256: str
    source_size: int
    output_size: int
    operation_count: int
    changes: tuple[RevisionPreviewBlock, ...]
    blocks: tuple[RevisionBlock, ...]


@dataclass(frozen=True)
class _ResolvedEdit:
    operation_index: int
    block_id: str
    start: int
    end: int
    target_text: str
    replacement_text: str


@dataclass
class _BlockReference:
    block: RevisionBlock
    paragraph: Paragraph


def inspect_docx(source: bytes, *, filename: str | None = None) -> DocumentInspection:
    """Validate and inspect one clean DOCX without changing ``source``.

    Unsupported or unsafe package capabilities raise
    :class:`DocumentCapabilityError`.  Block identifiers are structural paths;
    they remain stable across ``replace_text`` revisions because that operation
    cannot add, remove, or reorder document blocks.
    """

    source_bytes = _validated_source_bytes(source)
    _validate_filename(filename)
    _validate_docx_package(source_bytes)
    document = _open_document(source_bytes)
    references = _collect_block_references(document)
    return DocumentInspection(
        capability=CAPABILITY_BOUNDED_TEXT_REVISION,
        source_sha256=_sha256(source_bytes),
        source_size=len(source_bytes),
        blocks=tuple(reference.block for reference in references),
    )


def apply_docx_revision(
    source: bytes,
    operations: list[ReplaceTextOperation] | tuple[ReplaceTextOperation, ...],
    *,
    filename: str | None = None,
) -> DocumentRevisionResult:
    """Apply exact block-scoped replacements and return new DOCX bytes.

    Each target must occur exactly once in its named block.  All spans are
    resolved against the same source inspection before any mutation, and
    overlapping spans are rejected.  The caller's immutable ``bytes`` object is
    never modified.
    """

    source_bytes = _validated_source_bytes(source)
    source_sha256 = _sha256(source_bytes)
    _validate_filename(filename)
    _validate_docx_package(source_bytes)
    document = _open_document(source_bytes)
    references = _collect_block_references(document)
    by_id = {reference.block.block_id: reference for reference in references}
    resolved = _validate_and_resolve_operations(operations, by_id)

    expected_text_by_block = {
        block_id: reference.block.text for block_id, reference in by_id.items()
    }
    edits_by_block: dict[str, list[_ResolvedEdit]] = {}
    for edit in resolved:
        edits_by_block.setdefault(edit.block_id, []).append(edit)

    for block_id, edits in edits_by_block.items():
        reference = by_id[block_id]
        expected_text_by_block[block_id] = _text_after_edits(
            reference.block.text,
            edits,
        )
        _apply_edits_to_paragraph(reference.paragraph, edits)
        actual = _paragraph_text(reference.paragraph)
        if actual != expected_text_by_block[block_id]:
            raise DocumentOperationError(
                "in_memory_validation_failed",
                f"Revision of block {block_id!r} did not produce the expected text.",
            )

    output = io.BytesIO()
    try:
        document.save(output)
    except Exception as exc:  # pragma: no cover - python-docx failure boundary
        raise DocumentOperationError(
            "serialization_failed",
            "The revised Word document could not be serialized safely.",
        ) from exc
    output_bytes = output.getvalue()

    # Re-run the complete capability gate and reopen check on serialized bytes.
    _validate_docx_package(output_bytes)
    reopened = _open_document(output_bytes)
    after_references = _collect_block_references(reopened)
    after_by_id = {
        reference.block.block_id: reference for reference in after_references
    }
    if set(after_by_id) != set(by_id):
        raise DocumentOperationError(
            "structure_changed",
            "The DOCX structure changed while applying a bounded text revision.",
        )

    for block_id, expected in expected_text_by_block.items():
        actual = after_by_id[block_id].block.text
        if actual != expected:
            raise DocumentOperationError(
                "result_validation_failed",
                f"Serialized block {block_id!r} did not reopen with the expected text.",
            )

    changes = tuple(
        RevisionPreviewBlock(
            block_id=block_id,
            kind=by_id[block_id].block.kind,
            before_text=by_id[block_id].block.text,
            after_text=expected_text_by_block[block_id],
        )
        for block_id in dict.fromkeys(edit.block_id for edit in resolved)
    )
    return DocumentRevisionResult(
        output_bytes=output_bytes,
        source_sha256=source_sha256,
        output_sha256=_sha256(output_bytes),
        source_size=len(source_bytes),
        output_size=len(output_bytes),
        operation_count=len(resolved),
        changes=changes,
        blocks=tuple(reference.block for reference in after_references),
    )


def _validated_source_bytes(source: bytes) -> bytes:
    if type(source) is not bytes:
        raise DocumentCapabilityError(
            "invalid_source_type",
            "DOCX source content must be an immutable bytes value.",
        )
    if not source:
        raise DocumentCapabilityError("empty_document", "The DOCX source is empty.")
    if len(source) > MAX_DOCX_BYTES:
        raise DocumentCapabilityError(
            "document_too_large",
            f"DOCX source exceeds the {MAX_DOCX_BYTES}-byte revision limit.",
        )
    return source


def _validate_filename(filename: str | None) -> None:
    if filename is None:
        return
    if not isinstance(filename, str) or not filename.strip():
        raise DocumentCapabilityError(
            "invalid_filename",
            "A supplied DOCX filename must be a non-empty string.",
        )
    normalized = filename.replace("\\", "/")
    suffix = PurePosixPath(normalized).suffix.lower()
    if suffix in {".docm", ".dotm"}:
        raise DocumentCapabilityError(
            "macro_enabled_document",
            "Macro-enabled Word documents are not supported for bounded revision.",
            capability="derivative_only",
        )
    if suffix != ".docx":
        raise DocumentCapabilityError(
            "unsupported_extension",
            "Only .docx documents are supported for bounded revision.",
        )


def _validate_docx_package(source: bytes) -> None:
    if source.startswith(_OLE_COMPOUND_FILE_MAGIC):
        raise DocumentCapabilityError(
            "encrypted_or_legacy_word_document",
            "Encrypted or legacy binary Word documents cannot be revised.",
            capability="read_only",
        )
    if not source.startswith(b"PK") or not zipfile.is_zipfile(io.BytesIO(source)):
        raise DocumentCapabilityError(
            "invalid_docx_package",
            "The source is not a valid Open XML DOCX package.",
        )

    try:
        with zipfile.ZipFile(io.BytesIO(source)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_ARCHIVE_ENTRIES:
                raise DocumentCapabilityError(
                    "docx_package_too_complex",
                    "The DOCX package contains too many parts to inspect safely.",
                )
            total_uncompressed = 0
            names: dict[str, str] = {}
            for info in infos:
                normalized = _normalized_archive_name(info.filename)
                lowered = normalized.lower()
                if lowered in names:
                    raise DocumentCapabilityError(
                        "duplicate_package_part",
                        "The DOCX package contains duplicate or case-conflicting parts.",
                    )
                names[lowered] = normalized
                if info.flag_bits & 0x1:
                    raise DocumentCapabilityError(
                        "encrypted_document",
                        "Password-encrypted Word documents cannot be revised.",
                        capability="read_only",
                    )
                if info.file_size < 0 or info.file_size > MAX_SINGLE_PART_BYTES:
                    raise DocumentCapabilityError(
                        "docx_part_too_large",
                        f"DOCX part {normalized!r} exceeds the safe inspection limit.",
                    )
                total_uncompressed += info.file_size
                if total_uncompressed > MAX_UNCOMPRESSED_BYTES:
                    raise DocumentCapabilityError(
                        "docx_package_too_large",
                        "The expanded DOCX package exceeds the safe inspection limit.",
                    )

            required = {"[content_types].xml", "_rels/.rels", "word/document.xml"}
            if not required.issubset(names):
                raise DocumentCapabilityError(
                    "incomplete_docx_package",
                    "The DOCX package is missing required Open XML parts.",
                )
            if "encryptedpackage" in names or "encryptioninfo" in names:
                raise DocumentCapabilityError(
                    "encrypted_document",
                    "Password-encrypted Word documents cannot be revised.",
                    capability="read_only",
                )

            _reject_unsafe_part_names(names)
            _validate_content_types(archive, names["[content_types].xml"])
            _validate_relationships(archive, names)
            _validate_word_xml_parts(archive, names)
    except DocumentRevisionError:
        raise
    except (OSError, zipfile.BadZipFile, RuntimeError) as exc:
        raise DocumentCapabilityError(
            "invalid_docx_package",
            "The DOCX package is damaged or could not be inspected.",
        ) from exc


def _normalized_archive_name(raw_name: str) -> str:
    name = str(raw_name or "").replace("\\", "/")
    normalized = posixpath.normpath(name)
    raw_parts = name.split("/")
    if (
        not name
        or name.startswith("/")
        or any(part in {"", ".", ".."} for part in raw_parts[:-1])
        or normalized in {"", ".", ".."}
        or normalized.startswith("../")
        or ":" in normalized.split("/", 1)[0]
    ):
        raise DocumentCapabilityError(
            "unsafe_package_path",
            "The DOCX package contains an unsafe part path.",
        )
    return normalized


def _reject_unsafe_part_names(names: dict[str, str]) -> None:
    lowered_names = set(names)
    if any(
        name.startswith("word/vbaproject")
        or name.startswith("word/vbadata")
        or name.endswith("/vbaproject.bin")
        for name in lowered_names
    ):
        raise DocumentCapabilityError(
            "macro_enabled_document",
            "Word documents containing VBA or macro parts are not supported.",
            capability="derivative_only",
        )
    if any(name.startswith("_xmlsignatures/") for name in lowered_names):
        raise DocumentCapabilityError(
            "digitally_signed_document",
            "Digitally signed Word documents are read-only in this revision engine.",
            capability="read_only",
        )
    if any(name.startswith("word/activex/") for name in lowered_names):
        raise DocumentCapabilityError(
            "activex_content",
            "Word documents containing ActiveX controls are not supported.",
            capability="derivative_only",
        )
    if any(name.startswith("word/embeddings/") for name in lowered_names):
        raise DocumentCapabilityError(
            "embedded_ole_content",
            "Word documents containing embedded OLE/package objects are not supported.",
            capability="derivative_only",
        )


def _validate_content_types(archive: zipfile.ZipFile, part_name: str) -> None:
    root = _parse_xml_part(archive, part_name)
    content_types = {
        str(element.attrib.get("ContentType") or "").strip().lower()
        for element in root.iter()
        if _local_name(element.tag) in {"Default", "Override"}
    }
    if any("macroenabled" in item or "vbaproject" in item for item in content_types):
        raise DocumentCapabilityError(
            "macro_enabled_document",
            "Macro-enabled Word content types are not supported.",
            capability="derivative_only",
        )
    if any("digital-signature" in item for item in content_types):
        raise DocumentCapabilityError(
            "digitally_signed_document",
            "Digitally signed Word documents are read-only in this revision engine.",
            capability="read_only",
        )
    if _DOCX_MAIN_CONTENT_TYPE not in content_types:
        raise DocumentCapabilityError(
            "unsupported_word_package_type",
            "The package is not a standard, non-macro DOCX document.",
        )


def _validate_relationships(
    archive: zipfile.ZipFile,
    names: dict[str, str],
) -> None:
    for lowered, part_name in names.items():
        if not lowered.endswith(".rels"):
            continue
        root = _parse_xml_part(archive, part_name)
        for relation in root.iter():
            if _local_name(relation.tag) != "Relationship":
                continue
            relation_type = str(relation.attrib.get("Type") or "").lower()
            target_mode = str(relation.attrib.get("TargetMode") or "").lower()
            if target_mode == "external":
                raise DocumentCapabilityError(
                    "external_relationship",
                    "DOCX files with external links or external relationships are not supported.",
                    capability="derivative_only",
                )
            if "vbaproject" in relation_type:
                raise DocumentCapabilityError(
                    "macro_enabled_document",
                    "Word documents containing VBA relationships are not supported.",
                    capability="derivative_only",
                )
            if "digital-signature" in relation_type:
                raise DocumentCapabilityError(
                    "digitally_signed_document",
                    "Digitally signed Word documents are read-only in this revision engine.",
                    capability="read_only",
                )
            if relation_type.endswith("/oleobject") or relation_type.endswith(
                "/package"
            ):
                raise DocumentCapabilityError(
                    "embedded_ole_content",
                    "Word documents containing embedded OLE/package objects are not supported.",
                    capability="derivative_only",
                )
            if relation_type.endswith("/control"):
                raise DocumentCapabilityError(
                    "activex_content",
                    "Word documents containing ActiveX controls are not supported.",
                    capability="derivative_only",
                )
            if relation_type.endswith("/afchunk"):
                raise DocumentCapabilityError(
                    "altchunk_content",
                    "Word documents containing imported altChunk content are not supported.",
                    capability="derivative_only",
                )


def _validate_word_xml_parts(
    archive: zipfile.ZipFile,
    names: dict[str, str],
) -> None:
    for lowered, part_name in names.items():
        if not lowered.startswith("word/") or not lowered.endswith(".xml"):
            continue
        root = _parse_xml_part(archive, part_name)
        for element in root.iter():
            local_name = _local_name(element.tag)
            if local_name in _TRACKED_REVISION_ELEMENTS:
                raise DocumentCapabilityError(
                    "tracked_revisions",
                    "Accept or reject tracked revisions before using bounded DOCX revision.",
                    capability="derivative_only",
                )
            if local_name == "trackRevisions":
                raise DocumentCapabilityError(
                    "tracked_revisions_enabled",
                    "Disable Track Changes before using bounded DOCX revision.",
                    capability="derivative_only",
                )
            if local_name in _PROTECTION_ELEMENTS or local_name == "lock":
                raise DocumentCapabilityError(
                    "protected_document",
                    "Protected or restricted Word content is read-only in this revision engine.",
                    capability="read_only",
                )
            if local_name in _UNSUPPORTED_STRUCTURED_ELEMENTS:
                raise DocumentCapabilityError(
                    "unsupported_structured_content",
                    "Content controls, custom XML, and text boxes require a derivative workflow.",
                    capability="derivative_only",
                )
            if local_name == "altChunk":
                raise DocumentCapabilityError(
                    "altchunk_content",
                    "Imported altChunk content is not supported for bounded revision.",
                    capability="derivative_only",
                )
            if local_name in {"OLEObject", "object"}:
                raise DocumentCapabilityError(
                    "embedded_ole_content",
                    "Embedded OLE objects are not supported for bounded revision.",
                    capability="derivative_only",
                )
            if local_name == "control":
                raise DocumentCapabilityError(
                    "activex_content",
                    "ActiveX controls are not supported for bounded revision.",
                    capability="derivative_only",
                )
            if local_name.lower() == "signatureline":
                raise DocumentCapabilityError(
                    "signature_line_content",
                    "Word signature-line content is read-only in this revision engine.",
                    capability="read_only",
                )


def _parse_xml_part(archive: zipfile.ZipFile, part_name: str) -> ElementTree.Element:
    try:
        data = archive.read(part_name)
        return ElementTree.fromstring(data)
    except (KeyError, ElementTree.ParseError, OSError) as exc:
        raise DocumentCapabilityError(
            "invalid_ooxml",
            f"Open XML part {part_name!r} is missing or malformed.",
        ) from exc


def _local_name(tag: Any) -> str:
    value = str(tag)
    return value.rsplit("}", 1)[-1].rsplit(":", 1)[-1]


def _open_document(source: bytes) -> DocumentObject:
    try:
        return Document(io.BytesIO(source))
    except Exception as exc:
        raise DocumentCapabilityError(
            "invalid_docx_document",
            "The DOCX could not be opened as a Word document.",
        ) from exc


def _collect_block_references(document: DocumentObject) -> list[_BlockReference]:
    references: list[_BlockReference] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        _append_block(
            references,
            paragraph,
            block_id=f"body/p/{paragraph_index}",
            kind="body_paragraph",
            scope="body",
            path=f"p/{paragraph_index}",
        )
    _collect_table_blocks(
        references,
        document.tables,
        scope="body",
        path_prefix="",
        kind="table_cell_paragraph",
    )

    seen_parts: set[str] = set()
    header_footer_proxies: list[tuple[str, Any]] = []
    for section in document.sections:
        header_footer_proxies.extend(
            (
                ("header", section.header),
                ("header", section.first_page_header),
                ("header", section.even_page_header),
                ("footer", section.footer),
                ("footer", section.first_page_footer),
                ("footer", section.even_page_footer),
            )
        )

    for scope_type, proxy in header_footer_proxies:
        # Accessing .part on an undefined first-section header/footer creates a
        # new part.  The private predicate is intentionally checked first so
        # inspection and unrelated body edits cannot alter document structure.
        if not proxy._has_definition:
            continue
        part_name = str(proxy.part.partname).lstrip("/")
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        scope = f"{scope_type}/{part_name}"
        kind = f"{scope_type}_paragraph"
        for paragraph_index, paragraph in enumerate(proxy.paragraphs):
            _append_block(
                references,
                paragraph,
                block_id=f"{scope}/p/{paragraph_index}",
                kind=kind,
                scope=scope,
                path=f"p/{paragraph_index}",
            )
        _collect_table_blocks(
            references,
            proxy.tables,
            scope=scope,
            path_prefix="",
            kind="table_cell_paragraph",
        )

    if len(references) > MAX_BLOCKS:
        raise DocumentCapabilityError(
            "too_many_document_blocks",
            "The DOCX contains too many paragraph blocks for bounded revision.",
        )
    if (
        sum(len(reference.block.text) for reference in references)
        > MAX_TOTAL_BLOCK_TEXT_CHARS
    ):
        raise DocumentCapabilityError(
            "document_text_too_large",
            "The DOCX contains too much text for bounded revision.",
        )
    block_ids = [reference.block.block_id for reference in references]
    if len(block_ids) != len(set(block_ids)):
        raise DocumentCapabilityError(
            "unstable_document_structure",
            "The DOCX structure could not be assigned unique revision blocks.",
        )
    return references


def _collect_table_blocks(
    references: list[_BlockReference],
    tables: Any,
    *,
    scope: str,
    path_prefix: str,
    kind: Literal["table_cell_paragraph"],
) -> None:
    for table_index, table in enumerate(tables):
        table_path = _join_path(path_prefix, f"tbl/{table_index}")
        seen_cells: set[int] = set()
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                marker = id(cell._tc)
                if marker in seen_cells:
                    continue
                seen_cells.add(marker)
                cell_path = _join_path(
                    table_path,
                    f"row/{row_index}/cell/{cell_index}",
                )
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    paragraph_path = _join_path(cell_path, f"p/{paragraph_index}")
                    _append_block(
                        references,
                        paragraph,
                        block_id=f"{scope}/{paragraph_path}",
                        kind=kind,
                        scope=scope,
                        path=paragraph_path,
                    )
                _collect_table_blocks(
                    references,
                    cell.tables,
                    scope=scope,
                    path_prefix=cell_path,
                    kind=kind,
                )


def _join_path(prefix: str, suffix: str) -> str:
    return f"{prefix}/{suffix}" if prefix else suffix


def _append_block(
    references: list[_BlockReference],
    paragraph: Paragraph,
    *,
    block_id: str,
    kind: Any,
    scope: str,
    path: str,
) -> None:
    restriction = _paragraph_restriction(paragraph)
    references.append(
        _BlockReference(
            block=RevisionBlock(
                block_id=block_id,
                kind=kind,
                scope=scope,
                path=path,
                text=_paragraph_text(paragraph),
                editable=restriction is None,
                restriction=restriction,
            ),
            paragraph=paragraph,
        )
    )


def _paragraph_restriction(paragraph: Paragraph) -> str | None:
    for element in paragraph._p.iter():
        if _local_name(element.tag) in _UNEDITABLE_INLINE_ELEMENTS:
            return "complex_inline_content"
    return None


def _paragraph_runs(paragraph: Paragraph) -> list[Run]:
    # Paragraph.runs excludes runs nested in internal hyperlinks.  Descendant
    # w:r traversal keeps visible text and offsets complete while external
    # hyperlink relationships have already been rejected at the package gate.
    return [Run(element, paragraph) for element in paragraph._p.iter(qn("w:r"))]


def _paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in _paragraph_runs(paragraph))


def _validate_and_resolve_operations(
    operations: Any,
    by_id: dict[str, _BlockReference],
) -> list[_ResolvedEdit]:
    if not isinstance(operations, (list, tuple)):
        raise DocumentOperationError(
            "invalid_operations",
            "Revision operations must be a list or tuple of typed dictionaries.",
        )
    if not operations:
        raise DocumentOperationError(
            "empty_operations",
            "At least one replace_text operation is required.",
        )
    if len(operations) > MAX_OPERATIONS:
        raise DocumentOperationError(
            "too_many_operations",
            f"A revision may contain at most {MAX_OPERATIONS} operations.",
        )

    resolved: list[_ResolvedEdit] = []
    total_replacement_chars = 0
    required_keys = {"type", "block_id", "target_text", "replacement_text"}
    for operation_index, operation in enumerate(operations):
        if type(operation) is not dict:
            raise DocumentOperationError(
                "invalid_operation",
                f"Operation {operation_index} must be a plain dictionary.",
            )
        keys = set(operation)
        if keys != required_keys:
            missing = sorted(required_keys - keys)
            unknown = sorted(keys - required_keys)
            details = []
            if missing:
                details.append("missing " + ", ".join(missing))
            if unknown:
                details.append("unknown " + ", ".join(unknown))
            raise DocumentOperationError(
                "invalid_operation_keys",
                f"Operation {operation_index} has invalid keys ({'; '.join(details)}).",
            )
        if type(operation["type"]) is not str:
            raise DocumentOperationError(
                "invalid_operation_value",
                f"Operation {operation_index} field 'type' must be a string.",
            )
        if operation["type"] != "replace_text":
            raise DocumentOperationError(
                "unsupported_operation_type",
                f"Operation {operation_index} type must be 'replace_text'.",
            )
        for field_name in ("block_id", "target_text", "replacement_text"):
            if type(operation[field_name]) is not str:
                raise DocumentOperationError(
                    "invalid_operation_value",
                    f"Operation {operation_index} field {field_name!r} must be a string.",
                )

        block_id = operation["block_id"]
        target_text = operation["target_text"]
        replacement_text = operation["replacement_text"]
        if not block_id:
            raise DocumentOperationError(
                "empty_block_id",
                f"Operation {operation_index} block_id cannot be empty.",
            )
        if not target_text:
            raise DocumentOperationError(
                "empty_target_text",
                f"Operation {operation_index} target_text cannot be empty.",
            )
        if replacement_text == target_text:
            raise DocumentOperationError(
                "no_effect_operation",
                f"Operation {operation_index} replacement must change the target text.",
            )
        if len(target_text) > MAX_TARGET_TEXT_CHARS:
            raise DocumentOperationError(
                "target_text_too_large",
                f"Operation {operation_index} target_text exceeds {MAX_TARGET_TEXT_CHARS} characters.",
            )
        if len(replacement_text) > MAX_REPLACEMENT_TEXT_CHARS:
            raise DocumentOperationError(
                "replacement_text_too_large",
                f"Operation {operation_index} replacement_text exceeds {MAX_REPLACEMENT_TEXT_CHARS} characters.",
            )
        if not _is_valid_xml_text(replacement_text):
            raise DocumentOperationError(
                "invalid_replacement_text",
                f"Operation {operation_index} replacement_text contains characters that are invalid in DOCX XML.",
            )
        total_replacement_chars += len(replacement_text)
        if total_replacement_chars > MAX_TOTAL_REPLACEMENT_TEXT_CHARS:
            raise DocumentOperationError(
                "total_replacement_text_too_large",
                "Combined replacement text exceeds the bounded revision limit.",
            )

        reference = by_id.get(block_id)
        if reference is None:
            raise DocumentOperationError(
                "unknown_block",
                f"Operation {operation_index} references unknown block {block_id!r}.",
            )
        if not reference.block.editable:
            raise DocumentOperationError(
                "block_not_editable",
                f"Block {block_id!r} contains {reference.block.restriction or 'unsupported content'}.",
            )

        starts = _all_occurrences(reference.block.text, target_text)
        if not starts:
            raise DocumentOperationError(
                "target_not_found",
                f"Operation {operation_index} target text is not present in block {block_id!r}.",
            )
        if len(starts) > 1:
            raise DocumentOperationError(
                "ambiguous_target",
                f"Operation {operation_index} target text occurs more than once in block {block_id!r}.",
            )
        start = starts[0]
        resolved.append(
            _ResolvedEdit(
                operation_index=operation_index,
                block_id=block_id,
                start=start,
                end=start + len(target_text),
                target_text=target_text,
                replacement_text=replacement_text,
            )
        )

    edits_by_block: dict[str, list[_ResolvedEdit]] = {}
    for edit in resolved:
        edits_by_block.setdefault(edit.block_id, []).append(edit)
    for block_id, edits in edits_by_block.items():
        ordered = sorted(edits, key=lambda item: (item.start, item.end))
        for previous, current in zip(ordered, ordered[1:]):
            if current.start < previous.end:
                raise DocumentOperationError(
                    "overlapping_operations",
                    f"Operations {previous.operation_index} and {current.operation_index} overlap in block {block_id!r}.",
                )
    return resolved


def _all_occurrences(text: str, target: str) -> list[int]:
    starts: list[int] = []
    cursor = 0
    while cursor <= len(text) - len(target):
        start = text.find(target, cursor)
        if start < 0:
            break
        starts.append(start)
        cursor = start + 1
    return starts


def _text_after_edits(text: str, edits: list[_ResolvedEdit]) -> str:
    result = text
    for edit in sorted(edits, key=lambda item: item.start, reverse=True):
        if result[edit.start : edit.end] != edit.target_text:
            raise DocumentOperationError(
                "stale_operation",
                f"Target text for operation {edit.operation_index} became stale.",
            )
        result = result[: edit.start] + edit.replacement_text + result[edit.end :]
    return result


def _apply_edits_to_paragraph(
    paragraph: Paragraph,
    edits: list[_ResolvedEdit],
) -> None:
    for edit in sorted(edits, key=lambda item: item.start, reverse=True):
        _replace_paragraph_span(
            paragraph,
            start=edit.start,
            end=edit.end,
            expected=edit.target_text,
            replacement=edit.replacement_text,
        )


def _replace_paragraph_span(
    paragraph: Paragraph,
    *,
    start: int,
    end: int,
    expected: str,
    replacement: str,
) -> None:
    runs = _paragraph_runs(paragraph)
    combined = "".join(run.text for run in runs)
    if combined[start:end] != expected:
        raise DocumentOperationError(
            "stale_operation",
            "The paragraph changed before the revision could be applied.",
        )

    offsets: list[tuple[int, int]] = []
    offset = 0
    for run in runs:
        right = offset + len(run.text)
        offsets.append((offset, right))
        offset = right
    affected = [
        index
        for index, (left, right) in enumerate(offsets)
        if right > start and left < end
    ]
    if not affected:
        raise DocumentOperationError(
            "unaddressable_target",
            "The exact target could not be mapped to editable Word runs.",
        )

    first = affected[0]
    last = affected[-1]
    first_left = offsets[first][0]
    last_left = offsets[last][0]
    prefix = runs[first].text[: start - first_left]
    suffix = runs[last].text[end - last_left :]
    if first == last:
        runs[first].text = prefix + replacement + suffix
        return
    runs[first].text = prefix + replacement
    for index in range(first + 1, last):
        runs[index].text = ""
    runs[last].text = suffix


def _sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _is_valid_xml_text(value: str) -> bool:
    return all(
        character in "\t\n\r"
        or 0x20 <= ord(character) <= 0xD7FF
        or 0xE000 <= ord(character) <= 0xFFFD
        or 0x10000 <= ord(character) <= 0x10FFFF
        for character in value
    )


__all__ = [
    "CAPABILITY_BOUNDED_TEXT_REVISION",
    "MAX_OPERATIONS",
    "MAX_REPLACEMENT_TEXT_CHARS",
    "DocumentCapabilityError",
    "DocumentInspection",
    "DocumentOperationError",
    "DocumentRevisionError",
    "DocumentRevisionResult",
    "ReplaceTextOperation",
    "RevisionBlock",
    "RevisionPreviewBlock",
    "apply_docx_revision",
    "inspect_docx",
]
