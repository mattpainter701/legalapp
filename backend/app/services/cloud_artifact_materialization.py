"""Create verified tenant-cloud working copies for generated legal drafts."""

from __future__ import annotations

import hashlib
import io
import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any

from docx import Document
from sqlalchemy import select

from app.database import async_session_maker, set_tenant_context
from app.models.document_storage_operation import DocumentStorageOperation
from app.models.generated_artifact import GeneratedArtifact, GeneratedArtifactRevision
from app.models.matter_document import MatterDocument
from app.models.plugin import Matter
from app.models.task import Task
from app.models.tenant import TenantSettings
from app.services.cloud_docx_snapshot import (
    MAX_DOCX_BYTES,
    CloudDocxSnapshotError,
    inspect_cloud_docx_snapshot,
)
from app.services.document_accountability import (
    append_document_integrity_event,
    ensure_document_storage_operation,
)
from app.services.matter_file_store import MatterFileStore, StorageResult

_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_CLOUD_BACKENDS = frozenset({"onedrive", "sharepoint", "google_drive"})
_UNSAFE_FILENAME = re.compile(r"[^A-Za-z0-9._ -]+")
_WHITESPACE = re.compile(r"\s+")


class CloudArtifactMaterializationError(RuntimeError):
    """A generated artifact could not become a verified cloud working copy."""

    code = "cloud_materialization_failed"
    retryable = False


class CloudNotConfiguredError(CloudArtifactMaterializationError):
    code = "tenant_cloud_not_configured"
    retryable = True


class CloudUploadError(CloudArtifactMaterializationError):
    code = "cloud_write_unconfirmed"
    retryable = False


class CloudIntegrityError(CloudArtifactMaterializationError):
    code = "cloud_integrity_conflict"
    retryable = False


class CloudReconciliationRequired(CloudArtifactMaterializationError):
    code = "cloud_reconciliation_required"
    retryable = False


@dataclass(frozen=True)
class MaterializedArtifact:
    document: MatterDocument
    storage: StorageResult
    operation: DocumentStorageOperation
    sha256: str
    reused: bool = False


@dataclass(frozen=True)
class MatterFolderBinding:
    backend: str
    parent_id: str
    drive_id: str | None


def canonical_docx_filename(
    title: str,
    *,
    revision_no: int = 1,
    artifact_id: uuid.UUID | str | None = None,
) -> str:
    """Return a deterministic collision-resistant provider-safe DOCX name."""
    value = _UNSAFE_FILENAME.sub("-", str(title or "draft"))
    value = _WHITESPACE.sub(" ", value).strip(" .-") or "draft"
    if value.casefold().endswith(".docx"):
        value = value[:-5].rstrip(" .-") or "draft"
    artifact_marker = ""
    if artifact_id is not None:
        marker = re.sub(r"[^0-9a-f]", "", str(artifact_id).casefold())[:12]
        artifact_marker = f"-{marker}" if marker else ""
    suffix = f"{artifact_marker}-r{max(1, int(revision_no))}.docx"
    max_stem = max(1, 220 - len(suffix))
    return f"{value[:max_stem].rstrip(' .-')}{suffix}"


def render_revision_docx(*, title: str, content: str) -> bytes:
    """Render conservative OOXML that Word and LibreOffice can round-trip."""
    document = Document()
    clean_title = str(title or "Draft")
    document.core_properties.title = clean_title
    document.core_properties.subject = "LawHand generated working draft"
    document.add_heading(clean_title, level=1)
    text_value = str(content or "")
    for paragraph in text_value.split("\n"):
        clean = "".join(
            char for char in paragraph if char in "\t\n\r" or ord(char) >= 32
        )
        document.add_paragraph(clean)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _configured_provider(settings: TenantSettings | None) -> str:
    value = (
        str(getattr(settings, "primary_cloud_provider", "") or "")
        .strip()
        .casefold()
        .replace("-", "_")
    )
    aliases = {
        "google": "google_drive",
        "gdrive": "google_drive",
        "one_drive": "onedrive",
        "share_point": "sharepoint",
    }
    value = aliases.get(value, value)
    if value not in _CLOUD_BACKENDS:
        raise CloudNotConfiguredError(
            "Choose and connect a primary tenant cloud provider before generating "
            "matter documents"
        )
    return value


def _matter_folder_binding(
    cloud_folder: dict | None,
    *,
    backend: str,
) -> MatterFolderBinding:
    provider = (cloud_folder or {}).get(backend)
    if not isinstance(provider, dict):
        raise CloudNotConfiguredError(
            "The matter does not have a provisioned folder in the tenant cloud"
        )
    subfolders = provider.get("subfolders")
    if not isinstance(subfolders, dict):
        raise CloudNotConfiguredError(
            "The matter cloud folder does not have a documents subfolder"
        )
    parent_id = str(
        subfolders.get("documents")
        or subfolders.get("uploads")
        or subfolders.get("pleadings")
        or ""
    ).strip()
    drive_id = str(provider.get("drive_id") or "").strip() or None
    if not parent_id or (backend == "sharepoint" and not drive_id):
        raise CloudNotConfiguredError(
            "The matter cloud destination is incomplete and must be reprovisioned"
        )
    return MatterFolderBinding(backend=backend, parent_id=parent_id, drive_id=drive_id)


def _provider_datetime(value: Any) -> datetime | None:
    if isinstance(value, datetime):
        return value if value.tzinfo else value.replace(tzinfo=timezone.utc)
    clean = str(value or "").strip()
    if not clean:
        return None
    try:
        parsed = datetime.fromisoformat(clean.replace("Z", "+00:00"))
    except ValueError:
        return None
    return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)


class CloudArtifactMaterializer:
    """Materialize or verify one artifact revision without hosting its bytes."""

    def __init__(
        self,
        file_store: MatterFileStore | None = None,
        provider_db_factory: Any = None,
    ) -> None:
        self.file_store = file_store or MatterFileStore()
        self.provider_db_factory = provider_db_factory or async_session_maker

    async def read_current_cloud_bytes(
        self,
        *,
        tenant_id: uuid.UUID,
        document: MatterDocument,
    ) -> bytes:
        """Read a mutable working copy without trusting its prior size/hash."""
        async with self.provider_db_factory() as provider_db:
            await set_tenant_context(provider_db, str(tenant_id))
            return await self.file_store.read_matter_file_bytes(
                db=provider_db,
                tenant_id=str(tenant_id),
                document=document,
                expected_sha256=None,
                expected_size=None,
                max_bytes=MAX_DOCX_BYTES,
                enforce_persisted_size=False,
            )

    async def materialize(
        self,
        *,
        db: Any,
        tenant_id: uuid.UUID,
        artifact_id: uuid.UUID,
        revision_id: uuid.UUID,
        task_id: uuid.UUID,
        uploaded_by_user_id: uuid.UUID,
        supersedes_document_id: uuid.UUID | None = None,
        source_docx_bytes: bytes | None = None,
    ) -> MaterializedArtifact:
        artifact = await db.scalar(
            select(GeneratedArtifact)
            .where(
                GeneratedArtifact.tenant_id == tenant_id,
                GeneratedArtifact.id == artifact_id,
            )
            .with_for_update()
        )
        revision = await db.scalar(
            select(GeneratedArtifactRevision).where(
                GeneratedArtifactRevision.tenant_id == tenant_id,
                GeneratedArtifactRevision.artifact_id == artifact_id,
                GeneratedArtifactRevision.id == revision_id,
            )
        )
        if artifact is None or revision is None:
            raise CloudArtifactMaterializationError(
                "Artifact revision is not available in this tenant"
            )
        if artifact.current_revision_no != revision.revision_no:
            raise CloudIntegrityError(
                "Only the artifact's current revision can become the working copy"
            )
        task = await db.scalar(
            select(Task)
            .where(
                Task.id == task_id,
                Task.tenant_id == tenant_id,
                Task.matter_id == artifact.matter_id,
            )
            .with_for_update()
        )
        if task is None or artifact.task_id not in {None, task.id}:
            raise CloudArtifactMaterializationError(
                "Artifact and task are not bound to the same tenant matter"
            )
        matter = await db.scalar(
            select(Matter).where(
                Matter.id == artifact.matter_id,
                Matter.tenant_id == tenant_id,
            )
        )
        settings = await db.scalar(
            select(TenantSettings).where(TenantSettings.tenant_id == tenant_id)
        )
        if matter is None:
            raise CloudArtifactMaterializationError(
                "Artifact matter is not available in this tenant"
            )
        backend = _configured_provider(settings)
        folder = _matter_folder_binding(matter.cloud_folder, backend=backend)

        source_mode = "lawhand_text_renderer"
        if source_docx_bytes is None:
            content = render_revision_docx(
                title=artifact.title,
                content=revision.content_text,
            )
        else:
            try:
                snapshot = inspect_cloud_docx_snapshot(source_docx_bytes)
            except CloudDocxSnapshotError as exc:
                raise CloudIntegrityError(exc.message) from exc
            if snapshot.review_text != revision.content_text:
                raise CloudIntegrityError(
                    "The artifact preview does not match the source DOCX snapshot"
                )
            content = source_docx_bytes
            source_mode = "external_cloud_docx_snapshot"
        digest = hashlib.sha256(content).hexdigest()
        filename = canonical_docx_filename(
            artifact.title,
            revision_no=revision.revision_no,
            artifact_id=artifact.id,
        )
        operation = await ensure_document_storage_operation(
            db,
            tenant_id=tenant_id,
            matter_id=artifact.matter_id,
            task_id=task.id,
            artifact_id=artifact.id,
            artifact_revision_id=revision.id,
            actor_user_id=uploaded_by_user_id,
            content_sha256=digest,
            content_size=len(content),
            target_provider=("google" if backend == "google_drive" else "microsoft"),
            target_backend=backend,
            target_drive_id=folder.drive_id,
            target_parent_id=folder.parent_id,
        )

        existing = await db.scalar(
            select(MatterDocument)
            .where(
                MatterDocument.tenant_id == tenant_id,
                MatterDocument.generated_artifact_id == artifact.id,
                MatterDocument.generated_artifact_revision_id == revision.id,
            )
            .with_for_update()
        )
        if existing is not None:
            return await self._reuse_existing(
                db=db,
                tenant_id=tenant_id,
                artifact=artifact,
                revision=revision,
                operation=operation,
                document=existing,
                expected_sha256=digest,
                task=task,
            )
        if operation.status in {"writing", "provider_accepted", "ambiguous"}:
            raise CloudReconciliationRequired(
                "A prior cloud write may have reached the provider; reconcile it "
                "before retrying"
            )
        if operation.status == "linked":
            raise CloudReconciliationRequired(
                "The cloud operation is linked but its matter document is unavailable"
            )

        operation.status = "writing"
        operation.delivery_certainty = "unknown"
        operation.attempts += 1
        operation.error_code = None
        operation.error_message = None
        await db.flush()

        storage = await self._upload(
            tenant_id=tenant_id,
            matter=matter,
            backend=backend,
            filename=filename,
            content=content,
        )
        if (
            not storage.succeeded
            or storage.backend not in _CLOUD_BACKENDS
            or not storage.provider_item_id
        ):
            operation.status = "ambiguous"
            operation.delivery_certainty = "unknown"
            operation.error_code = "cloud_write_unconfirmed"
            operation.error_message = (
                storage.error or "Cloud provider did not return a durable item"
            )[:1_000]
            await append_document_integrity_event(
                db,
                tenant_id=tenant_id,
                matter_id=artifact.matter_id,
                task_id=task.id,
                artifact_id=artifact.id,
                artifact_revision_id=revision.id,
                operation_id=operation.id,
                event_type="cloud_write_unconfirmed",
                actor_type="service",
                actor_user_id=uploaded_by_user_id,
                content_sha256=digest,
                metadata={
                    "storage_backend": backend,
                    "attempt": operation.attempts,
                },
            )
            raise CloudUploadError(
                "The tenant cloud write could not be confirmed; the draft is blocked "
                "for reconciliation"
            )

        operation.status = "provider_accepted"
        operation.delivery_certainty = "provider_accepted"
        operation.provider_object_id = storage.provider_item_id
        operation.provider_etag = storage.provider_etag
        operation.provider_version_id = storage.provider_version_id
        operation.target_drive_id = storage.drive_id or folder.drive_id
        operation.target_parent_id = storage.parent_id or folder.parent_id
        await db.flush()

        document = MatterDocument(
            tenant_id=tenant_id,
            matter_id=artifact.matter_id,
            task_id=task.id,
            generated_artifact_id=artifact.id,
            generated_artifact_revision_id=revision.id,
            supersedes_document_id=supersedes_document_id,
            uploaded_by_user_id=uploaded_by_user_id,
            filename=filename,
            content_type=_DOCX_CONTENT_TYPE,
            file_size=len(content),
            storage_path=storage.web_url or storage.storage_path,
            storage_provider=storage.provider,
            provider_object_id=storage.provider_item_id,
            provider_drive_id=storage.drive_id or folder.drive_id,
            provider_parent_id=storage.parent_id or folder.parent_id,
            provider_etag=storage.provider_etag,
            provider_version_id=storage.provider_version_id,
            provider_checksum=storage.provider_checksum,
            provider_modified_at=_provider_datetime(storage.provider_modified_at),
            description=(
                "Cloud-edited working copy; human review required"
                if source_docx_bytes is not None
                else "AI-generated working copy; human review required"
            ),
            document_category="generated_draft",
            document_role="working_copy",
            document_status="in_review",
            storage_state="pending",
            document_sha256=digest,
        )
        document.storage_backend = storage.backend

        try:
            async with db.begin_nested():
                db.add(document)
                await db.flush()
                readback = await self._readback(
                    tenant_id=tenant_id,
                    document=document,
                    expected_sha256=digest,
                    expected_size=len(content),
                )
                if not hashlib.sha256(readback).hexdigest() == digest:
                    raise CloudIntegrityError(
                        "Cloud read-back did not match the generated DOCX"
                    )
                document.storage_verified_at = datetime.now(timezone.utc)
                document.storage_state = "verified"
                artifact.task_id = task.id
                artifact.output_document_id = document.id
                operation.status = "linked"
                operation.delivery_certainty = "verified"
                operation.document_id = document.id
                operation.error_code = None
                operation.error_message = None

                if supersedes_document_id is not None:
                    previous = await db.scalar(
                        select(MatterDocument)
                        .where(
                            MatterDocument.id == supersedes_document_id,
                            MatterDocument.tenant_id == tenant_id,
                            MatterDocument.matter_id == artifact.matter_id,
                            MatterDocument.generated_artifact_id == artifact.id,
                        )
                        .with_for_update()
                    )
                    if previous is None:
                        raise CloudIntegrityError(
                            "Superseded document is outside this artifact lineage"
                        )
                    previous.document_status = "superseded"

                await append_document_integrity_event(
                    db,
                    tenant_id=tenant_id,
                    matter_id=artifact.matter_id,
                    task_id=task.id,
                    artifact_id=artifact.id,
                    artifact_revision_id=revision.id,
                    document_id=document.id,
                    operation_id=operation.id,
                    event_type="cloud_working_copy_verified",
                    actor_type="service",
                    actor_user_id=uploaded_by_user_id,
                    content_sha256=digest,
                    provider_object_id=storage.provider_item_id,
                    provider_etag=storage.provider_etag,
                    provider_version_id=storage.provider_version_id,
                    metadata={
                        "storage_backend": storage.backend,
                        "byte_count": len(content),
                        "renderer": revision.renderer_version,
                        "source_mode": source_mode,
                        "revision_no": revision.revision_no,
                        "supersedes_document_id": supersedes_document_id,
                    },
                )
        except Exception as exc:
            compensated = await self._compensate(
                tenant_id=tenant_id,
                storage=storage,
            )
            operation.status = "failed" if compensated else "ambiguous"
            operation.delivery_certainty = "not_delivered" if compensated else "unknown"
            operation.error_code = (
                "cloud_materialization_compensated"
                if compensated
                else "cloud_reconciliation_required"
            )
            operation.error_message = str(exc)[:1_000]
            await append_document_integrity_event(
                db,
                tenant_id=tenant_id,
                matter_id=artifact.matter_id,
                task_id=task.id,
                artifact_id=artifact.id,
                artifact_revision_id=revision.id,
                operation_id=operation.id,
                event_type=(
                    "cloud_write_compensated"
                    if compensated
                    else "cloud_reconciliation_required"
                ),
                actor_type="service",
                actor_user_id=uploaded_by_user_id,
                content_sha256=digest,
                provider_object_id=storage.provider_item_id,
                provider_etag=storage.provider_etag,
                provider_version_id=storage.provider_version_id,
                metadata={
                    "storage_backend": storage.backend,
                    "attempt": operation.attempts,
                },
            )
            if not compensated:
                raise CloudReconciliationRequired(
                    "Cloud storage and workflow persistence diverged; operator "
                    "reconciliation is required"
                ) from exc
            if isinstance(exc, CloudArtifactMaterializationError):
                raise
            raise CloudArtifactMaterializationError(
                "Cloud draft verification failed and the staged object was removed"
            ) from exc

        return MaterializedArtifact(
            document=document,
            storage=storage,
            operation=operation,
            sha256=digest,
        )

    async def _reuse_existing(
        self,
        *,
        db: Any,
        tenant_id: uuid.UUID,
        artifact: GeneratedArtifact,
        revision: GeneratedArtifactRevision,
        operation: DocumentStorageOperation,
        document: MatterDocument,
        expected_sha256: str,
        task: Task,
    ) -> MaterializedArtifact:
        if (
            document.tenant_id != tenant_id
            or document.matter_id != artifact.matter_id
            or document.task_id != task.id
            or document.generated_artifact_id != artifact.id
            or document.generated_artifact_revision_id != revision.id
            or document.storage_backend not in _CLOUD_BACKENDS
            or not document.provider_object_id
            or document.document_sha256 != expected_sha256
            or document.storage_state != "verified"
        ):
            raise CloudIntegrityError(
                "Existing artifact binding is not the verified tenant-matter revision"
            )
        try:
            await self._readback(
                tenant_id=tenant_id,
                document=document,
                expected_sha256=expected_sha256,
                expected_size=document.file_size,
            )
        except Exception as exc:
            document.storage_state = "conflict"
            await append_document_integrity_event(
                db,
                tenant_id=tenant_id,
                matter_id=artifact.matter_id,
                task_id=task.id,
                artifact_id=artifact.id,
                artifact_revision_id=revision.id,
                document_id=document.id,
                operation_id=operation.id,
                event_type="cloud_working_copy_conflict",
                actor_type="service",
                content_sha256=expected_sha256,
                provider_object_id=document.provider_object_id,
                provider_etag=document.provider_etag,
                provider_version_id=document.provider_version_id,
                metadata={"storage_backend": document.storage_backend},
            )
            raise CloudIntegrityError(
                "The cloud working copy changed outside the reviewed revision"
            ) from exc

        operation.status = "linked"
        operation.delivery_certainty = "verified"
        operation.document_id = document.id
        artifact.output_document_id = document.id
        storage = StorageResult(
            provider=document.storage_provider or "",
            backend=document.storage_backend,
            storage_path=document.storage_path,
            web_url=document.cloud_url,
            provider_item_id=document.provider_object_id,
            provider_etag=document.provider_etag,
            provider_version_id=document.provider_version_id,
            provider_modified_at=(
                document.provider_modified_at.isoformat()
                if document.provider_modified_at
                else None
            ),
            provider_checksum=document.provider_checksum,
            drive_id=document.provider_drive_id,
            parent_id=document.provider_parent_id,
        )
        return MaterializedArtifact(
            document=document,
            storage=storage,
            operation=operation,
            sha256=expected_sha256,
            reused=True,
        )

    async def _upload(
        self,
        *,
        tenant_id: uuid.UUID,
        matter: Matter,
        backend: str,
        filename: str,
        content: bytes,
    ) -> StorageResult:
        async with self.provider_db_factory() as provider_db:
            await set_tenant_context(provider_db, str(tenant_id))
            return await self.file_store.store_matter_file_result(
                db=provider_db,
                tenant_id=str(tenant_id),
                matter_slug=matter.slug,
                category="documents",
                filename=filename,
                content=content,
                content_type=_DOCX_CONTENT_TYPE,
                matter_cloud_folder=matter.cloud_folder,
                preferred_provider=backend,
                require_cloud=True,
            )

    async def _readback(
        self,
        *,
        tenant_id: uuid.UUID,
        document: MatterDocument,
        expected_sha256: str,
        expected_size: int | None,
    ) -> bytes:
        async with self.provider_db_factory() as provider_db:
            await set_tenant_context(provider_db, str(tenant_id))
            return await self.file_store.read_matter_file_bytes(
                db=provider_db,
                tenant_id=str(tenant_id),
                document=document,
                expected_sha256=expected_sha256,
                expected_size=expected_size,
            )

    async def _compensate(
        self,
        *,
        tenant_id: uuid.UUID,
        storage: StorageResult,
    ) -> bool:
        if not storage.provider_item_id:
            return False
        try:
            async with self.provider_db_factory() as provider_db:
                await set_tenant_context(provider_db, str(tenant_id))
                await self.file_store.delete_stored_result(
                    db=provider_db,
                    tenant_id=str(tenant_id),
                    result=storage,
                )
            return True
        except Exception:
            return False


cloud_artifact_materializer = CloudArtifactMaterializer()
