import io
from pathlib import Path
from typing import List

import tiktoken


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    Chunk text into overlapping token-based segments using tiktoken.
    Uses the cl100k_base encoding (compatible with text-embedding-3-small).
    """
    if not text or not text.strip():
        return []

    enc = tiktoken.get_encoding("cl100k_base")
    tokens = enc.encode(text)

    if len(tokens) <= chunk_size:
        return [text]

    chunks: List[str] = []
    start = 0

    while start < len(tokens):
        end = min(start + chunk_size, len(tokens))
        chunk_tokens = tokens[start:end]
        chunk_text_decoded = enc.decode(chunk_tokens)
        chunks.append(chunk_text_decoded)

        if end == len(tokens):
            break

        # Move forward by (chunk_size - overlap) tokens
        start += chunk_size - overlap

    return chunks


def extract_text_from_pdf(
    file_bytes: bytes,
    *,
    max_pages: int | None = None,
    max_chars: int | None = None,
) -> str:
    """Extract text from PDF bytes using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    return extract_text_from_pdf_reader(
        reader,
        max_pages=max_pages,
        max_chars=max_chars,
    )


def extract_text_from_pdf_reader(
    reader,
    *,
    max_pages: int | None = None,
    max_chars: int | None = None,
) -> str:
    """Extract bounded text from an existing PDF reader without reparsing it."""

    text_parts = []
    extracted_chars = 0

    for page_num, page in enumerate(reader.pages):
        if max_pages is not None and page_num >= max_pages:
            break
        page_text = page.extract_text()
        if page_text:
            if max_chars is not None:
                remaining = max_chars - extracted_chars
                if remaining <= 0:
                    break
                page_text = page_text[:remaining]
            text_parts.append(page_text)
            extracted_chars += len(page_text)
            if max_chars is not None and extracted_chars >= max_chars:
                break

    extracted = "\n\n".join(text_parts)
    return extracted[:max_chars] if max_chars is not None else extracted


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    return _extract_text_from_docx_document(doc)


def _extract_text_from_docx_document(doc) -> str:
    """Extract body, nested-table, header, and footer text from a DOCX object."""

    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    text_parts: list[str] = []
    seen_paragraphs: set[object] = set()

    def add_paragraphs(paragraphs) -> None:
        for para in paragraphs:
            marker = para._p
            if marker in seen_paragraphs:
                continue
            seen_paragraphs.add(marker)
            if para.text.strip():
                text_parts.append(para.text.strip())

    def add_tables(tables) -> None:
        for table in tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                nonblank = [value for value in values if value]
                if len(values) == 2 and values[0] and values[1]:
                    text_parts.append(f"{values[0].rstrip(':')}: {values[1]}")
                elif nonblank:
                    text_parts.append(" | ".join(nonblank))
                for cell in row.cells:
                    # The row-level representation above is authoritative for
                    # ordinary table cells. Mark their high-level paragraphs
                    # as seen so the low-level text-box/content-control pass
                    # does not append duplicate, decontextualized labels.
                    seen_paragraphs.update(
                        paragraph._p for paragraph in cell.paragraphs
                    )
                    add_tables(cell.tables)

    def add_missing_xml(root, parent) -> None:
        for paragraph_element in root.iter(qn("w:p")):
            marker = paragraph_element
            if marker in seen_paragraphs:
                continue
            paragraph = Paragraph(paragraph_element, parent)
            seen_paragraphs.add(marker)
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

    add_paragraphs(doc.paragraphs)
    add_tables(doc.tables)
    for section in doc.sections:
        for container in (section.header, section.footer):
            add_paragraphs(container.paragraphs)
            add_tables(container.tables)
    add_missing_xml(doc.element.body, doc)
    for section in doc.sections:
        for container in (section.header, section.footer):
            add_missing_xml(container._element, container)

    return "\n\n".join(text_parts)


def _reject_legacy_doc(content_type: str, filename: str) -> None:
    """Refuse legacy .doc for both extractors with one message.

    The two entry points below used to disagree: one raised this, the other
    handed the file to python-docx, which fails on a non-zip container and
    surfaces as an opaque indexing error instead of an actionable one.
    """
    if content_type == "application/msword" or filename.endswith(".doc"):
        raise ValueError(
            "Legacy .doc files are not supported; convert to DOCX, PDF, or TXT."
        )


def extract_text(
    file_bytes: bytes,
    content_type: str,
    filename: str,
    *,
    max_pdf_pages: int | None = None,
    max_pdf_chars: int | None = None,
) -> str:
    """Route to the correct text extractor based on content type or filename."""
    ct_lower = (content_type or "").lower()
    fn_lower = (filename or "").lower()

    _reject_legacy_doc(ct_lower, fn_lower)

    if ct_lower == "application/pdf" or fn_lower.endswith(".pdf"):
        return extract_text_from_pdf(
            file_bytes, max_pages=max_pdf_pages, max_chars=max_pdf_chars
        )

    if ct_lower in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ) or fn_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    if ct_lower.startswith("text/") or fn_lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")

    # Fallback: attempt UTF-8 decode
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception:
        return ""


def extract_text_from_path(
    file_path: str | Path,
    content_type: str,
    filename: str,
    *,
    max_pdf_pages: int | None = None,
    max_pdf_chars: int | None = None,
) -> str:
    """Extract from a staged file without first copying the binary into memory."""

    path = Path(file_path)
    ct_lower = (content_type or "").lower()
    fn_lower = (filename or "").lower()

    _reject_legacy_doc(ct_lower, fn_lower)

    if ct_lower == "application/pdf" or fn_lower.endswith(".pdf"):
        from pypdf import PdfReader

        return extract_text_from_pdf_reader(
            PdfReader(str(path)),
            max_pages=max_pdf_pages,
            max_chars=max_pdf_chars,
        )

    if (
        ct_lower
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or fn_lower.endswith(".docx")
    ):
        from docx import Document

        return _extract_text_from_docx_document(Document(str(path)))

    return path.read_text(encoding="utf-8", errors="replace")
